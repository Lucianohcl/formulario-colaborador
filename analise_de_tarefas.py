# ============================================================
# IMPORTS
# ============================================================

import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from statistics import mean

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from datetime import datetime
import pytz
import time
from zoneinfo import ZoneInfo
import plotly.express as px
import streamlit as st
import json
from datetime import datetime
from github import Github
import time
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px


# ============================================================

# CONFIGURAÇÃO E INICIALIZAÇÃO ÚNICA

# ============================================================

st.set_page_config(

    page_title="Sistema de Análise de Tarefas",

    page_icon="📊",

    layout="wide",

    initial_sidebar_state="expanded"

)


# 1. CONEXÃO GLOBAL (FORA DE QUALQUER IF OU FUNÇÃO)
# Isso garante que 'g' e 'repo' existam em qualquer parte do script
TOKEN = st.secrets["DB_TOKEN"]
g = Github(TOKEN)
repo = g.get_repo("lucianohcl/formulario-colaborador")


# 2. TRAVA DE SEGURANÇA (Vem logo em seguida)
# ============================================================
# st.error("### 🚧 O FORMULÁRIO ENCONTRA-SE INDISPONÍVEL NO MOMENTO.")
# st.stop() 
# ============================================================


# No topo do script, após os imports
if 't' not in locals(): t = None

# Inicialização centralizada

if "logged_in" not in st.session_state: st.session_state.logged_in = False

if "pagina" not in st.session_state:
    st.session_state["pagina"] = "script2"

if "formularios" not in st.session_state: st.session_state["formularios"] = []

       



# Leitura da URL (Prioridade total para permitir acesso ao formulário)

query_params = st.query_params

if "page" in query_params:

    st.session_state.pagina = query_params["page"]

st.markdown("""
    <style>
    /* Oculta a coluna de índice do data_editor */
    div[data-testid="stDataEditor"] > div > div > div > div:first-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================================
# 🛡️ INICIALIZAÇÃO DE VARIÁVEIS (PREVINE ERRO 'NOT DEFINED')
# ============================================================
# Aqui dizemos ao Python que essas variáveis existem, mesmo que vazias.
if "nome_f" not in locals(): nome_f = ""
if "cargo" not in locals(): cargo = ""
if "depto" not in locals(): depto = ""
if "setor" not in locals(): setor = ""
if "chefe" not in locals(): chefe = ""
if "unidade" not in locals(): unidade = ""
if "escolaridade" not in locals(): escolaridade = ""
if "devolver_em" not in locals(): devolver_em = ""
if "cursos" not in locals(): cursos = ""
if "objetivo" not in locals(): objetivo = ""
# Caso seu código ainda procure pelo nome antigo em algum lugar:
nome_digitado = st.session_state.get("usuario_atual", "")
# ============================================================


def gerar_word(form):
    from docx import Document
    import io

    doc = Document()

    c_json = form.get('campos', {})
    t_json = form.get('tabelas', {})

    doc.add_heading(f"Relatório: {form.get('colaborador','Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('timestamp','N/A')}")

    doc.add_heading("Informações de Identificação", level=1)
    for c, valor in c_json.items():
        doc.add_paragraph(f"{c}: {valor}")

    for chave, dados in t_json.items():

        if isinstance(dados, dict):
            dados = list(dados.values())

        doc.add_heading(chave.upper(), level=1)

        if dados:
            cols = list(dados[0].keys())

            table = doc.add_table(rows=1, cols=len(cols))

            for i, col in enumerate(cols):
                table.rows[0].cells[i].text = col

            for item in dados:
                row = table.add_row().cells
                for i, col in enumerate(cols):
                    row[i].text = str(item.get(col, '')).replace('\n', ' ')
        else:
            doc.add_paragraph("Sem dados")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()

def gerar_pdf_html(form):
    campos = form.get("campos", {})
    tabelas = form.get("tabelas", {})
    disc = form.get("disc", {})
    nome = form.get("colaborador", "Colaborador")
    data = form.get("timestamp", "")

    html = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial;
                padding: 20px;
            }}
            h1 {{
                color: #2c3e50;
            }}
            h2 {{
                margin-top: 20px;
                border-bottom: 1px solid #ccc;
                padding-bottom: 5px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 6px;
                font-size: 12px;
            }}
            th {{
                background: #f2f2f2;
            }}
        </style>
    </head>

    <body>
        <h1>Relatório: {nome}</h1>
        <p><b>Data:</b> {data}</p>

        <h2>Campos</h2>
        {"".join([f"<p><b>{k}:</b> {v}</p>" for k, v in campos.items()])}
    """

    # 🔹 TABELAS
    for secao, lista in tabelas.items():
        html += f"<h2>{secao.upper()}</h2>"

        if isinstance(lista, dict):
            lista = list(lista.values())

        if lista:
            html += "<table><tr>"

            cols = list(lista[0].keys())
            for c in cols:
                html += f"<th>{c}</th>"
            html += "</tr>"

            for item in lista:
                html += "<tr>"
                for c in cols:
                    html += f"<td>{item.get(c, '')}</td>"
                html += "</tr>"

            html += "</table>"

    # 🔥 DISC (CORRIGIDO E LIMPO)
    perguntas_disc = [
        "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
        "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
        "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda regras",
        "No trabalho em equipe: (A) Lidera | (B) Motiva | (C) Apoia | (D) Organiza",
        "Em reuniões: (A) Direto ao ponto | (B) Interativo | (C) Escuta | (D) Detalhista",
        "Ao lidar com conflitos: (A) Enfrenta | (B) Apazigua | (C) Evita | (D) Analisa",
        "Seu ritmo de trabalho: (A) Rápido | (B) Entusiasmado | (C) Calmo | (D) Metódico",
        "Prefere tarefas: (A) Desafiadoras | (B) Sociais | (C) Rotina | (D) Técnicas",
        "Seu foco principal: (A) Resultado | (B) Pessoas | (C) Estabilidade | (D) Qualidade",
        "Ao decidir: (A) Rápido | (B) Intuitivo | (C) Cauteloso | (D) Analítico",
        "Confia mais em: (A) Intuição | (B) Opinião | (C) Experiência | (D) Dados",
        "Prefere decisões: (A) Independentes | (B) Grupo | (C) Consenso | (D) Normas",
        "Estilo de organização: (A) Prático | (B) Criativo | (C) Tradicional | (D) Organizado",
        "Lida melhor com: (A) Mudanças | (B) Ideias | (C) Rotina | (D) Regras",
        "Prefere trabalhar: (A) Sozinho | (B) Social | (C) Calmo | (D) Silencioso",
        "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
        "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Analítico",
        "Se motiva por: (A) Poder | (B) Reconhecimento | (C) Segurança | (D) Conhecimento",
        "Reação a cobranças: (A) Esforço | (B) Criatividade | (C) Ansiedade | (D) Técnica",
        "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
        "Feedback: (A) Ajusta | (B) Debate | (C) Planeja | (D) Segue regras",
        "Como aprende: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando",
        "Gestão de tempo: (A) Resultado | (B) Relações | (C) Planejamento | (D) Processo",
        "Como se comunica: (A) Direto | (B) Amigável | (C) Calmo | (D) Técnico"
    ]

    html += "<h2>DISC - Perguntas e Respostas</h2>"

    for i, pergunta in enumerate(perguntas_disc):
        resposta = disc.get(str(i), "NÃO RESPONDIDO")

        html += f"""
        <p>
            <b>{i+1}. {pergunta}</b><br>
            <b>Resposta marcada:</b> {resposta}
        </p>
        """

    html += "</body></html>"

    return html.encode("utf-8")


# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]


# --- FUNÇÕES DE EXPORTAÇÃO (COLE NO TOPO DO SEU ARQUIVO) ---
from docx import Document
from fpdf import FPDF
import io



def extrair_num(texto):
    """Transforma '10 h' ou '5 min' em apenas o número 10 ou 5."""
    try:
        if isinstance(texto, str):
            # Pega apenas os dígitos do texto
            num = "".join(filter(str.isdigit, texto))
            return int(num) if num else 0
        return int(texto)
    except:
        return 0

def limpar_para_rascunho(*args, **kwargs):
    # O (*args, **kwargs) permite que a função receba QUALQUER coisa 
    # e não reclame mais de "arguments".
    st.rerun()

if "respostas_disc" not in st.session_state:
    st.session_state["respostas_disc"] = {}

# ============================================================
# FUNÇÃO DE SUPORTE PARA AS TABELAS (O QUE ESTAVA FALTANDO)
# ============================================================
def preparar_df(chave_rascunho, colunas, fonte_dict):
    """
    Esta função verifica se existe um rascunho carregado.
    Se não, usa os dados oficiais. Se não, cria linhas vazias.
    """
    # 1. Tenta pegar do rascunho carregado no session_state (f_alta_v2, etc)
    # Mapeamento das chaves de rascunho para os nomes das tabelas
    mapa_chaves = {
        "atividades_alta": "f_alta_v2",
        "atividades_normal": "f_normal_v2",
        "atividades_baixa": "f_baixa_v2",
        "dificuldades": "f_dif_v2",
        "sugestoes": "f_sug_v2"
    }
    
    chave_sessao = mapa_chaves.get(chave_rascunho)
    dados_v2 = st.session_state.get(chave_sessao)

    if dados_v2 is not None and len(dados_v2) > 0:
        return pd.DataFrame(dados_v2)

    # 2. Se não tem rascunho, tenta nos dados oficiais (fonte)
    dados_fonte = fonte_dict.get(chave_rascunho, [])
    if dados_fonte:
        # Converte lista simples em DataFrame se necessário
        if isinstance(dados_fonte[0], str):
            return pd.DataFrame([{colunas[0]: item} for item in dados_fonte])
        return pd.DataFrame(dados_fonte)

    # 3. Fallback: Retorna 3 linhas vazias
    return pd.DataFrame([{colunas[0]: ""} for _ in range(3)])

# Configurações de colunas para os editores
col_atv = ["Tarefa"]
col_dif = ["Dificuldade"]
col_sug = ["Sugestão"]
config_col = {"Tarefa": st.column_config.TextColumn("Descrição", width="large")}


# 1. Deixe apenas as variáveis de endereço (sem espaços antes do G)
GITHUB_USER = "lucianohcl"
GITHUB_REPO = "formulario-colaborador"
FOLDER_PATH = "rascunhos" 
GITHUB_TOKEN = st.secrets["DB_TOKEN"]

# 2. Deixe a definição da função (o "def")
# Ela pode ficar aqui em cima, porque o Python só "lê", não "executa"
@st.cache_data(ttl=10)
def atualizar_rascunhos_do_github():
    import requests
    import json
    url = f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/contents/{FOLDER_PATH}"
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}

    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            arquivos = response.json()
            rascunhos_temp = {}
            for arquivo in arquivos:
                if arquivo["name"].endswith(".json"):
                    try:
                        conteudo_res = requests.get(arquivo["download_url"], headers=headers)
                        if conteudo_res.status_code == 200:
                            dados = conteudo_res.json()
                            # ... (mantenha sua lógica de nome aqui dentro) ...
                            nome_colaborador = dados.get("colaborador")
                            if isinstance(nome_colaborador, dict):
                                nome_colaborador = nome_colaborador.get("nome")
                            if not nome_colaborador:
                                nome_colaborador = dados.get("campos", {}).get("nome")
                            
                            if isinstance(nome_colaborador, str):
                                rascunhos_temp[nome_colaborador.strip().upper()] = dados
                    except:
                        continue
            
            if rascunhos_temp:
                st.session_state["rascunhos"] = rascunhos_temp
                return True
        return False
    except:
        return False

# --- O QUE VOCÊ DEVE APAGAR ---
# Se houver uma linha escrita assim: atualizar_rascunhos_do_github() 
# AQUI NA PAREDE ESQUERDA, APAGUE ELA!





def gerar_word(form):
    import io
    from docx import Document

    doc = Document()

    # Título (CORRIGIDO)
    doc.add_heading(f"Relatório: {form.get('colaborador', 'Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('timestamp', 'N/A')}")

    # CAMPOS (CORRETO: dentro de 'campos')
    doc.add_heading("Informações de Identificação", level=1)
    campos = form.get("campos", {})

    for chave, valor in campos.items():
        doc.add_paragraph(f"{chave}: {valor}")

    # TABELAS (CORRETO: dentro de 'tabelas')
    tabelas = form.get("tabelas", {})

    for nome_secao, lista in tabelas.items():

        # 🔥 CORREÇÃO DO FORMATO (dict → list)
        if isinstance(lista, dict):
            lista = list(lista.values())

        doc.add_heading(f"📋 {nome_secao.upper()}", level=1)

        if lista:
            colunas = list(lista[0].keys())

            table = doc.add_table(rows=1, cols=len(colunas))
            table.style = 'Table Grid'

            # Cabeçalho
            for i, col in enumerate(colunas):
                table.rows[0].cells[i].text = col

            # Linhas
            for item in lista:
                row = table.add_row().cells
                for i, col in enumerate(colunas):
                    row[i].text = str(item.get(col, '')).replace('\n', ' ')
        else:
            doc.add_paragraph("Sem dados")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()

def gerar_pdf(form):
    import io
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    # DEBUG
    elementos.append(Paragraph("DEBUG PDF OK", styles['Normal']))

    # Título (AJUSTADO PRO SEU JSON)
    elementos.append(Paragraph(f"Relatório: {form.get('colaborador', 'Colaborador')}", styles['Title']))
    elementos.append(Paragraph(f"Data: {form.get('timestamp', 'N/A')}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    # CAMPOS CORRETOS (SEU JSON TEM 'campos')
    elementos.append(Paragraph("Informações Gerais", styles['Heading2']))
    campos = form.get("campos", {})

    for chave, valor in campos.items():
        elementos.append(Paragraph(f"<b>{chave}:</b> {str(valor).replace('\n','<br/>')}", styles['Normal']))

    elementos.append(Spacer(1, 12))

    # TABELAS CORRETAS (SEU JSON TEM 'tabelas')
    tabelas = form.get("tabelas", {})

    for nome_secao, lista in tabelas.items():

        # 🔥 CORREÇÃO DO SEU FORMATO (vem como dict numerado)
        if isinstance(lista, dict):
            lista = list(lista.values())

        elementos.append(Paragraph(nome_secao.upper(), styles['Heading2']))

        if lista:
            colunas = list(lista[0].keys())
            data = [colunas]

            for item in lista:
                linha = [str(item.get(c, '')).replace('\n', ' ') for c in colunas]
                data.append(linha)

            tabela = Table(data, repeatRows=1)
            tabela.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTSIZE', (0,0), (-1,-1), 8)
            ]))

            elementos.append(tabela)
        else:
            elementos.append(Paragraph("Sem dados", styles['Normal']))

        elementos.append(Spacer(1, 12))

    # 🔥 ÚNICO RETURN — NO FINAL
    doc.build(elementos)

    buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# IMPORTS
# ============================================================

import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from statistics import mean

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from datetime import datetime
import pytz
import time
from zoneinfo import ZoneInfo
import plotly.express as px
# ============================================================

# CONFIGURAÇÃO E INICIALIZAÇÃO ÚNICA

# ============================================================

st.set_page_config(

    page_title="Sistema de Análise de Tarefas",

    page_icon="📊",

    layout="wide",

    initial_sidebar_state="expanded"

)



# Inicialização centralizada

if "logged_in" not in st.session_state: st.session_state.logged_in = False

if "pagina" not in st.session_state: st.session_state.pagina = "home"

if "formularios" not in st.session_state: st.session_state["formularios"] = []



# Leitura da URL (Prioridade total para permitir acesso ao formulário)

query_params = st.query_params

if "page" in query_params:

    st.session_state.pagina = query_params["page"]

st.markdown("""
    <style>
    /* Oculta a coluna de índice do data_editor */
    div[data-testid="stDataEditor"] > div > div > div > div:first-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

# DEFINE O DIRETÓRIO (Isso resolve o problema da função não achar os arquivos)
dados_dir = "dados"
if not os.path.exists(dados_dir):
    os.makedirs(dados_dir)


# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]

# --- FUNÇÕES DE EXPORTAÇÃO (COLE NO TOPO DO SEU ARQUIVO) ---
from docx import Document
from fpdf import FPDF
import io

def salvar_no_github(payload, nome_arquivo, pasta="rascunhos"):
    from github import Github
    import json

    # ✅ Usa o token corretamente do st.secrets
    GITHUB_TOKEN = st.secrets["DB_TOKEN"]
    REPO_NOME = "lucianohcl/formulario-colaborador"

    g = Github(GITHUB_TOKEN)
    repo = g.get_repo(REPO_NOME)

    caminho = f"{pasta}/{nome_arquivo}"  # junta a pasta + nome do arquivo
    try:
        conteudo = json.dumps(payload, indent=4, ensure_ascii=False)
        try:
            file = repo.get_contents(caminho)
            repo.update_file(caminho, f"Atualizando {nome_arquivo}", conteudo, file.sha)
        except:
            repo.create_file(caminho, f"Criando {nome_arquivo}", conteudo)
        return True
    except Exception as e:
        st.error(f"❌ Erro ao salvar no GitHub: {e}")
        return False

def gerar_word(form):
    doc = Document()
    doc.add_heading(f"Relatório: {form.get('Nome', 'Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('DataEnvio', 'N/A')}")
    
    # 1. Informações Gerais
    doc.add_heading("Informações de Identificação", level=1)
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        doc.add_paragraph(f"{campo}: {form.get(campo, 'N/A')}")
    
    # 2. Tabelas (Atividades, Dificuldades, Sugestões)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for chave, colunas in secoes.items():
        if chave in form and isinstance(form[chave], list):
            doc.add_heading(f"📋 {chave}", level=1)
            # Filtra apenas itens que tenham conteúdo real
            dados = [item for item in form[chave] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                table = doc.add_table(rows=1, cols=len(colunas))
                table.style = 'Table Grid'
                # Cabeçalho
                for i, col in enumerate(colunas):
                    table.rows[0].cells[i].text = col
                # Linhas
                for item in dados:
                    row = table.add_row().cells
                    for i, col in enumerate(colunas):
                        row[i].text = str(item.get(col, ''))
            else:
                doc.add_paragraph("Nenhum dado preenchido nesta seção.")

    # 3. Avaliação DISC
    doc.add_heading("📊 Avaliação DISC (Perguntas e Respostas)", level=1)
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        doc.add_paragraph(f"{i}. {pergunta}", style='Heading 2')
        doc.add_paragraph(f"Resposta: {valor_resposta}")
        doc.add_paragraph("-" * 20)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def gerar_pdf(form):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    # Título
    elementos.append(Paragraph(f"Relatório: {form.get('Nome', 'Colaborador')}", styles['Title']))
    elementos.append(Paragraph(f"Data: {form.get('DataEnvio', 'N/A')}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    # Informações Gerais
    elementos.append(Paragraph("Informações Gerais", styles['Heading2']))
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        elementos.append(Paragraph(f"<b>{campo}:</b> {form.get(campo, 'N/A')}", styles['Normal']))
    
    elementos.append(Spacer(1, 12))

    # Tabelas (Atividades, Dificuldades, Sugestoes)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for titulo, colunas in secoes.items():
        if titulo in form and isinstance(form[titulo], list):
            elementos.append(Paragraph(titulo, styles['Heading2']))
            dados = [item for item in form[titulo] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                data = [colunas] # Cabeçalho
                for item in dados:
                    data.append([str(item.get(c, '')) for c in colunas])
                
                tabela = Table(data, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                    ('FONTSIZE', (0,0), (-1,-1), 8)
                ]))
                elementos.append(tabela)
            else:
                elementos.append(Paragraph("Nenhum dado preenchido.", styles['Normal']))
            elementos.append(Spacer(1, 12))

    # DISC
    elementos.append(Paragraph("Avaliação DISC", styles['Heading2']))
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        elementos.append(Paragraph(f"<b>{i}. {pergunta}</b>", styles['Normal']))
        elementos.append(Paragraph(f"Resposta: {valor_resposta}", styles['Italic']))
        elementos.append(Spacer(1, 6))

    doc.build(elementos)
    buffer.seek(0)
    return buffer

# ============================================================
# CALCULAR DISC PERCENTUAL E DOMINANTE
# ============================================================

def calcular_disc(respostas_disc):
    contagem = {"D":0, "I":0, "S":0, "C":0}
    for r in respostas_disc.values():
        if r in contagem:
            contagem[r] += 1
    total = sum(contagem.values())
    if total > 0:
        percentuais = {k: round(v/total*100,1) for k,v in contagem.items()}
        dominante = max(percentuais, key=percentuais.get)
    else:
        percentuais = contagem
        dominante = None
    return percentuais, dominante

# ============================================================
# SCORE DISC PONDERADO
# ============================================================

def score_disc(disc):
    pesos = {"D":1.0,"I":0.9,"S":0.85,"C":0.95}
    total = sum(disc.values())
    if total == 0:
        return 0
    calculo = sum(disc[k]*pesos.get(k,1) for k in disc)
    return round((calculo/total)*100,2)


# ============================================================
# DEFINIÇÃO E CARREGAMENTO DO BANCO DE DADOS (VERSÃO NUVEM)
# ============================================================
import streamlit as st
import os
import json

# --- CAMINHO FIXO PARA O REPOSITÓRIO ---
# 'os.getcwd()' retorna a raiz do projeto onde o arquivo .py está rodando no Streamlit Cloud
base_dir = os.getcwd()
dados_dir = os.path.join(base_dir, "dados")

# Criamos a pasta 'dados' se ela não existir (evita erro de diretório inexistente)
if not os.path.exists(dados_dir):
    os.makedirs(dados_dir, exist_ok=True)

def carregar_todos_formularios(repo_conectado):
    """
    Lê todos os arquivos .json da pasta /dados/ usando a conexão do repo.
    """
    lista_formularios = []
    try:
        # Puxa o conteúdo da pasta /dados/ direto do GitHub
        conteudos = repo_conectado.get_contents("dados")
        
        for item in conteudos:
            if item.path.endswith(".json"):
                try:
                    dados = json.loads(item.decoded_content.decode('utf-8'))
                    if isinstance(dados, dict):
                        lista_formularios.append(dados)
                except Exception:
                    continue
    except Exception as e:
        st.error(f"Erro ao acessar a pasta /dados/ no GitHub: {e}")
        
    return lista_formularios

# ============================================================
# LOGIN (CENTRALIZADO E COM BOTÃO VERMELHO FORÇADO)
# ============================================================
if not st.session_state.logged_in and st.session_state.pagina != "formulario":
    
    # CSS Injetado para forçar a cor e o alinhamento
    st.markdown("""
        <style>
        /* Força a cor do botão vermelho */
        div.stButton > button {
            background-color: #FF4B4B !important;
            color: white !important;
            border-radius: 10px !important;
            height: 3em !important;
            width: 100% !important;
            font-weight: bold !important;
        }
        
        /* Centraliza os inputs de texto */
        .stTextInput > div > div > input {
            text-align: center;
        }
        </style>
    """, unsafe_allow_html=True)

    # Banner de Boas-vindas com a Mensagem Estratégica
    st.markdown("""
        <div style="background: linear-gradient(to right, #4facfe, #00f2fe); padding: 10px 20px; border-radius: 15px; text-align: center; margin-bottom: 20px;">
            <h2 style="color: white; margin: 0;">🚀 Bem-vindo ao Sistema</h2>
            <p style="color: white; font-size: 1.1rem; opacity: 0.9; margin-top: 10px; font-weight: 300;">
                Gestão Estratégica de Análise de Processos e Colaboradores
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Centralização Real usando Colunas Largas
    col1, col2, col3 = st.columns([1, 1.5, 1])

    with col2:
        st.markdown("<h3 style='text-align: center;'>🔐 Acesso</h3>", unsafe_allow_html=True)
        
        usuario = st.text_input("Usuário", placeholder="Digite seu usuário")
        senha = st.text_input("Senha", type="password", placeholder="Digite sua senha")
        
        st.write("") # Espaçador
        
        if st.button("Acessar Painel"):
            if (usuario == "admin" and senha == "admin123") or (usuario == "Luciano" and senha == "123") or (usuario == "JV" and senha == "123"):
                st.session_state.logged_in = True
                st.session_state.user_nome = usuario
                st.session_state.is_admin = True
                st.session_state["usuario_logado"] = "Luciano 123" if usuario == "Luciano" else usuario
                st.rerun()
            else:
                st.error("❌ Usuário ou senha incorretos")
        
    
    st.stop()


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.title("📌 Menu de Navegação")

btn_home = st.sidebar.button("🏠 Home")
btn_analise = st.sidebar.button("📊 Análise Inteligente")
btn_comparar = st.sidebar.button("⚖️ Comparar Colaboradores")
btn_disc = st.sidebar.button("🧠 Perfil DISC")
btn_parecer = st.sidebar.button("📄 Parecer Estratégico")
btn_visualizar = st.sidebar.button("👁️ Visualizar Dados")
btn_produtividade = st.sidebar.button("🚀 Produtividade")


st.sidebar.markdown("---")

btn_logout = st.sidebar.button("🚪 Logout")

pagina_anterior = st.session_state.pagina

if btn_home:
    st.session_state.pagina = "home"
elif btn_analise:
    st.session_state.pagina = "analise"
elif btn_comparar:
    st.session_state.pagina = "comparar"
elif btn_disc:
    st.session_state.pagina = "disc"
elif btn_parecer:
    st.session_state.pagina = "parecer"
elif btn_visualizar:
    st.session_state.pagina = "visualizar"
elif btn_produtividade: 
    st.session_state.pagina = "produtividade"
# O elif abaixo verifica a URL sem precisar de botão
elif st.session_state.pagina == "formulario":
    pass # Este comando é obrigatório para não dar erro de sintaxe
elif btn_logout:
    st.session_state.logged_in = False
    st.session_state.pagina = "home"

if pagina_anterior != st.session_state.pagina:
    st.rerun()


# ============================================================
# EXIBIÇÃO DO CONTEÚDO (PÓS-NAVEGAÇÃO)
# ============================================================

# --- PÁGINA HOME ---
if st.session_state.pagina == "home" or st.session_state.pagina == "script2":
    nome_user = st.session_state.get("user_nome", "Colaborador")
    
    # Banner Principal Chamativo
    st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 40px;
            border-radius: 20px;
            text-align: center;
            color: white;
            box-shadow: 0 15px 35px rgba(0,0,0,0.2);
            margin-bottom: 30px;
        ">
            <h1 style="margin:0; font-size: 3rem; font-weight: 800; border: none;">
                👋 Olá, {nome_user}!
            </h1>
            <p style="font-size: 1.4rem; opacity: 0.9; font-weight: 300; margin-top: 15px;">
                Bem-vindo ao seu Painel Estratégico. 
                Selecione uma opção no menu ao lado para começar.
            </p>
        </div>
    """, unsafe_allow_html=True)

   

# ============================================================
# PÁGINA PERFIL DISC (VERSÃO SINCRONIZADA TOTAL)
# ============================================================
if st.session_state.pagina == "disc":
    import plotly.express as px
    import pandas as pd
    from github import Github

    st.title("🧠 Análise de Perfil DISC")

    # 1. CONEXÃO E CARREGAMENTO (AJUSTADO PARA GARANTIR O REPO)
    try:
        # Pega o repo da sessão
        repo = st.session_state.get('repo_conectado')
        
        # SE O REPO FOR NONE, FORÇA A RECONEXÃO IMEDIATAMENTE
        if repo is None:
            g = Github(st.secrets["DB_TOKEN"])
            repo = g.get_repo("lucianohcl/formulario-colaborador")
            st.session_state.repo_conectado = repo
            
        # Agora sim, chama a lista com a garantia de que o repo existe
        lista_fresca = carregar_todos_formularios(repo)
        
    except Exception as e:
        # Se qualquer coisa falhar acima, tenta uma última vez do zero absoluto
        try:
            g = Github(st.secrets["DB_TOKEN"])
            repo = g.get_repo("lucianohcl/formulario-colaborador")
            st.session_state.repo_conectado = repo
            lista_fresca = carregar_todos_formularios(repo)
        except Exception as e_final:
            st.error(f"❌ Erro crítico de conexão: {e_final}")
            st.stop()

    # 2. VALIDAÇÃO DOS DADOS
    if not lista_fresca:
        st.warning("⚠️ Nenhum formulário encontrado na pasta /dados/.")
        st.stop()
    else:
        st.success(f"📊 {len(lista_fresca)} registros sincronizados para análise.")

    # 3. MAPEAMENTO PARA O SELECTBOX (EXTRAÇÃO DINÂMICA)
    opcoes_colaboradores = {}
    
    for idx, f in enumerate(lista_fresca):
        c = f.get('campos', f)
        nome_bruto = (f.get('colaborador') or 
                      f.get('nome') or 
                      c.get('nome') or 
                      f"Colaborador {idx+1}")
        
        cargo = c.get('cargo', 'Sem Cargo')
        chave_selectbox = f"{str(nome_bruto).upper()} ({cargo})"
        
        opcoes_colaboradores[chave_selectbox] = f

    # 4. INTERFACE DE SELEÇÃO
    colaborador_chave = st.selectbox(
        "🎯 Escolha o colaborador para analisar o perfil:",
        options=list(opcoes_colaboradores.keys())
    )

    # Este é o formulário que será usado para gerar os gráficos abaixo
    formulario_sel = opcoes_colaboradores.get(colaborador_chave)

    # ABAIXO VOCÊ CONTINUA COM O BLOCO DO BOTÃO "GERAR ANÁLISE DISC"
    # ============================================================
    # BOTÃO GERAR ANÁLISE
    # ============================================================

    if formulario_sel and st.button("🔎 Gerar análise DISC"):
        # A partir daqui o seu processamento continua normal
        form = formulario_sel
        
        mapa_disc = {
           "A": "D",
           "B": "I",
           "C": "S",
           "D": "C"
        }
        

        # Extraímos as respostas garantindo que o dicionário 'disc' existe no JSON
        respostas_raw = form.get("disc", {})
        respostas_disc = {}

        for k, v in respostas_raw.items():
            if v in mapa_disc:
                respostas_disc[k] = mapa_disc[v]

        # ============================================================
        # PAINEL DISC DO COLABORADOR (AJUSTADO)
        # ============================================================

        # 1️⃣ Função de Índice de Confirmação (Com detecção de Equilíbrio)
        def score_disc(percentuais):
            if not percentuais:
                return 0
            valores = sorted(percentuais.values(), reverse=True)
            p1 = valores[0]
            p2 = valores[1] if len(valores) > 1 else 0
            
            fator_destaque = max(0, (p1 - 25) / 25)
            fator_distancia = (p1 - p2) / p1 if p1 > 0 else 0
            
            # Ponderação ajustada: 60% destaque, 40% distância
            confirmacao = ((fator_destaque * 0.6) + (fator_distancia * 0.4)) * 100
            score_final = round(confirmacao * 1.6, 1) 
            return max(0, min(score_final, 100))

        # 2️⃣ Cálculos
        percentuais, _ = calcular_disc(respostas_disc)
        ranking = sorted(percentuais.items(), key=lambda x: x[1], reverse=True)
        p1, v1 = ranking[0]
        p2, v2 = ranking[1]

        # Define se é Puro ou Híbrido
        dominante = f"{p1}/{p2}" if (v1 - v2) < 8 else p1
        score = score_disc(percentuais)

        # Prepara os textos ANTES de exibir (Garante que info['nome'] nunca seja N/A)
        letra_busca = p1 
        # Define o dicionário antes para o Python não dar NameError
        textos_disc = {
           "D": {"nome": "Dominante", "estilo": "Resultados e Assertividade", "desc": "Decidido e direto. Busca desafios e rapidez.", "cor": "red", "tarefas": "Tomada de decisão, Gestão de crises, Metas."},
           "I": {"nome": "Influente", "estilo": "Pessoas e Comunicação", "desc": "Entusiasmado e otimista. Busca conexão social.", "cor": "orange", "tarefas": "Apresentações, Networking, Motivação."},
           "S": {"nome": "Estável", "estilo": "Colaboração e Persistência", "desc": "Paciente e leal. Busca harmonia e segurança.", "cor": "green", "tarefas": "Apoio operacional, Suporte, Processos."},
           "C": {"nome": "Conformidade", "estilo": "Precisão e Qualidade", "desc": "Analítico e detalhista. Busca lógica e regras.", "cor": "blue", "tarefas": "Auditoria, Análise de dados, Padronização."}
        }

        # Agora a linha que dava erro vai funcionar:
        letra_busca = dominante.split('/')[0] if '/' in dominante else dominante
        info = textos_disc.get(letra_busca, {"nome": "N/A", "estilo": "N/A", "desc": "", "cor": "gray", "tarefas": ""}).copy()

        # Ajusta o nome para perfis híbridos
        if '/' in dominante:
            letras_hibridas = dominante.split('/')
            nomes_compostos = [textos_disc.get(l, {}).get('nome', l) for l in letras_hibridas]
            info['nome'] = " / ".join(nomes_compostos)
        
        if '/' in dominante:
            nomes_hibridos = [textos_disc.get(l, {}).get('nome', l) for l in dominante.split('/')]
            info['nome'] = " / ".join(nomes_hibridos)

        st.markdown("## 🔹 Painel DISC do Colaborador")

        # 3️⃣ Gráfico e Métricas lado a lado
        col_graf, col_met = st.columns([2,1])

        with col_graf:
            fig = px.bar(
                x=list(percentuais.keys()),
                y=list(percentuais.values()),
                labels={'x':'Tipo','y':'Percentual (%)'},
                text=list(percentuais.values()),
                color=list(percentuais.keys()),
                color_discrete_map={"D":"#FF4136","I":"#FF851B","S":"#2ECC40","C":"#0074D9"}
            )
            fig.update_layout(
                yaxis_range=[0,100], 
                height=350, 
                margin=dict(l=20, r=20, t=30, b=20), 
                template="plotly_white",
                showlegend=False
            )
            st.plotly_chart(fig, use_container_width=True)

        with col_met:
            st.metric("Perfil Dominante", dominante)
            st.metric("Intensidade (Score)", f"{score}%")
            
            # Interpretação com detecção de Perfil Híbrido/Equilibrado
            def interpretar_valor(p, percentuais):
                try:
                    v = float(str(p).replace('%',''))
                    valores = sorted(percentuais.values(), reverse=True)
                    
                    # Se a diferença for pequena, é híbrido. Se for grande (caso do Adson), é Puro.
                    if (valores[0] - valores[1]) < 10:
                        return "⚖️ **Perfil Híbrido / Versátil**"
                    
                    if v > 80: return "🎯 **Intensidade Muito Alta**"
                    if v > 55: return "✅ **Intensidade Alta**"
                    if v > 30: return "⚖️ **Intensidade Moderada**"
                    return "⚠️ **Intensidade Baixa (Indefinido)**"
                except:
                    return ""
            
            st.write(interpretar_valor(score, percentuais))

            st.caption("ℹ️ Este índice indica o nível de precisão e nitidez do perfil identificado. Scores moderados com perfis equilibrados confirmam uma natureza versátil e adaptável do colaborador.")


            st.markdown("---")

            

        # 2. INTERPRETAÇÃO DETALHADA (Substitui a Base de Conhecimento e o Parecer)
        textos_disc = {
           "D": {"nome": "Dominante", "estilo": "Resultados e Assertividade", "desc": "Decidido e direto. Busca desafios e rapidez.", "cor": "red", "tarefas": "Tomada de decisão, Gestão de crises, Metas."},
           "I": {"nome": "Influente", "estilo": "Pessoas e Comunicação", "desc": "Entusiasmado e otimista. Busca conexão social.", "cor": "orange", "tarefas": "Apresentações, Networking, Motivação."},
           "S": {"nome": "Estável", "estilo": "Colaboração e Persistência", "desc": "Paciente e leal. Busca harmonia e segurança.", "cor": "green", "tarefas": "Apoio operacional, Suporte, Processos."},
           "C": {"nome": "Conformidade", "estilo": "Precisão e Qualidade", "desc": "Analítico e detalhista. Busca lógica e regras.", "cor": "blue", "tarefas": "Auditoria, Análise de dados, Padronização."}
        }

        # Identifica a letra principal para busca no dicionário
        letra_busca = dominante.split('/')[0] if '/' in dominante else dominante
        info = textos_disc.get(letra_busca, {"nome": "N/A", "estilo": "", "desc": "", "cor": "gray", "tarefas": ""}).copy()

        # Se for híbrido, monta o nome composto (ex: Dominante / Conformidade)
        if '/' in dominante:
            letras = dominante.split('/')
            nomes = [textos_disc.get(l, {}).get('nome', l) for l in letras]
            info['nome'] = " / ".join(nomes)

        st.markdown(f"### Análise do Perfil: :{info['cor']}[{info['nome']}]")
        st.write(f"**Foco Principal:** {info['estilo']}")
        
        col_desc, col_tar = st.columns(2)
        with col_desc:
            st.info(info['desc'])
        with col_tar:
            st.warning(f"**Tarefas Sugeridas:**\n{info['tarefas']}")

        # 3. LEGENDA DETALHADA (Final da página)
        with st.expander("🔍 Legenda Geral DISC - Detalhada", expanded=False):
            textos_disc_legenda = {
                "D": {
                    "nome": "Dominante",
                    "estilo": "Resultados e Assertividade",
                    "desc": "Decidido e direto. Busca desafios, rapidez e liderança.",
                    "cargos": "Gerente, Líder de Projeto, Coordenador",
                    "tarefas_mais": "Tomada de decisão, Gestão de crises, Definir metas",
                    "tarefas_menos": "Atendimento de rotina, Processos detalhados, Documentação"
                },
                "I": {
                    "nome": "Influente",
                    "estilo": "Pessoas e Comunicação",
                    "desc": "Entusiasmado, sociável e persuasivo. Busca conexão e motivação do grupo.",
                    "cargos": "Marketing, Vendas, Comunicação, Treinamento",
                    "tarefas_mais": "Apresentações, Networking, Reuniões de equipe, Motivação",
                    "tarefas_menos": "Tarefas repetitivas, Processos rígidos, Detalhes técnicos"
                },
                "S": {
                    "nome": "Estável",
                    "estilo": "Colaboração e Persistência",
                    "desc": "Paciente, leal e confiável. Busca harmonia e segurança.",
                    "cargos": "Suporte, Administrativo, RH, Atendimento ao Cliente",
                    "tarefas_mais": "Suporte operacional, Atendimento, Organizar processos",
                    "tarefas_menos": "Mudanças constantes, Pressão por resultados rápidos, Competição intensa"
                },
                "C": {
                    "nome": "Conformidade",
                    "estilo": "Precisão e Qualidade",
                    "desc": "Analítico, detalhista e criterioso. Busca lógica, regras e perfeição.",
                    "cargos": "Auditoria, Contabilidade, TI, Qualidade",
                    "tarefas_mais": "Análise de dados, Relatórios, Controle de qualidade, Padronização",
                    "tarefas_menos": "Decisões rápidas sem dados, Interações sociais constantes, Ambiguidade"
                }
            }

            for key, info_leg in textos_disc_legenda.items():
                st.markdown(f"### **{key} - {info_leg['nome']}**")
                st.write(f"**Estilo de trabalho:** {info_leg['estilo']}")
                st.write(f"**Descrição:** {info_leg['desc']}")
                st.write(f"**Cargos mais compatíveis:** {info_leg['cargos']}")
                st.write(f"**Atividades que combinam mais:** {info_leg['tarefas_mais']}")
                st.write(f"**Atividades que combinam menos:** {info_leg['tarefas_menos']}")
                st.markdown("---")

        
                
        # ============================================================
        # MÓDULO DE INTELIGÊNCIA DE RH: BENCHMARK POR CARGO
        # ============================================================
        # 0. GARANTE A DEFINIÇÃO DA VARIÁVEL (CORREÇÃO DO ERRO)
        c_internos = formulario_sel.get('campos', {})
        cargo_bruto = c_internos.get('cargo') or formulario_sel.get('cargo') or "N/A"
        cargo_limpo = str(cargo_bruto).lower()

        st.markdown(f"### 📑 Consultoria de Perfil: {cargo_bruto.upper()}")

        # 1. BASE DE CONHECIMENTO DE RH (Desejado por Cargo)
        job_benchmarks = {
            "gestor": {
                "perfis": "Híbrido Multidirecional (D-I-C-S)",
                "competencias": [
                    "Liderança e Tomada de Decisão (D)", 
                    "Comunicação, Empatia e Engajamento (I)", 
                    "Visão Estratégica e Planejamento (S)", 
                    "Rigor Técnico, Compliance e Processos (C)"
                ]
            },
            "vendas": {
                "perfis": "Influente (I-D)",
                "competencias": ["Poder de Persuasão (I)", "Foco em Metas e Resultados (D)", "Comunicação Assertiva (I)"]
            },
            "analista": {
                "perfis": "Analítico (C-S)",
                "competencias": ["Atenção Minuciosa a Detalhes (C)", "Organização e Método (C)", "Constância e Foco em Entrega (S)"]
            },
            "rh": {
                "perfis": "Estável (S-I)",
                "competencias": ["Escuta Ativa e Empatia (S)", "Mediação de Conflitos (S)", "Comunicação Interpessoal (I)"]
            }
        }

        # Busca dinâmica: se não achar o cargo exato, tenta por palavra-chave ou usa um padrão Geral
        cargo_key = next((k for k in job_benchmarks if k in cargo_limpo), "gestor")
        benchmark = job_benchmarks[cargo_key]

        # 2. EXPOSIÇÃO DO "GABARITO" DE RH
        st.subheader("🎯 Perfil e Competências Desejadas")
        col_bench1, col_bench2 = st.columns(2)
        
        with col_bench1:
            st.markdown("**Competências Exigidas pela Função:**")
            for comp in benchmark["competencias"]:
                st.markdown(f"• {comp}")
        
        with col_bench2:
            st.info(f"**Perfil Ideal para o Cargo:**\n{benchmark['perfis']}")

        # ============================================================
        # ⚖️ 3. ANÁLISE COMPARATIVA PONDERADA (LÓGICA UNIVERSAL)
        # ============================================================
        st.write("---")
        st.subheader("🔍 Diagnóstico do Colaborador")

        # 1. MOTOR DE PONDERAÇÃO POR AMPLITUDE
        # Diferença entre o maior e o menor score define se é Híbrido ou Especialista
        valores_disc = list(percentuais.values())
        maior_score = max(valores_disc)
        menor_score = min(valores_disc)
        amplitude = maior_score - menor_score 

        # 2. DEFINIÇÃO DA NATUREZA DO PERFIL
        # Se a amplitude for <= 12%, é Equilibrado (Caso Adson). Se > 12%, é Especialista (Caso Pedro).
        is_equilibrado = amplitude <= 12
        perfil_desejado = benchmark["perfis"] 

        # 3. EXIBIÇÃO DO PARECER REFINADO
        if is_equilibrado:
            # FOCO ADSON (Equilíbrio Nativo)
            st.success(f"✅ **Alta Aderência: Perfil Híbrido Multidirecional**")
            st.write(f"""
            **Parecer do Especialista:** Identificamos uma ponderação de percentuais rara e altamente equilibrada 
            (Amplitude de apenas {amplitude:.1f}%). Para o cargo de **{cargo_bruto.upper()}**, este perfil é o ideal, 
            pois a **mitigação de conflitos entre eixos é nativa**. O colaborador possui versatilidade para agir com comando, 
            influência, segurança e análise sem o desgaste mental da adaptação forçada.
            """)
        else:
            # FOCO PEDRO (Especialista Concentrado)
            eixos_dominantes = [letra for letra, valor in percentuais.items() if valor >= 32]
            if not eixos_dominantes: 
                eixos_dominantes = sorted(percentuais, key=percentuais.get, reverse=True)[:2]
            
            perfil_real = "/".join(eixos_dominantes)
            
            if any(p in perfil_desejado for p in eixos_dominantes):
                st.info(f"⚖️ **Aderência Funcional: Perfil Especialista ({perfil_real})**")
                st.write(f"""
                **Parecer do Especialista:** O colaborador possui foco concentrado no eixo **{perfil_real}**. 
                Embora possua a competência técnica exigida para **{cargo_bruto.upper()}**, sua natureza não é híbrida 
                (Amplitude de {amplitude:.1f}%). Ele entregará resultados excelentes nas tarefas de sua dominância, 
                mas poderá sentir fadiga em atividades que exijam os eixos secundários.
                """)
            else:
                st.warning(f"⚠️ **Baixa Aderência: Perfil Desalinhado ({perfil_real})**")
                st.write(f"""
                **Parecer do Especialista:** O perfil dominante (**{perfil_real}**) diverge do benchmark esperado 
                para o cargo de **{cargo_bruto.upper()}**. A alta concentração em um único eixo exigirá um esforço 
                de adaptação elevado para esta função específica.
                """)

        
        # ============================================================
        # 💡 NOTA DO CONSULTOR DINÂMICA (IDENTAÇÃO CORRIGIDA)
        # ============================================================
        st.write("---")
        st.write("👉 **ANÁLISE DE ADAPTAÇÃO ÀS TAREFAS**")

        if is_equilibrado:
            st.info(f"""
            💡 **Nota do Consultor (Perfil Híbrido):** Sua distribuição de energia é equilibrada (Amplitude: {amplitude:.1f}%). 
            A necessidade de adaptação para tarefas de alto rigor técnico ou auditoria é **significativamente mitigada**. 
            Diferente de perfis concentrados, você possui **flexibilidade nativa** para lidar com dados frios e processos 
            rígidos com baixo desgaste mental e sem a fadiga típica de perfis extremos.
            """)
        else:
            # 1. Primeiro identifica o perfil e o desafio (DENTRO do else)
            perfil_primario = max(percentuais, key=percentuais.get)
            
            if perfil_primario in ['I', 'D']:
                desafio = "alto rigor técnico, auditoria e análise de dados frios"
            else:
                desafio = "ritmo acelerado de comunicação, networking e exposição pública"

            # 2. Exibe o aviso usando a variável 'desafio' que acabamos de criar
            st.warning(f"""
            💡 **Nota do Consultor (Perfil Especialista):** Como seu perfil é concentrado no eixo **{perfil_primario}**, 
            tarefas que exigem **{desafio}** demandam um **esforço consciente maior**. 
            Para evitar fadiga, recomenda-se intercalar essas atividades com outras mais naturais ao seu perfil.
            """)

        # GARANTE QUE A VARIÁVEL DOMINANTE EXISTE PARA A BUSCA FUNCIONAR
        if 'dominante' not in locals():
            dominante = max(percentuais, key=percentuais.get)


        # ============================================================
        # 🔍 BUSCA ATIVA: SUGESTÕES E DIFICULDADES (DADOS REAIS)
        # ============================================================
        st.markdown("---")
        # 1. Tenta pegar o nome por 'colaborador' ou 'nome', se não houver, usa 'Colaborador'
        nome_completo = form.get('colaborador') or form.get('nome') or "Colaborador"

        # 2. Pega o primeiro nome com segurança
        primeiro_nome = nome_completo.split()[0] if nome_completo.split() else "Colaborador"

        st.markdown(f"### 🧠 Diagnóstico de Coerência: {primeiro_nome}")
        
        # 1. ACESSA AS TABELAS DIRETAMENTE NO OBJETO BUSCADO
        t_raiz = form.get('tabelas', {})
        sugestoes_lista = t_raiz.get('sugestoes', [])
        dificuldades_lista = t_raiz.get('dificuldades', [])
        
        # Lista de termos para ignorar (limpeza de dados)
        bloqueio_total = ["nenhuma dificuldade", "nenhuma melhoria", "nenhuma", "não tenho", "n/a", "não há", "0", "nenhum", "nada", "", "ok", "n"]

        # 2. MAPEAMENTO DE DORES (Para o sistema usar na busca lógica)
        dores_perfil = {
            "I": ["processo", "planilha", "burocracia", "rotina", "detalhe", "sistema", "repetitivo"],
            "S": ["pressão", "mudança", "conflito", "urgência", "improviso", "rápido"],
            "D": ["lentidão", "burocracia", "espera", "autonomia", "parado", "lento"],
            "C": ["falta de dados", "erro", "improviso", "desorganização", "bagunça", "cliente"]
        }

        col_dif, col_sug = st.columns(2)

        # --- DICIONÁRIO DE BLOQUEIO (REDE LARGA) ---
        bloqueio_check = [
            "nenhuma", "nenhum", "nada", "n", "ñ", "nao", "não", "0", "zero",
            "n/a", "na", "ok", "tudo ok", "tudo certo", "sem sugestões", 
            "sem sugestao", "sem dificuldades", "sem dificuldade",
            "nenhuma melhoria", "nenhuma dificuldade", "não tenho", 
            "não há", "nanhuma", "nemhum", "vazio", "---", ".", ".."
        ]
        
        # Pega o que foi escrito, limpa espaços e joga para minúsculo
        difs_texto = [str(d.get('Dificuldade', '')).lower().strip() for d in dificuldades_lista]
        sugs_texto = [str(s.get('Sugestão', '')).lower().strip() for s in form.get('tabelas', {}).get('sugestoes', [])]

        # Verifica se existe pelo menos UMA resposta que NÃO esteja no dicionário de bloqueio
        tem_algo_dif = any(d for d in difs_texto if d not in bloqueio_check and len(d) > 3)
        tem_algo_sug = any(s for s in sugs_texto if s not in bloqueio_check and len(s) > 3)

        # SE AMBOS FOREM "NADA" -> DISPARA O ALERTA DE RESISTÊNCIA
        if not tem_algo_dif and not tem_algo_sug:
            st.warning("⚠️ **Alerta de Resistência:** O colaborador não reportou dificuldades nem sugeriu melhorias. Isso pode indicar resistência a mudanças, postura defensiva ou proteção de falhas nos processos atuais.")
        
        # Sincroniza o filtro das colunas com o novo dicionário
        bloqueio_total = bloqueio_check
        
        # --- SEU CÓDIGO COMEÇA ABAIXO (NÃO MEXA) ---
        bloqueio_total = bloqueio_check # Garante que as duas listas de filtro sejam iguais

        with col_dif:
            st.subheader("⚠️ Dificuldades vs Perfil")
            # Busca e filtra dificuldades válidas do banco
            difs_validas = [d.get('Dificuldade', '') for d in dificuldades_lista 
                            if str(d.get('Dificuldade', '')).lower() not in bloqueio_total and len(str(d.get('Dificuldade', ''))) > 3]
            
            if difs_validas:
                for dif in difs_validas:
                    st.warning(f"**Relatado:** {dif}")
                    # Usa a variável de perfil (letra_busca ou perfil_primario) para validar
                    letra_v = perfil_primario if 'perfil_primario' in locals() else dominante[0]
                    matches = [w for w in dores_perfil.get(letra_v, []) if w in dif.lower()]
                    
                    if matches:
                        st.success(f"✅ **Coerente:** Gatilhos `{', '.join(matches)}` confirmam o perfil.")
                    else:
                        st.error("⚠️ **Estrutural:** Queixa não ligada ao perfil comportamental.")
            else:
                st.info("Nenhuma dificuldade detectada na busca.")

        with col_sug:
            st.subheader("💡 Sugestões vs Inovação")
            
            # 1. RESGATE DE DADOS (Tabelas)
            t_origem = form.get('tabelas', {})
            sug_banco = t_origem.get('sugestoes', [])
            
            bloqueio = ["nenhuma", "não tenho", "n/a", "não há", "0", "nada", "", "ok", "n", "nenhum", "não"]

            sugestoes_reais = [s.get('Sugestão', '') for s in sug_banco 
                               if str(s.get('Sugestão', '')).lower() not in bloqueio 
                               and len(str(s.get('Sugestão', ''))) > 3]

            # 2. EXIBIÇÃO DAS SUGESTÕES
            if sugestoes_reais:
                for texto in sugestoes_reais:
                    st.info(f"📝 **Sugestão:** {texto}")
                    
                    # Análise de gatilhos para o ícone de foguete
                    gatilhos = ["otimizar", "sistema", "automação", "melhorar", "digitalização", "processo", "kpi", "indicadores"]
                    if any(w in texto.lower() for w in gatilhos):
                        st.success("🚀 **Foco em Eficiência:** Evolução detectada.")
            else:
                st.warning("🚨 Nenhuma sugestão encontrada.")

            

        # ============================================================
        # 💡 SUGESTÃO FINAL PARA O GESTOR (CONTEÚDO DE RH)
        # ============================================================
        st.markdown("---")
        dicionario_gestao = {
            "D": "Dar autonomia e focar em resultados. Evite microgestão.",
            "I": "Promover interação e reconhecimento público. Utilize sua persuasão.",
            "S": "Manter rotinas claras e dar apoio em momentos de mudanças.",
            "C": "Alocar em tarefas de análise, precisão técnica e conformidade."
        }
        
        # Puxa a letra para a sugestão final
        letra_final = perfil_primario if 'perfil_primario' in locals() else dominante[0]
        txt_gestao = dicionario_gestao.get(letra_final, "Acompanhar adaptação.")
        st.info(f"**Dica para o Gestor de {cargo_bruto.upper()}:** {txt_gestao}")

        # ============================================================
        # PERFIL DISC EXIGIDO PELAS ATIVIDADES (VERSÃO INTELIGENTE)
        # ============================================================
        st.markdown("### 🔹 Perfil DISC Exigido pelas Atividades")

        # 1. Extração de todas as atividades das tabelas (Alta, Normal, Baixa)
        tabelas = form.get("tabelas", {})
        todas_atividades = []
        for nivel in ["alta", "normal", "baixa"]:
            for item in tabelas.get(nivel, []):
                todas_atividades.append(item.get("Atividade", "").lower())

        atividades_texto = " ".join(todas_atividades)

        # 2. Dicionário de Palavras-Chave Refinado para RH/DP
        compatibilidade_ativ = {
            "D": ["gerenciar", "liderar", "tomada de decisão", "estratégico", "diretoria", "negociar", "metas", "implementar"],
            "I": ["treinamento", "conduzir", "comunicar", "clientes", "reunião", "atendimento consultivo", "apresentar"],
            "S": ["suporte", "acompanhar", "organizar", "rotina", "processos internos", "apoio", "planejar", "cronograma"],
            "C": ["auditoria", "conferir", "esocial", "legislac", "juridica", "calculos", "analise", "validar", "fiscal", "folha", "inconsistências"]
        }

        # 3. Cálculo de Pontuação
        scores_atividades = {}
        for perfil, palavras in compatibilidade_ativ.items():
            pontos = sum(atividades_texto.count(p) for p in palavras)
            scores_atividades[perfil] = pontos

        # Define os perfis exigidos (Pega os dois maiores para ser Híbrido)
        perfis_ordenados = sorted(scores_atividades.items(), key=lambda x: x[1], reverse=True)
        perfil_exigido_1 = perfis_ordenados[0][0]
        perfil_exigido_2 = perfis_ordenados[1][0] if perfis_ordenados[1][1] > 0 else ""
        
        exigencia_final = f"{perfil_exigido_1}-{perfil_exigido_2}" if perfil_exigido_2 else perfil_exigido_1

        # 4. Cálculo de Compatibilidade Real (Lógica de Ranking)
        # 'dominante' já foi calculado lá em cima com a trava de 15%
        match_real = False
        if dominante == exigencia_final:
            porcentagem_comp = "100%"
            match_real = True
        elif any(letra in dominante for letra in exigencia_final):
            # Se o colaborador é D/C e a tarefa pede C (ou vice-versa)
            porcentagem_comp = "85%"
            match_real = True
        else:
            # Se o colaborador é D puro e a tarefa exige C
            porcentagem_comp = "60%"
            match_real = False

        # 5. Exibição Visual
        col1, col2, col3 = st.columns(3)
        col1.metric("Perfil do Colaborador", dominante)
        col2.metric("Exigência das Tarefas", exigencia_final)
        col3.metric("Compatibilidade", porcentagem_comp)

        if match_real:
            st.success(f"As atividades de **{cargo_bruto.upper()}** exigem um perfil **{exigencia_final}**, o que é compatível com o comportamento do colaborador.")
        else:
            st.warning(f"As atividades atuais puxam muito para o eixo **{exigencia_final}**. O colaborador precisará adaptar seu estilo natural para atender essas demandas técnicas.")

        # ============================================================
        # MENSAGEM PRINCIPAL (VERSÃO INTEGRADA)
        # ============================================================
        # Puxa as descrições do dicionário textos_disc que você já tem no código
        info_dominante = textos_disc.get(dominante, {"nome": "N/A", "estilo": "N/A"})
        
        # Usa o primeiro perfil exigido para buscar a descrição
        info_exigido = textos_disc.get(perfil_exigido_1, {"nome": "N/A", "estilo": "N/A"})

        if match_real:
            # Pega o nome de forma segura, se falhar usa "Colaborador"
            nome_raw = form.get('colaborador', form.get('nome', 'Colaborador'))
            primeiro_nome = nome_raw.split()[0] if nome_raw and nome_raw.split() else "Colaborador"

            st.success(f"✅ **Alta Aderência Inteligente:** O colaborador {primeiro_nome} possui um perfil com foco em resultados.")
        else:
            # 1. Garante que o nome existe antes de tentar o split (12 espaços)
            nome_raw = form.get('nome', 'Colaborador')
            nome_curto = nome_raw.split()[0] if nome_raw and nome_raw.strip() else "Colaborador"

            # 2. A linha corrigida (12 espaços):
            st.warning(f"⚠️ **Ponto de Atenção:** As atividades pedem um foco em **{info_exigido['estilo']}**, mas o colaborador {nome_curto} terá que se esforçar para sair do seu estilo natural de **{info_dominante['estilo']}**.")    


        # ============================================================
        # SUGESTÕES DE GESTÃO (FINAL)
        # ============================================================
        st.markdown("---")
        st.markdown("#### 💡 Sugestão para o Gestor")
        
        dicionario_sugestoes = {
            "D": "Focar em autonomia e entrega de resultados rápidos. Evitar microgestão.",
            "I": "Promover interação com o time e utilizar sua capacidade de convencimento.",
            "S": "Manter rotinas claras e dar suporte em momentos de mudanças bruscas.",
            "C": "Alocar em tarefas que exijam conformidade, análise de dados e precisão técnica."
        }
        
        sugestao_final = dicionario_sugestoes.get(dominante, "Acompanhar adaptação às tarefas.")
        st.info(f"**Dica de Performance:** {sugestao_final}")

        
        # ============================================================
        # ANÁLISE DE ESFORÇO ADAPTATIVO (FOCO NO PERFIL PRINCIPAL)
        # ============================================================
        
        # 0. INICIALIZAÇÃO (Evita NameError)
        atividades_desafio = []
        perfil_principal = dominante if 'dominante' in locals() else "N/A"
        
        # 1. MAPEAMENTO DE OPOSIÇÃO
        opostos = {
            "I": ["C", "D"], "S": ["D", "I"], 
            "D": ["S", "C"], "C": ["I", "D"]
        }
        eixos_desafiadores = opostos.get(perfil_principal, ["C", "D"])

        # 2. BUSCA NAS TABELAS (Garante que 'tabelas' e 'compatibilidade_ativ' existem)
        if 'tabelas' in locals() and 'compatibilidade_ativ' in locals():
            for nivel in ["alta", "normal", "baixa"]:
                for item in tabelas.get(nivel, []):
                    texto_ativ = item.get("Atividade", "")
                    texto_lower = texto_ativ.lower()
                    
                    for eixo in eixos_desafiadores:
                        if any(p in texto_lower for p in compatibilidade_ativ.get(eixo, [])):
                            atividades_desafio.append(texto_ativ)
                            break

        # ============================================================
        # 3. EXIBIÇÃO UNIVERSAL COM EXPLICAÇÃO COMPORTAMENTAL
        # ============================================================
        
        # Garante que as variáveis existam para não quebrar o código
        perfil_dominante = dominante if 'dominante' in locals() else "N/A"
        hibrido_status = is_hibrido if 'is_hibrido' in locals() else False

        if atividades_desafio:
            st.markdown("#### ⚠️ ALGUNS PONTOS DE ATENÇÃO EM RELAÇÃO ÀS TAREFAS DESCRITAS QUE PODEM EXIGIR UM NÍVEL MAIOR DE ADAPTAÇÃO:")
            
            # Remove duplicatas e exibe as 3 principais
            unicas_desafio = list(dict.fromkeys(atividades_desafio))
            for ativ in unicas_desafio[:3]:
                ativ_limpa = str(ativ).replace("\n", " ").strip()
                st.info(f"👉 {ativ_limpa}")

            # 1. DEFINE A RAZÃO DO CONFLITO (Lógica de I/S vs D/C)
            if perfil_dominante in ['I', 'S']:
                razao_conflito = (
                    f"Como seu perfil principal (**{perfil_dominante}**) é orientado a pessoas, ritmo e comunicação, "
                    f"tarefas que exigem alto rigor técnico, auditoria, análise de dados frios e conformidade rígida "
                    f"demandam um esforço consciente maior para manter a precisão e o foco por longos períodos."
                )
            else:
                razao_conflito = (
                    f"Como seu perfil principal (**{perfil_dominante}**) é orientado a resultados rápidos e processos objetivos, "
                    f"tarefas que exigem alta diplomacia, paciência para ritos sociais lentos ou manutenção de rotinas "
                    f"repetitivas podem ser percebidas como menos produtivas, exigindo alta adaptação comportamental."
                )

            # 2. CONSTRUÇÃO DA NOTA FINAL (Com a trava para Perfis Equilibrados)
            if is_equilibrado:
                nota_final = (
                    f"Identificamos que seu perfil é altamente equilibrado (Amplitude: {amplitude:.1f}%). "
                    f"Isso significa que a fadiga em tarefas técnicas ou sociais é mitigada pela sua flexibilidade nativa, "
                    f"permitindo transitar entre diferentes exigências com baixo desgaste mental."
                )
            else:
                nota_final = (
                    f"{razao_conflito} Como seu perfil é mais concentrado, essas tarefas podem gerar fadiga ao longo do dia. "
                    f"Recomenda-se atenção redobrada e a organização da agenda para intercalar essas atividades com outras "
                    f"que sejam mais naturais ao seu perfil."
                )

            # 3. EXIBE A NOTA NA TELA
            st.markdown(f"> **💡 Nota do Consultor:** {nota_final}")

        else:
            st.success("✅ As atividades descritas estão em total harmonia com seu perfil natural.")

        

        
        
        # ============================================================
        # 4. DIAGNÓSTICO INTELIGENTE - ÚNICO E BLINDADO
        # ============================================================
        st.markdown("---")
        
        # 1. RESET E CAPTURA LIMPA
        t_raiz = form.get('tabelas', {})
        
        # Lista de termos evasivos que devem ser ignorados
        bloqueio_total = [
            "nenhuma dificuldade", "nenhuma melhoria", "nenhuma", "não tenho", 
            "n/a", "não há", "0", "nenhum", "dp", "null", "nada", "n / a", ""
        ]

        def filtrar_conteudo(lista, chave):
            validos = []
            for d in lista:
                valor = str(d.get(chave, '')).strip()
                # Só aceita se não estiver no bloqueio e tiver mais de 3 letras
                if valor.lower() not in bloqueio_total and len(valor) > 3:
                    validos.append(valor)
            return " ".join(validos)

        # Gera as strings finais - Se for Adson, ficarão vazias ""
        final_dif = filtrar_conteudo(t_raiz.get('dificuldades', []), 'Dificuldade')
        final_sug = filtrar_conteudo(t_raiz.get('sugestoes', []), 'Sugestão')

        # 2. TRAVA DE EXIBIÇÃO: SE AMBAS VAZIAS -> ALERTA DE RESISTÊNCIA (ADSON)
        if not final_dif and not final_sug:
            st.error(f"🚨 **ALERTA DE RESISTÊNCIA À MUDANÇA (STATUS QUO)**")
            st.markdown(f"A ausência de Sugestões ou Dificuldades por um perfil **{perfil_dominante}** pode indicar resistência passiva ou preenchimento evasivo.")
        
        else:
            pass  # O 'pass' diz ao Python: "não faça nada, mas siga as regras"

        

        # ============================================================
        # 🧠 DIAGNÓSTICO DE ADERÊNCIA AO CARGO (SCORE DE MATCH)
        # ============================================================
        st.markdown("---")
        with st.container(border=True):
            st.subheader("🏆 Score de Alinhamento: Perfil vs. Cargo")

            # 1. CAPTURA DOS DADOS
            dados_pessoais = form.get("campos", {})
            txt_cursos = str(dados_pessoais.get("cursos", "")).lower()
            txt_objetivo = str(dados_pessoais.get("objetivo", "")).lower()
            cargo = str(dados_pessoais.get("cargo", "")).upper()

            # 2. MOTOR DE ANÁLISE INTERNA (SCORE BASEADO EM 3 PILARES)
            score_alinhamento = 0
            analise_interna = []

            # PILAR A: QUALIFICAÇÃO TÉCNICA (Peso 40)
            # Verifica se os cursos dão base para a complexidade do cargo
            if any(x in txt_cursos for x in ["pós", "mba", "especialização", "graduação"]):
                score_alinhamento += 20
                analise_interna.append("Base acadêmica sólida para o nível do cargo.")
            if any(x in txt_cursos for x in ["esocial", "reinf", "legislação", "auditoria", "fiscal"]):
                score_alinhamento += 20
                analise_interna.append("Domínio técnico específico identificado.")

            # PILAR B: CLAREZA E ALINHAMENTO DO OBJETIVO (Peso 40)
            # Verifica se o que ele quer fazer é o que o cargo exige (Gestão/Estratégia)
            if any(x in txt_objetivo for x in ["estratégica", "mitigação", "indicadores", "processos", "liderança"]):
                score_alinhamento += 40
                analise_interna.append("Objetivos estão alinhados com a visão estratégica do cargo.")
            elif len(txt_objetivo) > 100: # Se for longo mas não tiver as palavras-chave acima
                score_alinhamento += 20
                analise_interna.append("Objetivo detalhado, porém com foco muito operacional.")

            # PILAR C: SOBREQUALIFICAÇÃO / SENIORIDADE (Peso 20)
            # Verifica se ele entrega mais do que o "básico"
            if "auditoria" in txt_objetivo and "gestão" in txt_cursos:
                score_alinhamento += 20
                analise_interna.append("Potencial de entrega superior (Perfil Consultivo/Sênior).")

            # 3. EXIBIÇÃO DO SCORE FINAL
            col_score, col_veredito = st.columns([1, 2])
            
            with col_score:
                st.write("**Índice de Aderência**")
                cor_score = "#2ecc71" if score_alinhamento >= 70 else "#f1c40f" if score_alinhamento >= 40 else "#e74c3c"
                st.markdown(f"<h1 style='text-align: center; color: {cor_score};'>{score_alinhamento}%</h1>", unsafe_allow_html=True)
                st.progress(score_alinhamento / 100)

            with col_veredito:
                st.write("**Veredito de Qualificação:**")
                if score_alinhamento >= 80:
                    st.success(f"💎 **ALTO NÍVEL:** O colaborador possui qualificação superior e objetivos perfeitamente alinhados às exigências de {cargo}.")
                elif score_alinhamento >= 50:
                    st.info(f"📊 **ALINHADO:** Perfil técnico adequado. Os objetivos coincidem com as rotinas padrão de {cargo}.")
                else:
                    st.warning(f"⚠️ **DESALINHADO:** A qualificação ou os objetivos declarados estão aquém da complexidade esperada para {cargo}.")

            # 4. LISTA DE JUSTIFICATIVAS DO SCORE
            if analise_interna:
                st.markdown("---")
                st.write("**Justificativa do Score:**")
                for item in analise_interna:
                    st.markdown(f"🔹 {item}")

        # ============================================================
        # 📥 LAUDO PERICIAL MASTER (VERSÃO REVISADA - PONDERAÇÃO GERCINO)
        # ============================================================

        # 0. PROCESSAMENTO DO GRÁFICO (CONVERSÃO PARA HTML)
        grafico_html_div = "" 
        try:
            import plotly.io as pio
            if 'fig' in locals():
                grafico_html_div = pio.to_html(fig, full_html=False, include_plotlyjs='cdn')
        except Exception:
            grafico_html_div = "<p style='text-align:center; color:gray;'>Gráfico Indisponível</p>"

        # 1. TRATAMENTO DE VARIÁVEIS E BUSCA NO JSON
        lista_sugestoes = [s.get("Sugestão", "") for s in tabelas.get("sugestoes", []) if s.get("Sugestão")]
        lista_dificuldades = [d.get("Dificuldade", "") for d in tabelas.get("dificuldades", []) if d.get("Dificuldade")]
        
        # 2. NOVA PONDERAÇÃO TÉCNICA (NEXO CAUSAL REVISADO)
        # Amplitude alta (>50%) = Perfil Especialista/Pico (Menos Equilíbrio, Mais Foco)
        if amplitude > 50:
            status_perfil = "Especialista de Alto Impacto"
            diagnostico_fadiga = f"A amplitude de {amplitude:.1f}% indica um perfil com picos comportamentais definidos. Ao contrário de perfis equilibrados, GERCINO possui um 'trilho' de atuação muito claro, o que gera <b>fadiga severa</b> quando exposto a tarefas multifuncionais ou que fujam de sua especialidade técnica."
        else:
            status_perfil = "Generalista Versátil"
            diagnostico_fadiga = f"A amplitude de {amplitude:.1f}% indica um perfil equilibrado, onde a flexibilidade nativa permite a transição entre tarefas técnicas e sociais com menor desgaste funcional."

        # 3. CONSTRUÇÃO DOS BLOCOS DE TEXTO
        alerta_resistencia = ""
        if not lista_sugestoes and not lista_dificuldades:
            alerta_resistencia = f"""
            <div style='background: #fff5f5; border-left: 6px solid #e74c3c; padding: 20px; border-radius: 8px; margin-top: 20px;'>
                <b style='color: #c0392b;'>🚨 ALERTA DE RESISTÊNCIA À MUDANÇA:</b><br>
                A ausência de reportes de dificuldades sugere uma postura de <b>autopreservação</b>. Em perfis especialistas, isso pode mascarar gargalos que levam ao burnout técnico.
            </div>
            """

        nota_consultor = f"""
        <div style='background: #f8f9fa; border: 1px solid #e9ecef; padding: 20px; border-radius: 8px; margin-top: 20px; font-style: italic; border-left: 5px solid #1B1E5D;'>
            <b>💡 Nota do Consultor:</b> Identificamos que o perfil de {primeiro_nome} é <b>{status_perfil}</b>. 
            {diagnostico_fadiga}
        </div>
        """

        # 4. MONTAGEM DA STRING HTML
        html_final_estendido = f"""
        <!DOCTYPE html>
        <html lang='pt-br'>
        <head>
            <meta charset='utf-8'>
            <style>
                body {{ font-family: 'Segoe UI', Arial, sans-serif; padding: 40px; color: #2c3e50; line-height: 1.6; background: #eceff1; }}
                .page {{ background: white; padding: 50px; border-radius: 15px; max-width: 1000px; margin: auto; box-shadow: 0 10px 30px rgba(0,0,0,0.1); border-top: 15px solid #1B1E5D; }}
                .header {{ text-align: center; border-bottom: 3px solid #f1f1f1; padding-bottom: 25px; margin-bottom: 35px; }}
                .header h1 {{ margin: 0; color: #1B1E5D; font-size: 28px; text-transform: uppercase; }}
                .grid-info {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px; margin-bottom: 35px; }}
                .stat-card {{ background: #fdfdfd; padding: 15px; border-radius: 10px; border: 1px solid #e0e6ed; text-align: center; border-bottom: 4px solid #1B1E5D; }}
                .stat-card label {{ font-size: 10px; text-transform: uppercase; color: #95a5a6; font-weight: 800; display: block; }}
                .section-title {{ background: #1B1E5D; color: white; padding: 12px 20px; border-radius: 8px; margin-top: 40px; font-size: 14px; font-weight: bold; text-transform: uppercase; }}
                .chart-box {{ border: 1px solid #eee; margin: 20px 0; border-radius: 12px; padding: 15px; background: #fff; text-align: center; }}
                .parecer-box {{ background: #f0f7ff; border-left: 6px solid #1B1E5D; padding: 25px; border-radius: 8px; margin: 20px 0; }}
                .footer {{ margin-top: 60px; text-align: center; font-size: 11px; color: #bdc3c7; border-top: 1px solid #eee; padding-top: 25px; }}
            </style>
        </head>
        <body>
            <div class='page'>
                <div class='header'>
                    <h1>LAUDO PERICIAL COMPORTAMENTAL 360°</h1>
                    <p>AUDITORIA TÉCNICA E ANÁLISE DE NEXO CAUSAL</p>
                </div>

                <div class='grid-info'>
                    <div class='stat-card'><label>Colaborador</label><b>{primeiro_nome}</b></div>
                    <div class='stat-card'><label>Cargo Atual</label><b>{cargo_bruto.upper()}</b></div>
                    <div class='stat-card'><label>Perfil Resultante</label><b>{dominante}</b></div>
                    <div class='stat-card'><label>Índice Fit</label><b style='color:#27ae60'>{porcentagem_comp}</b></div>
                </div>

                <div class='section-title'>1. ANÁLISE QUANTITATIVA (GRÁFICO DISC)</div>
                <div class='chart-box'>{grafico_html_div}</div>

                <div class='section-title'>2. DIAGNÓSTICO DE COERÊNCIA E ADAPTAÇÃO</div>
                <div class='parecer-box'>
                    <h4 style='margin-top:0;'>Parecer do Especialista:</h4>
                    {info.get('desc', 'Análise técnica em processamento.')}
                    <br><br>
                    <b>Veredito:</b> {primeiro_nome} possui as competências críticas para a cadeira atual, exigindo apenas monitoramento de carga cognitiva.
                </div>

                {nota_consultor}
                {alerta_resistencia}

                <div class='section-title'>3. PONTOS DE ATENÇÃO (NEXO CAUSAL)</div>
                <div style='margin-top: 20px; font-size: 14px;'>
                    <p>⚠️ <b>Tarefas de Alto Risco de Esgotamento para este Perfil:</b></p>
                    <ul>
                        <li>Intervenções sociais não planejadas.</li>
                        <li>Gestão multifocal de processos sem POP definido.</li>
                        <li>Demandas que exijam alta flexibilidade comportamental imediata.</li>
                    </ul>
                </div>

                <div class='footer'>
                    <b>GERADO POR NETEXAME AUDITORIA ESTRATÉGICA - 2026</b>
                </div>
            </div>
        </body>
        </html>
        """

        # 5. RENDERIZAÇÃO DO BOTÃO
        st.download_button(
            label="📥 BAIXAR LAUDO PERICIAL COMPLETO",
            data=html_final_estendido,
            file_name=f"LAUDO_PERICIAL_{primeiro_nome}.html",
            mime="text/html",
            key="btn_laudo_final_deploy",
            use_container_width=True
        )             
        

# --- VISUALIZAÇÃO ---
if st.session_state.get("pagina") == "visualizar":
    st.title("👁️ Visualização de Registros")
    
    # 1. CONEXÃO E CARREGAMENTO
    try:
        lista_de_arquivos = carregar_todos_formularios(repo)
    except NameError:
        repo = g.get_repo("lucianohcl/formulario-colaborador")
        lista_de_arquivos = carregar_todos_formularios(repo)
    
    if not lista_de_arquivos:
        st.warning("⚠️ Nenhum formulário encontrado na pasta /dados/.")
    else:
        st.success(f"Foram encontrados {len(lista_de_arquivos)} formulários.")
        st.session_state['base_auditoria'] = lista_de_arquivos
        # ESTA LINHA SALVA A VARIÁVEL T NA MEMÓRIA PARA A AUDITORIA NÃO DAR ERRO
        st.session_state['t'] = lista_de_arquivos

        # 2. INICIALIZAÇÃO DA LISTA DE OCULTOS
        if "arquivos_escondidos" not in st.session_state:
            st.session_state["arquivos_escondidos"] = []

        # 3. LOOP DE EXIBIÇÃO COM FILTRO VIRTUAL
        for idx, form in enumerate(lista_de_arquivos, 1):
            # Identificador único para ocultar (usa o timestamp ou index)
            id_atual = form.get('timestamp') or f"form_{idx}"
            
            # --- FILTRO VIRTUAL: PULA SE ESTIVER NA LISTA ---
            if id_atual in st.session_state.get("arquivos_escondidos", []):
                continue

            # Extração de Identificação
            nome_extraido = (form.get('colaborador') or 
                             form.get('nome') or 
                             form.get('campos', {}).get('nome') or 
                             f'Colaborador {idx}')
            
            nome_exibir = str(nome_extraido).upper()
            c = form.get('campos', form) 
            
            with st.expander(f"👤 FORMULÁRIO DE: {nome_exibir}", expanded=False):
                
                # --- CABEÇALHO ---
                st.subheader("📝 Informações de Identificação")
                col1, col2 = st.columns(2)
                col1.write(f"**Data de Envio:** {form.get('timestamp') or form.get('data_envio', 'N/A')}")
                col2.write(f"**Devolver em:** {c.get('devolver_em') or c.get('devolucao', 'N/A')}")
                
                col_a, col_b = st.columns(2)
                col_a.write(f"**Setor:** {c.get('setor', 'N/A')}")
                col_b.write(f"**Departamento:** {c.get('departamento', 'N/A')}")
                col_a.write(f"**Cargo:** {c.get('cargo', 'N/A')}")
                col_b.write(f"**Chefe Imediato:** {c.get('chefe', 'N/A')}")
                col_a.write(f"**Empresa/Unidade:** {c.get('unidade') or c.get('empresa', 'N/A')}")
                col_b.write(f"**Escolaridade:** {c.get('escolaridade', 'N/A')}")
                
                st.subheader("🎓 Cursos Obrigatórios ou Diferenciais")
                st.info(c.get("cursos") or "Não informado")

                st.subheader("🎯 Trabalho e Principal Objetivo")
                st.info(c.get("objetivo") or "Não informado")
                
                # --- TABELAS ---
                st.markdown("---")
                t_raiz = form.get('tabelas', form)
                secoes = {
                    "alta": "🚀 Alta Complexidade",
                    "normal": "📋 Complexidade Normal",
                    "baixa": "⏳ Baixa Complexidade",
                    "atividades": "📋 Atividades Executadas (Antigo)",
                    "dificuldades": "⚠️ Dificuldades e Bloqueios",
                    "sugestoes": "💡 Sugestões de Melhoria"
                }
                
                for chave, titulo in secoes.items():
                    dados_tabela = t_raiz.get(chave)
                    if dados_tabela and isinstance(dados_tabela, list):
                        try:
                            df = pd.DataFrame(dados_tabela)
                            df = df.replace("", None).dropna(how='all').fillna("")
                            if not df.empty:
                                st.subheader(titulo)
                                st.table(df)
                        except:
                            continue

                # --- QUESTIONÁRIO DISC ---
                st.markdown("---")
                st.subheader("📊 Avaliação DISC Detalhada")
                respostas_json = form.get("disc", {})

                if respostas_json:
                    for i, pergunta in enumerate(perguntas_disc):
                        letra_resposta = respostas_json.get(str(i)) or respostas_json.get(i)
                        st.write(f"**{i+1}. {pergunta}**")
                        if letra_resposta:
                            st.info(f"✅ Resposta selecionada: **{letra_resposta}**")
                        else:
                            st.warning("⚠️ Resposta não encontrada.")
                        st.divider() 
                else:
                    st.error("❌ Nenhuma resposta DISC encontrada.")

                # --- BLOCO DE EXPORTAÇÃO (DENTRO DO EXPANDER) ---
                if st.session_state.get("usuario_logado", "").strip().upper() in ["LUCIANO 123", "JV"]:
                    st.markdown("---")
                    st.subheader("⚙️ Painel de Exportação")

                    col1_exp, col2_exp = st.columns(2)

                    data_raw = form.get('timestamp') or 'sem_data'
                    data_clean = str(data_raw).replace('/', '').replace(' ', '_').replace(':', '')

                    nome_raw = form.get('colaborador') or 'Colaborador'
                    nome_clean = str(nome_raw).replace(' ', '_')

                    nome_arquivo = f"Relatorio_{nome_clean}_{data_clean}"

                    word_file = gerar_word(form)
                    pdf_file = gerar_pdf(form)

                    st.download_button(
                        label="📑 Baixar PDF",
                        data=gerar_pdf_html(form),
                        file_name=f"{nome_arquivo}.html",
                        mime="text/html",
                        key=f"pdf_unico_{id_atual}_{idx}"
                    )

        # --- SEÇÃO DE EXCLUSÃO VIRTUAL (FORA DO LOOP) ---
        st.markdown("---")
        st.subheader("🚫 Ocultar formulário da visualização")

        opcoes_para_esconder = []
        for i, f in enumerate(lista_de_arquivos):
            id_f = f.get('timestamp') or f"form_{i}"
            if id_f not in st.session_state.get("arquivos_escondidos", []):
                nome_f = (f.get('colaborador') or f.get('nome') or f"Registro {i}").upper()
                opcoes_para_esconder.append({"id": id_f, "label": nome_f})

        if opcoes_para_esconder:
            labels = [o["label"] for o in opcoes_para_esconder]
            
            escolha = st.selectbox(
                "Selecione para ocultar desta sessão:", 
                labels, 
                key="key_exclusiva_selectbox_visualizacao_v1"
            )

            if st.button("👁️‍🗨️ Ocultar Registro", key="key_exclusiva_botao_visualizacao_v1"):
                id_sel = opcoes_para_esconder[labels.index(escolha)]["id"]
                st.session_state.setdefault("arquivos_escondidos", []).append(id_sel)
                st.success("Ocultado!")
                st.rerun()
        
        if st.session_state.get("arquivos_escondidos"):
            if st.button("Resetar Visualização (Mostrar Todos)", key="key_exclusiva_reset_visualizacao_v1"):
                st.session_state["arquivos_escondidos"] = []
                st.rerun()                                    

       
# ============================================================
# CALCULAR CARGA HORÁRIA
# ============================================================

def calcular_carga(atividades):
    total_min = 0
    for at in atividades:
        try:
            tempo = float(at.get("tempo","0"))
        except:
            tempo = 0
        freq = at.get("frequencia","semanal").lower()
        if freq == "diaria":
            total_min += tempo * 5
        elif freq == "mensal":
            total_min += tempo / 4
        else:
            total_min += tempo
    horas = total_min / 60
    status = "Adequado"
    if horas > 44: status = "Sobrecarga"
    elif horas < 30: status = "Subutilização"
    return round(horas,2), status

# ============================================================
# GERAR ATIVIDADES IDEAIS (GPT)
# ============================================================

def gerar_atividades_ideais(cargo, setor, client=None):
    if client is None:
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]
    
    prompt = f"""
    Gere 12 atividades ideais para:
    Cargo: {cargo}
    Setor: {setor}
    Para cada atividade informe:
      - nome_atividade
      - descricao
      - frequencia_ideal
      - tempo_medio_minutos
      - justificativa_tecnica
    Responda SOMENTE JSON válido.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.3
        )
        
        # 1. Primeiro carregamos o conteúdo em uma variável
        dados_carregados = json.loads(response.choices[0].message.content)

        # 2. Agora injetamos o DISC na memória antes de sair da função
        if isinstance(dados_carregados, dict) and "disc" in dados_carregados:
            st.session_state["respostas_disc_fix"] = {
                str(k): v for k, v in dados_carregados["disc"].items()
            }
        
        # 3. SÓ AGORA damos o return com os dados prontos
        return dados_carregados

    except Exception as e:
        # Se der erro, retorna o padrão
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]

# ============================================================
# COMPARAÇÃO SEMÂNTICA
# ============================================================

def comparar_semanticamente(reais, ideais, client=None):
    if client is None:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

    prompt = f"""
    Compare semanticamente:
    Atividades reais: {reais}
    Atividades ideais: {ideais}
    Retorne JSON com:
      - score_aderencia (0-100)
      - tempo_gap_medio_percentual
      - atividades_desvio
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

# ============================================================
# CLASSIFICAR DIFICULDADES
# ============================================================

def classificar_dificuldades_gpt(dificuldades, client=None):
    if client is None:
        return {}
    
    prompt = f"""
    Classifique semanticamente as dificuldades abaixo em:
    - Processo
    - Tempo
    - Comunicação
    - Estrutura
    - Liderança
    - Sistema
    Retorne JSON com contagem por categoria.
    Dificuldades: {dificuldades}
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {}

# ============================================================
# ÍNDICE GERAL DO CARGO
# ============================================================

def indice_geral(score_aderencia, score_disc, status_carga):
    fator_carga = 100
    if status_carga == "Sobrecarga": fator_carga = 70
    elif status_carga == "Subutilização": fator_carga = 75
    return round(mean([score_aderencia, score_disc, fator_carga]),2)

# ============================================================
# MOTOR PRINCIPAL COMPLETO – ANÁLISE CORPORATIVA
# ============================================================

def gerar_analise_corporativa(dados, client=None):
    """
    Gera análise completa de um colaborador com base em:
    - Atividades reais
    - Perfil DISC
    - Dificuldades
    Retorna:
    - parecer (texto)
    - indicadores (dict)
    """
    # 1️⃣ Atividades ideais
    ideais = gerar_atividades_ideais(dados["cargo"], dados["setor"], client)

    # 2️⃣ Comparação semântica (reais x ideais)
    comparacao = comparar_semanticamente(dados["atividades"], ideais, client)

    # 3️⃣ Carga horária
    horas, status_carga = calcular_carga(dados["atividades"])

    # 4️⃣ Score DISC
    disc_score = score_disc(dados["disc"])

    # 5️⃣ Classificação de dificuldades
    dificuldades_classificadas = classificar_dificuldades_gpt(dados["dificuldades"], client)

    # 6️⃣ Score de aderência
    score_aderencia = comparacao.get("score_aderencia",0)

    # 7️⃣ Índice geral
    indice = indice_geral(score_aderencia, disc_score, status_carga)

    # 8️⃣ Classificação de risco
    risco = "Baixo" if indice < 60 else "Moderado" if indice < 75 else "Alto"

    # 9️⃣ Prompt final para parecer estratégico
    prompt_final = f"""
    Gere parecer estratégico completo considerando:
    - Score aderência: {score_aderencia}
    - Horas semanais: {horas}
    - Status carga: {status_carga}
    - Score DISC: {disc_score}
    - Dificuldades: {dificuldades_classificadas}
    - Índice geral do cargo: {indice}
    - Classificação de risco: {risco}
    
    Inclua:
    - Diagnóstico estrutural
    - Análise de desvios
    - Avaliação comportamental
    - Riscos organizacionais
    - Recomendação detalhada de redistribuição
    - Atividades corretas para o cargo com tempo e frequência ideais
    - Conclusão executiva
    """

    # 10️⃣ Obter parecer do GPT
    parecer = ""
    try:
        if client:
            resposta = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt_final}],
                temperature=0.3
            )
            parecer = resposta.choices[0].message.content
        else:
            parecer = "GPT não disponível. Retorno padrão: análise resumida."
    except:
        parecer = "Erro ao gerar parecer com GPT."

    # 11️⃣ Indicadores
    indicadores = {
        "score_aderencia": score_aderencia,
        "horas_semanais": horas,
        "status_carga": status_carga,
        "score_disc": disc_score,
        "indice_geral": indice,
        "risco": risco
    }

    return parecer, indicadores

# ============================================================
# GERAR PDF DO PARECER
# ============================================================

def gerar_pdf(parecer, nome):
    """
    Recebe:
    - parecer (texto)
    - nome do colaborador
    Cria arquivo PDF pronto para download
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    nome_arquivo = f"{nome}_parecer.pdf"
    doc = SimpleDocTemplate(nome_arquivo)
    elements = []
    styles = getSampleStyleSheet()

    # Título
    elements.append(Paragraph("PARECER ESTRATÉGICO ORGANIZACIONAL", styles["Title"]))
    elements.append(Spacer(1, 0.5*inch))

    # Conteúdo linha a linha
    for linha in parecer.split("\n"):
        if linha.strip():  # evita parágrafos vazios
            elements.append(Paragraph(linha, styles["Normal"]))
            elements.append(Spacer(1, 0.2*inch))

    doc.build(elements)
    return nome_arquivo

# ============================================================
# PASTA BASE PARA FORMULÁRIOS (JSON)
# ============================================================
# Usamos 'dados_dir' para manter o padrão que já criamos
json_master = os.path.join(dados_dir, "formularios.json")

# Inicializa arquivo JSON se não existir
if not os.path.exists(json_master):
    with open(json_master, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=4)


import streamlit as st
import json
from datetime import datetime
from github import Github

# =========================================================
# 1. CONFIGURAÇÕES DE ACESSO (VIA STREAMLIT SECRETS)
# =========================================================
try:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
  
    DB_TOKEN       = st.secrets["DB_TOKEN"]
    
    # Definimos o repositório direto aqui para evitar erro de Secret faltante
    REPO_NOME = "lucianohcl/formulario-colaborador"
    
except Exception as e:
    st.error(f"❌ Erro nos Secrets: A chave {e} não foi encontrada no painel do Streamlit.")
    st.stop()

if st.session_state.get("pagina") == "formulario":

    # 1. PRIMEIRO: Define as variáveis (O material)
    GITHUB_USER = "lucianohcl"
    GITHUB_REPO = "formulario-colaborador"
    FOLDER_PATH = "rascunhos" 
    GITHUB_TOKEN = st.secrets["DB_TOKEN"]

    # 2. SEGUNDO: Define a função (A ferramenta)
    @st.cache_data(ttl=10)
    def atualizar_rascunhos_do_github():
        import requests
        import json
        url = f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/contents/{FOLDER_PATH}"
        headers = {"Authorization": f"token {GITHUB_TOKEN}"}
        try:
            response = requests.get(url, headers=headers)
            if response.status_code == 200:
                # Aqui você coloca toda aquela sua lógica de rascunhos_temp...
                return response.json() 
        except:
            return {}
        return {}

    # 3. TERCEIRO: Agora sim, você usa a ferramenta! (A execução)
    if "rascunhos" not in st.session_state:
        st.session_state["rascunhos"] = {}
        # AGORA ele sabe quem é essa função
        atualizar_rascunhos_do_github()
    

    # =========================================================
    # 2. COLOQUE A FUNÇÃO AQUI (DEFINIÇÃO)
    # =========================================================
    def atualizar_rascunhos_do_github():
        # Inicializa rascunhos como dicionário vazio se não existir
        if "rascunhos" not in st.session_state:
            st.session_state["rascunhos"] = {}

        try:
            g = Github(DB_TOKEN)
            repo = g.get_repo(REPO_NOME)
            
            # Tenta acessar a pasta. Se falhar, tenta a raiz ""
            caminho = "rascunhos"
            try:
                contents = repo.get_contents(caminho)
            except:
                contents = repo.get_contents("")

            rascunhos_localizados = {}
            for content_file in contents:
                if content_file.name.endswith(".json"):
                    try:
                        file_data = content_file.decoded_content.decode("utf-8")
                        dados_json = json.loads(file_data)
                        
                        # Normalização da Chave (Upper Case)
                        nome_raw = dados_json.get("colaborador") or dados_json.get("nome")
                        if nome_raw:
                            nome_chave = str(nome_raw).strip().upper() 
                            rascunhos_localizados[nome_chave] = dados_json
                    except:
                        continue # Se um arquivo estiver com erro, não quebra os outros
            
            # IMPORTANTE: Só atualiza se encontrar algo, para não limpar o que já tem
            if rascunhos_localizados:
                st.session_state["rascunhos"] = rascunhos_localizados
                return True
            return False

        except Exception as e:
            # Se der erro de rede, mantém o que já estava na memória para não sumir tudo
            return False

    # CHAME A FUNÇÃO AUTOMATICAMENTE NA INICIALIZAÇÃO
    # Adicionamos uma trava para não ficar rodando toda hora sem necessidade
    if "rascunhos" not in st.session_state or not st.session_state["rascunhos"]:
        atualizar_rascunhos_do_github()


    # =========================================================
    # 3. CHAME A EXECUÇÃO AQUI (INICIALIZAÇÃO)
    # =========================================================
    if "rascunhos" not in st.session_state:
        atualizar_rascunhos_do_github()

    # 4. PARTE VISUAL (Daqui para baixo segue o seu st.title, etc)


    # =========================================================
    # 2. FUNÇÃO PARA SALVAR DADOS NO GITHUB
    # =========================================================
    def salvar_no_github(conteudo_dict, nome_arquivo):
        try:
            g = Github(DB_TOKEN)
            repo = g.get_repo(REPO_NOME)
            caminho_git = f"dados/{nome_arquivo}"
            
            json_string = json.dumps(conteudo_dict, ensure_ascii=False, indent=4)
            
            try:
                contents = repo.get_contents(caminho_git)
                repo.update_file(contents.path, f"Update: {nome_arquivo}", json_string, contents.sha)
            except:
                repo.create_file(caminho_git, f"Novo envio: {nome_arquivo}", json_string)
            
            return True
        except Exception as e:
            st.error(f"❌ Erro ao conectar com o GitHub: {e}")
            return False

    # =========================================================
    # 3. INTERFACE E LÓGICA DO FORMULÁRIO
    # =========================================================


    from github import Github
    import json
    import streamlit as st

    def salvar_no_github(conteudo_dict, nome_arquivo):
        try:
            g = Github(st.secrets["DB_TOKEN"])
            repo = g.get_repo("lucianohcl/formulario-colaborador")
            caminho_git = f"dados/{nome_arquivo}"
            
            json_string = json.dumps(conteudo_dict, ensure_ascii=False, indent=4)
            
            try:
                contents = repo.get_contents(caminho_git)
                repo.update_file(contents.path, f"Update: {nome_arquivo}", json_string, contents.sha)
            except:
                repo.create_file(caminho_git, f"Novo envio: {nome_arquivo}", json_string)
            
            return True
        except Exception as e:
            st.error(f"❌ Erro ao conectar com o GitHub: {e}")
            return False







        # ============================================================
        # GARANTIA DE PERSISTÊNCIA (CARGA DOS DADOS)
        # ============================================================

        # Recarregamos os dados diretamente do disco/nuvem para garantir persistência total
        st.session_state["formularios"] = carregar_todos_formularios()

    import streamlit as st
    import pandas as pd
    import plotly.express as px

    # ============================================================
    # 1. FUNÇÕES DE APOIO (CÁLCULOS E TRADUÇÃO)
    # ============================================================

    MAPA_DISC = {
        "A": "D", "B": "I", "C": "S", "D": "C"
    }

    def calcular_disc(respostas_disc):
        """Traduz A, B, C, D para D, I, S, C e calcula percentuais."""
        contagem = {"D": 0, "I": 0, "S": 0, "C": 0}
        for r in respostas_disc.values():
            perfil = MAPA_DISC.get(r)
            if perfil in contagem:
                contagem[perfil] += 1
        
        total = sum(contagem.values())
        if total > 0:
            percentuais = {k: round(v/total*100, 1) for k, v in contagem.items()}
            dominante = max(percentuais, key=percentuais.get)
        else:
            percentuais = {"D": 0, "I": 0, "S": 0, "C": 0}
            dominante = None
        return percentuais, dominante

    # ============================================================
    # 2. CARREGAMENTO E PERSISTÊNCIA
    # ============================================================

    # Garante que os dados estejam carregados na sessão
    formularios = carregar_todos_formularios()
    st.session_state["formularios"] = formularios

    
    import streamlit as st
    import pandas as pd
    import os
    import json
    import sys

    # ============================================================
    # CONFIGURAÇÃO DE DIRETÓRIO E CARREGAMENTO
    # ============================================================

    # Define o diretório base e a pasta de dados
    base_dir = os.path.dirname(os.path.abspath(__file__))
    dados_dir = os.path.join(base_dir, "dados")
    os.makedirs(dados_dir, exist_ok=True)

    # Função para carregar todos os JSONs da pasta 'dados'
    def carregar_todos_formularios():
        lista_formularios = []
        if os.path.exists(dados_dir):
            for nome_arquivo in os.listdir(dados_dir):
                if nome_arquivo.endswith(".json"):
                    caminho_completo = os.path.join(dados_dir, nome_arquivo)
                    try:
                        with open(caminho_completo, "r", encoding="utf-8") as f:
                            dados = json.load(f)
                            if isinstance(dados, dict):
                                lista_formularios.append(dados)
                    except Exception:
                        continue
        return lista_formularios

    # Inicializa o estado da sessão com os dados carregados
    if "formularios" not in st.session_state:
        st.session_state["formularios"] = carregar_todos_formularios()

    # --- BLOCO DE CSS PARA OCULTAÇÃO ---
    if st.query_params.get("page") == "formulario":
        st.markdown("""
        <style>
            [data-testid="stSidebar"] {display: none !important;}
            #MainMenu, footer, header {visibility: hidden !important;}
        </style>
        """, unsafe_allow_html=True)


    # =========================================================
    # 1. FUNÇÕES DE SUPORTE
    # =========================================================
    def preparar_df(chave_json, colunas, fonte_local, linhas_padrao=15):
        if not isinstance(fonte_local, dict): fonte_local = {}
        dados = fonte_local.get(chave_json, [])
        if dados and isinstance(dados, list):
            df = pd.DataFrame(dados)
            for col in colunas:
                if col not in df.columns: df[col] = ""
            return df[colunas]
        return pd.DataFrame({col: [""] * linhas_padrao for col in colunas})

    lista_frequencia = ["", "DVD", "D", "S", "Q", "M", "T", "A"]
    lista_horas = [f"{i} h" for i in range(25)]
    lista_minutos = [f"{i} min" for i in range(0, 60, 5)]

    col_atv = ["Atividade", "Frequência", "Horas", "Minutos"]
    # Altere para:
    col_dif = ["Dificuldade/Bloqueio", "Setor/Parceiro Envolvido", "Frequência", "Horas", "Minutos"]
    col_sug = ["Sugestão de Melhoria", "Impacto Esperado", "Frequência", "Horas", "Minutos"]


    # =========================================================
    # 🎭 CAPA DE INVISIBILIDADE (NÃO QUEBRA O RASCUNHO)
    # =========================================================

    # Criamos um lugar no app que pode ser "esvaziado"
    area_do_formulario = st.container()

    # Se NÃO estivermos na página do formulário, a gente limpa a área visual
    # mas deixa o código das 3000 linhas rodar "em silêncio" para o rascunho
    if st.session_state.get("pagina") != "formulario":
        area_do_formulario.empty() 

    # Agora, para o Título e as mensagens iniciais, usamos o 'with'
    with area_do_formulario:
        resgate = st.session_state.get("rascunho_atual", {})
        nome_titulo = resgate.get("colaborador", "Novo Formulário")

        st.title("📋 Formulário de Acompanhamento")

        if nome_titulo != "Novo Formulário":
            st.info(f"✨ **Editando Rascunho de:** {nome_titulo}")
        else:
            st.success("📝 **Criando Novo Registro**")

        st.markdown("---")

    # ABAIXO SEGUEM AS 3000 LINHAS SEM INDENTAÇÃO
    # O rascunho continuará funcionando porque o código está sendo lido,
    # mas o 'area_do_formulario.empty()' lá em cima ajuda a limpar o topo.

    # =========================================================
    # Perguntas DISC
    # =========================================================
    perguntas_disc = [
        "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
        "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
        "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
        "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
        "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
        "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
        "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constantemente | (D) Metódico/Cauteloso",
        "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
        "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
        "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
        "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
        "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
        "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
        "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
        "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
        "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
        "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
        "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
        "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
        "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
        "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
        "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
        "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
        "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
    ]


    # =========================================================
    # 👤 DADOS DE IDENTIFICAÇÃO (AJUSTADO PARA 5 TABELAS)
    # =========================================================
    st.subheader("👤 Dados de Identificação")

    st.write("DEBUG - O que tem no state agora:", st.session_state.get("f_cargo"))

    fonte = st.session_state.get("dados_oficiais", {})
    col1, col2 = st.columns(2)

    with col1:
        # Mostra os nomes encontrados no GitHub
        rascunhos_dict = st.session_state.get("rascunhos", {})
        nomes_disponiveis = list(rascunhos_dict.keys())
        st.write(f"🗂️ Rascunhos na Nuvem: **{', '.join(nomes_disponiveis) if nomes_disponiveis else 'Nenhum'}**")

        v = st.session_state.get("v_tab", 0)
        nome_f = st.text_input(
            "Nome do colaborador",
            value=st.session_state.get("f_nome_v2") or fonte.get("nome", ""),
            key=f"f_nome_{v}"
        )

        if st.button("📥 Carregar Rascunho", key="btn_carregar_rascunho_v3"):
            if nome_f:
                atualizar_rascunhos_do_github() 
                rascunhos_dict = st.session_state.get("rascunhos", {})
                
                nome_busca = nome_f.strip().upper()
                rascunho = rascunhos_dict.get(nome_busca)
                
                if rascunho:
                    # 1. SALVA O ESTADO GLOBAL
                    st.session_state["rascunho"] = rascunho  # <--- CRUCIAL para o motor de tabelas
                    st.session_state["f_nome_v2"] = nome_busca
                    st.session_state["v_tab"] = st.session_state.get("v_tab", 0) + 1
                    
                    # 2. DADOS BÁSICOS (Ajustado para a estrutura do seu JSON)
                    cp = rascunho.get("campos", {})
                    
                    # Mapeamento: "Chave_do_Widget": "Chave_dentro_do_JSON_campos"
                    campos_map = {
                        "f_cargo": "cargo",
                        "f_depto": "dep",         # No JSON está "dep", não "departamento"
                        "f_setor": "setor",
                        "f_chefe": "chefe",
                        "f_unidade": "unidade",
                        "f_esc": "escolaridade",
                        "f_dev": "devolver_em",    # No JSON está "devolver_em"
                        "f_cursos_area": "cursos",
                        "f_obj_area": "objetivo"
                    }
                    
                    for key_ui, key_json in campos_map.items():
                        valor = cp.get(key_json, "")
                        st.session_state[key_ui] = valor
                        st.session_state[f"{key_ui}_v2"] = valor

                    # 3. DISC - SINCRONIZAÇÃO TOTAL
                    disc_salvo = rascunho.get("disc", {})
                    if disc_salvo:
                        for i in range(24):
                            chave_json = str(i)
                            # Sincroniza a chave do rádio que você usa no loop do DISC
                            # Importante: a chave aqui deve ser idêntica à definida no st.radio
                            v = st.session_state["v_tab"]
                            st.session_state[f"disc_radio_{i}_{v}"] = disc_salvo.get(chave_json)


                    # --- 4. TABELAS - SINCRONIZAÇÃO ---
                    tabelas_salvas = rascunho.get("tabelas", {})
                    if tabelas_salvas:
                        # Lista das chaves de tabela que você tem no JSON
                        chaves_tabelas = ["alta", "normal", "baixa", "dificuldades", "sugestoes"]
                        
                        for t_key in chaves_tabelas:
                            dados = tabelas_salvas.get(t_key, [])
                            if dados:
                                # Converte a lista de dicionários do JSON em um DataFrame
                                df_carregado = pd.DataFrame(dados)
                                
                                # SALVA NO SESSION_STATE 
                                # Importante: A chave deve ser EXATAMENTE a 'key' que você usa no st.data_editor
                                # Se o seu data_editor usa key=f"editor_{t_key}_{v}", faça igual:
                                v = st.session_state["v_tab"]
                                st.session_state[f"editor_{t_key}_{v}"] = df_carregado

                    st.success(f"✅ Rascunho e DISC de {nome_busca} carregados!")
                    st.rerun()
                else:
                    st.error(f"⚠️ Rascunho de '{nome_busca}' não encontrado.")
            else:
                st.warning("⚠️ Digite um nome antes de carregar.")

    with col2:
        cargo_f = st.text_input("Cargo", value=st.session_state.get("f_cargo_v2") or fonte.get("cargo", ""), key="f_cargo")
        depto_f = st.text_input("Departamento", value=st.session_state.get("f_depto_v2") or fonte.get("departamento", ""), key="f_depto")
        esc_f = st.text_input("Escolaridade", value=st.session_state.get("f_esc_v2") or fonte.get("escolaridade", ""), key="f_esc")
        setor_f = st.text_input("Setor", value=st.session_state.get("f_setor_v2") or fonte.get("setor", ""), key="f_setor")
        chefe_f = st.text_input("Chefe imediato", value=st.session_state.get("f_chefe_v2") or fonte.get("chefe", ""), key="f_chefe")
        unidade_f = st.text_input("Empresa / Unidade", value=st.session_state.get("f_unidade_v2") or fonte.get("unidade", ""), key="f_unidade")
        dev_f = st.text_input("Devolver preenchido em", value=st.session_state.get("f_dev_v2") or fonte.get("devolucao", ""), key="f_dev")

    cursos_f = st.text_area("Cursos Obrigatórios e Diferenciais", value=st.session_state.get("f_cursos_v2") or fonte.get("cursos", ""), key="f_cursos_area")
    obj_f = st.text_area("Em que consiste seu Trabalho e qual seu Principal Objetivo", value=st.session_state.get("f_obj_v2") or fonte.get("objetivo", ""), key="f_obj_area")

        

    # =========================================================
    # 5. TABELAS DE TAREFAS (COM FUNÇÃO DE SUPORTE INTEGRADA)
    # =========================================================
    st.markdown("---")

    st.subheader("📋 Tabelas") # Título médio

    # --- INÍCIO DAS LEGENDAS ---
    col_leg1, col_leg2 = st.columns(2)

    with col_leg1:
        st.info("""
        **📋 LEGENDA DE FREQUÊNCIA:**
        * **DVD**: Diário Várias Vezes
        * **D**: Diário | **S**: Semanal
        * **Q**: Quinzenal | **M**: Mensal
        * **T**: Trimestral | **A**: Anual
        """)

    with col_leg2:
        st.warning("""
        **⏱️ COMO REGISTRAR O TEMPO:**
        * **Horas e Minutos**: Selecione o valor em cada coluna.
        * **Menos de 1 hora?**: Selecione **0 h** e o tempo real em minutos.
        * **Não se aplica?**: Selecione **0 h** e **0 min** em ambos.
        """)
    # --- FIM DAS LEGENDAS ---

    # --- FUNÇÃO AUXILIAR (Garante que o editor tenha linhas suficientes) ---
    def garantir_15_linhas(df, colunas):
        if df is None or df.empty:
            df = pd.DataFrame(columns=colunas)
        for col in colunas:
            if col not in df.columns: df[col] = ""
        while len(df) < 15:
            df.loc[len(df)] = [""] * len(colunas)
        return df.head(15)

    # 1. Configurações e rascunho
    lista_frequencia = ["", "DVD", "D", "S", "Q", "M", "T", "A"]
    lista_horas = [f"{i} h" for i in range(25)]
    lista_minutos = [f"{i} min" for i in range(0, 60, 5)]

    if "rascunho_atual" not in st.session_state:
        st.session_state["rascunho_atual"] = {}

    rascunho = st.session_state.get("rascunho", {})
    v_layout = st.session_state.get("v_tab", 0)

    # 2. Definição da função de renderização
    def gerar_tabela_final(titulo, chave_json, col_principal, col_extra=None, label_extra=None):
        st.subheader(titulo)
        dict_tabelas = rascunho.get("tabelas", {}) if isinstance(rascunho, dict) else {}
        dados_salvos = dict_tabelas.get(chave_json, [])
        
        colunas = [col_principal, "Horas", "Minutos", "Frequência"]
        if col_extra: 
            colunas.insert(1, col_extra)
        
        # Chama a função que estava dando NameError
        df_base = garantir_15_linhas(pd.DataFrame(dados_salvos), colunas)
        
        config_tab = {
            col_principal: st.column_config.TextColumn("Descrição", width="large"),
            "Frequência": st.column_config.SelectboxColumn("Frequência", options=lista_frequencia, width="small"),
            "Horas": st.column_config.SelectboxColumn("Horas", options=lista_horas, width="small"),
            "Minutos": st.column_config.SelectboxColumn("Minutos", options=lista_minutos, width="small"),
        }
        if col_extra: 
            config_tab[col_extra] = st.column_config.TextColumn(label_extra, width="medium")

        return st.data_editor(
            df_base, 
            key=f"editor_{chave_json}_v{v_layout}", 
            column_config=config_tab, 
            use_container_width=True, 
            num_rows="fixed"
        )

    # 3. Chamadas das Tabelas (Sincronização Total com o JSON)
    # Nota: O terceiro parâmetro deve ser EXATAMENTE a chave do JSON (ex: "Atividade")

    e_alta = gerar_tabela_final("🚀 Atividades de Alta Complexidade", "alta", "Atividade")

    e_normal = gerar_tabela_final("📋 Atividades de Complexidade Normal", "normal", "Atividade")

    e_baixa = gerar_tabela_final("⏳ Atividades de Baixa Complexidade", "baixa", "Atividade")

    # Aqui mudamos de "Setor/Parceiro Envolvido" para "Setor Envolvido"
    e_dif = gerar_tabela_final("⚠️ Dificuldades e Bloqueios", "dificuldades", "Dificuldade", "Setor Envolvido", "Setor Envolvido")

    # Aqui mudamos de "Impacto" para "Impacto Esperado"
    e_sug = gerar_tabela_final("💡 Sugestões de Melhoria", "sugestoes", "Sugestão", "Impacto Esperado", "Impacto Esperado")


    # =========================================================
    # 📊 7. QUESTIONÁRIO DISC (SINCRONIZADO COM O JSON)
    # =========================================================
    st.markdown("---")
    st.subheader("📊 Questionário")

    v = st.session_state.get("v_tab", 0) 

    # 1. BUSCA O RASCUNHO NO LUGAR CERTO (Onde o JSON que você mostrou reside)
    # O seu JSON mostra que o DISC está dentro de "rascunho" -> "disc"
    rascunho_disc = st.session_state.get("rascunho", {}).get("disc", {})

    respostas_disc_atual = {}

    for i, pergunta in enumerate(perguntas_disc):
        # 2. BUSCA PELA CHAVE DO JSON (O seu JSON usa apenas o número como string)
        chave_json = str(i)
        letra_salva = rascunho_disc.get(chave_json)
        
        # 3. DEFINE O ÍNDICE (A=0, B=1, C=2, D=3)
        opcoes = ["A", "B", "C", "D"]
        # Se letra_salva for None ou null no JSON, o index será None (fica desmarcado)
        idx_selecionado = opcoes.index(letra_salva) if letra_salva in opcoes else None
        
        # 4. O WIDGET
        escolha = st.radio(
            f"**{i+1}.** {pergunta}",
            options=opcoes,
            index=idx_selecionado,
            key=f"disc_radio_{i}_{v}",
            horizontal=True
        )
        
        # 5. GUARDA PARA SALVAR DEPOIS (Mantendo o padrão de string do JSON)
        respostas_disc_atual[chave_json] = escolha

    # 6. ATUALIZA O RASCUNHO GLOBAL NA HORA
    # Verifica se o rascunho existe; se não, inicializa para evitar o KeyError
    if "rascunho" not in st.session_state:
        st.session_state["rascunho"] = {}

    # Garante que a sub-chave "disc" também exista
    if "disc" not in st.session_state["rascunho"]:
        st.session_state["rascunho"]["disc"] = {}

    # Agora sim, salva as respostas sem risco de quebrar
    st.session_state["rascunho"]["disc"] = respostas_disc_atual



    # =========================================================
    # 6. VALIDAÇÃO UNIFICADA (TABELAS, DISC E CABEÇALHO)
    # =========================================================
    st.markdown("---")
    st.subheader("✅ Status de Validação do Formulário")

    pendencias = []

    # --- 1. VALIDAÇÃO DE CABEÇALHO ---
    campos_id = {
        "Nome": nome_f, "Cargo": cargo_f, "Departamento": depto_f,
        "Escolaridade": esc_f, "Setor": setor_f, "Chefe Imediato": chefe_f,
        "Empresa/Unidade": unidade_f, "Devolver em": dev_f,
        "Cursos": cursos_f, "Objetivo": obj_f
    }
    for campo, valor in campos_id.items():
        if not valor or str(valor).strip() == "":
            pendencias.append(f"Identificação: O campo **{campo}** está vazio.")

    # --- 2. VALIDAÇÃO DAS TABELAS (RIGOR TOTAL) ---
    dict_tabelas = {
        "Alta Complexidade": e_alta, 
        "Complexidade Normal": e_normal,
        "Baixa Complexidade": e_baixa, 
        "Dificuldades": e_dif,
        "Sugestões e Melhorias": e_sug
    }

    regras_colunas = {
        "Alta Complexidade": "Atividade", 
        "Complexidade Normal": "Atividade",
        "Baixa Complexidade": "Atividade", 
        "Dificuldades": "Dificuldade",
        "Sugestões e Melhorias": "Sugestão"
    }

    for nome_tab, df_validar in dict_tabelas.items():
        col_alvo = regras_colunas.get(nome_tab)
        
        if df_validar is not None and col_alvo in df_validar.columns:
            # Identifica linhas onde a descrição foi preenchida
            linhas_ativas = df_validar[df_validar[col_alvo].astype(str).str.strip() != ""]
            
            if len(linhas_ativas) == 0:
                pendencias.append(f"⚠️ A tabela **{nome_tab}** está totalmente vazia. Preencha pelo menos 1 item.")
            else:
                for i, row in linhas_ativas.iterrows():
                    # Extração limpa dos valores
                    h_str = str(row.get("Horas", "")).strip()
                    m_str = str(row.get("Minutos", "")).strip()
                    freq = str(row.get("Frequência", "")).strip()
                    
                    # Validação de Horas
                    if h_str == "":
                        pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar as **Horas**.")
                    
                    # Validação de Minutos
                    if m_str == "":
                        pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar os **Minutos**.")
                    
                    # Validação de Frequência
                    if freq == "":
                        pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar a **Frequência**.")
                    
                    # Validação extra para colunas específicas (Impacto / Setor)
                    if nome_tab == "Dificuldades":
                        if str(row.get("Setor Envolvido", "")).strip() == "":
                            pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Informe o **Setor Envolvido**.")
                    
                    if nome_tab == "Sugestões e Melhorias":
                        if str(row.get("Impacto Esperado", "")).strip() == "":
                            pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Informe o **Impacto Esperado**.")

    # --- 3. VALIDAÇÃO DO DISC ---
    respostas_vazias = [k for k, v in respostas_disc_atual.items() if v is None]
    if len(respostas_vazias) > 0:
        pendencias.append(f"Questionário: Faltam responder **{len(respostas_vazias)} questões**.")

    # --- EXIBIÇÃO FINAL DO STATUS ---
    if pendencias:
        st.warning(f"⚠️ **Existem {len(pendencias)} pendências obrigatórias:**")
        for p in pendencias:
            st.write(f"• {p}")
        st.session_state["confirmacao_final"] = False
    else:
        st.success("🎉 **Perfeito! Tudo preenchido corretamente. O envio está liberado.**")

    # =========================================================
    # 🚀 4. BOTÃO DE ENVIO E SALVAMENTO REAL (VERSÃO FINAL)
    # =========================================================

    # Centralizando o botão para dar mais destaque
    col_btn, _ = st.columns([2, 1])

    with col_btn:
        if st.button("🚀 FINALIZAR E ENVIAR FORMULÁRIO", type="primary", use_container_width=True):
            if pendencias:
                st.error("❌ Corrija as pendências listadas acima antes de enviar.")
                st.stop()
            
            # Sistema de confirmação dupla
            if not st.session_state.get("confirmacao_final", False):
                st.warning(f"⚠️ **{nome_f}**, clique novamente para confirmar o envio definitivo.")
                st.session_state["confirmacao_final"] = True
                st.stop()

            try:
                from datetime import datetime
                import json

                def preparar_dados(df):
                    if df is None or df.empty: return []
                    # Pega dinamicamente a coluna de Descrição (Atividade, Dificuldade ou Sugestão)
                    col_principal = df.columns[0] 
                    return df[df[col_principal].astype(str).str.strip() != ""].to_dict("records")

                # 1. EXTRAÇÃO FORÇADA (Garante que os dados saiam do rádio e entrem no código)
                dados_disc_final = {}
                for i in range(24):
                    # Tentamos todas as variações possíveis de nomes que você pode ter usado
                    valor = st.session_state.get(f"q_{i}") or st.session_state.get(f"q{i}") or st.session_state.get(f"p{i}") or ""
                    dados_disc_final[str(i)] = valor

                payload = {
                    "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                    "colaborador": nome_f,
                    "campos": {
                        "cargo": cargo_f,
                        "departamento": depto_f,
                        "setor": setor_f,
                        "chefe": chefe_f,
                        "unidade": unidade_f,
                        "escolaridade": esc_f,
                        "devolver_em": dev_f,
                        "cursos": cursos_f,
                        "objetivo": obj_f
                    },
                    "tabelas": {
                        "alta": preparar_dados(e_alta),
                        "normal": preparar_dados(e_normal),
                        "baixa": preparar_dados(e_baixa),
                        "dificuldades": preparar_dados(e_dif),
                        "sugestoes": preparar_dados(e_sug)
                    },
                    "disc": st.session_state.get("rascunho", {}).get("disc", {})
                }




                nome_arquivo = f"{nome_f.replace(' ', '_').upper()}.json"
                
                with st.spinner("Sincronizando..."):
                    sucesso = salvar_no_github(payload, nome_arquivo, pasta="dados")

                    if sucesso:
                        
                        st.success(f"✅ Formulário de {nome_f} enviado com sucesso!")
                        enviado = enviar_para_sheets(payload)

                        if enviado:
                            st.toast("📊 Enviado para Google Sheets!")
                        else:
                            st.warning("⚠️ Salvou, mas não enviou para o Sheets")



                        # Limpa os estados de controle
                        st.session_state["confirmacao_final"] = False
                        # Opcional: st.session_state["rascunho_atual"] = {} 
                    else:
                        st.error("⚠️ O GitHub não respondeu. Mas não se preocupe! Baixe o arquivo abaixo e envie por e-mail/WhatsApp.")

                    # Botão de download sempre visível após tentativa de envio (Backup)
                    st.download_button(
                        label="📥 Baixar Cópia de Segurança (JSON)",
                        data=json.dumps(payload, indent=4, ensure_ascii=False),
                        file_name=nome_arquivo,
                        mime="application/json",
                        use_container_width=True
                    )

            except Exception as e:
                st.error(f"❌ Erro crítico ao processar o envio: {e}")






    import streamlit as st
    import pandas as pd
    import json
    from datetime import datetime
    from github import Github

    # =========================================================

    # 1. CONFIGURAÇÕES E CONEXÃO

    # =========================================================

    st.set_page_config(page_title="Formulário Analítico", layout="wide")



    # Puxa o Token dos segredos

    TOKEN = st.secrets["DB_TOKEN"]



    try:

        DB_USERNAME = st.secrets["DB_USERNAME"]

    except Exception:

        DB_USERNAME = "lucianohcl"



    # Conecta ao Github

    g = Github(TOKEN)
    


    # ESSE É O PONTO: Coloque o texto direto aqui para não dar erro de "REPO_NOME not defined"

    repo = g.get_repo("lucianohcl/formulario-colaborador") 



    # Se precisar do username para outra coisa, use direto dos secrets ou fixo:



    if "rascunho" not in st.session_state: st.session_state["rascunho"] = {}

    if "logado" not in st.session_state: st.session_state["logado"] = False



    def val(chave, default=""):

        # Só tenta ler se o rascunho existir, senão ignora

        if "rascunho" in st.session_state:

            d = st.session_state["rascunho"]

            return d.get("campos", {}).get(chave, d.get(chave, default))

        return default


    # =========================================================
    # 2. IDENTIFICAÇÃO E CARREGAMENTO (VERSÃO BLINDADA)
    # =========================================================
    st.subheader("📋 Acesso ao Rascunho")
    nome_input = st.text_input("DIGITE SEU NOME COMPLETO:").strip().upper()
    # ADICIONE ESTA LINHA ABAIXO:
    nome_digitado = nome_input

    if not nome_input:
        st.info("Digite seu nome para começar.")
        st.stop()

    nome_arq = f"rascunhos/{nome_input.replace(' ', '_')}.json"

    if st.session_state.get("usuario_atual") != nome_input:
        st.session_state["usuario_atual"] = nome_input
        st.session_state["logado"] = False

    confirmar = st.checkbox("✅ CLIQUE PARA CARREGAR MEUS DADOS")

    if confirmar and not st.session_state.get("logado"):
        try:
            conteudo = repo.get_contents(nome_arq)
            dados_carregados = json.loads(conteudo.decoded_content.decode())
            # Garante que o rascunho seja EXATAMENTE o que está no GitHub
            st.session_state["rascunho"] = dados_carregados
            st.success("Dados recuperados!")
        except:
            st.session_state["rascunho"] = {"colaborador": nome_input, "campos": {}, "tabelas": {}, "disc": {}}
            st.info("Iniciando novo rascunho.")
        
        st.session_state["logado"] = True
        st.rerun()

    # Se não confirmou, para aqui
    if not st.session_state["logado"]:
        st.stop()

    # =========================================================
    # 3. FORMULÁRIO: CAMPOS DE TEXTO
    # =========================================================
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        cargo = st.text_input("Cargo:", value=val("cargo"))
        depto = st.text_input("Departamento:", value=val("dep"))
        setor = st.text_input("Setor:", value=val("setor"))
    with col2:
        chefe = st.text_input("Chefe imediato:", value=val("chefe"))
        unidade = st.text_input("Empresa / Unidade:", value=val("unidade"))
        escolaridade = st.text_input("Escolaridade:", value=val("escolaridade"))
        devolver_em = st.text_input("Devolver em:", value=val("devolver_em"))

    cursos = st.text_area("Cursos Obrigatórios e Diferenciais:", value=val("cursos"))
    objetivo = st.text_area("Em que consiste seu trabalho e qual seu Principal Objetivo:", value=val("objetivo"))


    # =========================================================
    # 4. MOTOR DE TABELAS (VERSÃO ULTRA-BLINDADA V3)
    # =========================================================
    def criar_editor(titulo, chave, col_p, col_e=None, nome_e=None):
        # --- 1. RESET DE MEMÓRIA (FORÇA O STREAMLIT A REDESENHAR O LAYOUT) ---
        # Se a v3 ainda não existe no estado da sessão, limpamos as versões antigas
        if f"ed_{chave}_v3" not in st.session_state:
            for k in list(st.session_state.keys()):
                if f"ed_{chave}" in k:
                    del st.session_state[k]

        st.write(f"**{titulo}**")
        
        # 2. Puxa os dados do rascunho
        dados = st.session_state["rascunho"].get("tabelas", {}).get(chave, [])
        df = pd.DataFrame(dados)
        
        # 3. Define a Ordem Rígida (O Setor/Impacto TEM que ser a segunda coluna)
        if col_e:
            cols_finais = [col_p, col_e, "Horas", "Minutos", "Frequência"]
        else:
            cols_finais = [col_p, "Horas", "Minutos", "Frequência"]
        
        # 4. Limpeza e Reindex (Remove colunas fantasmas e organiza a ordem)
        df = df.fillna("").astype(str)
        for c in df.columns:
            df[c] = df[c].str.strip()
        
        # O reindex com columns=cols_finais descarta qualquer coluna que não esteja na lista
        df = df.reindex(columns=cols_finais, fill_value="")
        
        # 5. Garante as 15 linhas fixas
        if len(df) < 15:
            faltam = 15 - len(df)
            extras = pd.DataFrame([{c: "" for c in cols_finais} for _ in range(faltam)])
            df = pd.concat([df, extras], ignore_index=True)
        
        # Trava em 15 linhas e na ordem correta
        df = df[cols_finais].head(15)

        # 6. Configuração Visual dos Seletores
        cfg = {
            col_p: st.column_config.TextColumn("Descrição", width="large"),
            "Frequência": st.column_config.SelectboxColumn(
                options=["", "DVD", "D", "S", "Q", "M", "T", "A"], 
                width="small"
            ),
            "Horas": st.column_config.SelectboxColumn(
                options=[""] + [f"{i} h" for i in range(25)], 
                width="small"
            ),
            "Minutos": st.column_config.SelectboxColumn(
                options=[""] + [f"{i} min" for i in range(0, 60, 5)], 
                width="small"
            ),
        }
        if col_e: 
            cfg[col_e] = st.column_config.TextColumn(nome_e, width="medium")
            
        # 7. Renderização Híbrida (Texto completo para análises longas)
        import pandas as pd
        pd.set_option('display.max_colwidth', None)

        st.markdown("""
            <style>
                .tabela-ajustada { overflow-x: auto; width: 100%; }
                table { min-width: 1000px !important; width: 100% !important; border-collapse: collapse; }
                th { background-color: #f0f2f6; color: #31333F; text-align: left; padding: 10px; }
                td { 
                    white-space: normal !important; 
                    word-wrap: break-word !important; 
                    padding: 10px !important;
                    font-size: 14px !important;
                    vertical-align: top !important;
                    border-bottom: 1px solid #ddd;
                }
            </style>
        """, unsafe_allow_html=True)

        st.markdown('<div class="tabela-ajustada">', unsafe_allow_html=True)
        # Usamos st.table para garantir que o texto quebre linha e apareça TODO
        st.table(df)
        st.markdown('</div>', unsafe_allow_html=True)
        
        return None # Como já renderizamos com st.table, não precisamos retornar o editor

    # Chamadas das tabelas
    e_alta = criar_editor("🚀 Alta Complexidade", "alta", "Atividade")
    e_normal = criar_editor("📋 Complexidade Normal", "normal", "Atividade")
    e_baixa = criar_editor("⏳ Baixa Complexidade", "baixa", "Atividade")
    e_dif = criar_editor("⚠️ Dificuldades", "dificuldades", "Dificuldade", "Setor Envolvido")
    e_sug = criar_editor("💡 Sugestões", "sugestoes", "Sugestão", "Impacto Esperado")


    # =========================================================
    # 6. PERFIL DISC (PERSISTÊNCIA GARANTIDA - CORRIGIDO)
    # =========================================================
    st.markdown("---")
    st.subheader("📊 Questionário DISC")

    # =========================================================
    # 🔥 GARANTE ESTRUTURA
    # =========================================================
    if "rascunho_atual" not in st.session_state:
        st.session_state["rascunho_atual"] = {}

    if "disc" not in st.session_state["rascunho_atual"]:
        st.session_state["rascunho_atual"]["disc"] = {}

    # Recupera o dicionário salvo do rascunho (Garante chaves como string)
    disc_data = {
        str(k): v for k, v in st.session_state["rascunho_atual"]["disc"].items()
    }

    # Pegamos o nome do colaborador para resetar os campos se mudar de pessoa
    nome_colab = st.session_state.get("nome_colaborador", "novo")

    # =========================================================
    # 📋 LISTA DE PERGUNTAS (EXATA PARA ESPELHAMENTO)
    # =========================================================
    perguntas_disc = [
        "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
        "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
        "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
        "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
        "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
        "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
        "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
        "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
        "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
        "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
        "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
        "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
        "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
        "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
        "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
        "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
        "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
        "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
        "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
        "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
        "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
        "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
        "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
        "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
    ]

    # =========================================================
    # 🎯 DEFINIÇÕES E RENDER (SINCRONIZADO)
    # =========================================================
    opcoes = ["A", "B", "C", "D"]
    respostas_disc_final = {} # Nome correto para evitar o NameError

    for i, pergunta in enumerate(perguntas_disc):
        chave = str(i)
        
        # Busca a letra no banco. Se não houver, retorna None para ficar desmarcado
        letra_banco = disc_data.get(chave)
        
        # Define o índice da bolinha (A=0, B=1, C=2, D=3)
        idx = opcoes.index(letra_banco) if letra_banco in opcoes else None

        respostas_disc_final[chave] = st.radio(
            f"**{i+1}. {pergunta}**",
            options=opcoes,
            index=idx,
            horizontal=True,
            key=f"disc_{nome_colab}_{i}" # Key dinâmica para resetar ao trocar colaborador
        )

    # =========================================================
    # 💾 PERSISTÊNCIA AUTOMÁTICA
    # =========================================================
    st.session_state["rascunho_atual"]["disc"] = respostas_disc_final



    # =========================================================
    # 6. SALVAMENTO (GITHUB)
    # =========================================================
    if st.button("💾 SALVAR TUDO", use_container_width=True):

        # --- VALIDAÇÃO CRÍTICA ---
        if not nome_input or len(nome_input) < 3:
            st.error("⚠️ Erro: Nome do colaborador está vazio ou inválido.")
            st.stop()

        # 1. Monta o payload (O "corpo" do arquivo)
        payload = {
            "colaborador": nome_input, 
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "campos": {
                "cargo": cargo, "dep": depto, "setor": setor, 
                "chefe": chefe, "unidade": unidade, 
                "escolaridade": escolaridade, "devolver_em": devolver_em, 
                "cursos": cursos, "objetivo": objetivo
            },
            "tabelas": {
                "alta": e_alta.to_dict("records"), 
                "normal": e_normal.to_dict("records"), 
                "baixa": e_baixa.to_dict("records"), 
                "dificuldades": e_dif.to_dict("records"), 
                "sugestoes": e_sug.to_dict("records")
            },
            "disc": respostas_disc_final  # <--- AGORA O NOME ESTÁ IGUAL AO QUE VOCÊ CRIOU
        }

        # --- CONFIGURAÇÃO DO NOME DO ARQUIVO (A CHAVE DO SUCESSO) ---
        nome_limpo = nome_input.strip().replace(" ", "_").upper()
        caminho_github = f"rascunhos/{nome_limpo}.json" # Garante que vai para a pasta rascunhos

        try:
            # --- AÇÃO 1: GITHUB (ENVIO REAL PARA O REPOSITÓRIO) ---
            conteudo_json = json.dumps(payload, indent=4, ensure_ascii=False)
            
            try:
                # Tenta atualizar se já existir
                f = repo.get_contents(caminho_github)
                repo.update_file(f.path, f"Update: {nome_input}", conteudo_json, f.sha)
            except:
                # Cria novo se não existir
                repo.create_file(caminho_github, f"Novo: {nome_input}", conteudo_json)
            
            # Salva no estado da sessão para uso imediato
            st.session_state["rascunho"] = payload
            st.success(f"✅ {nome_input} salvo com sucesso no GitHub!")

            # === COLOQUE AQUI O BOTÃO DE DOWNLOAD ===
            st.download_button(
                label="📥 Baixar Arquivo JSON do Rascunho",
                data=conteudo_json.encode('utf-8'), # Usa a variável que você já criou acima
                file_name=f"{nome_limpo}.json",
                mime="application/json",
                use_container_width=True
            )
            # =========================================

            # --- AÇÃO 2: SHEETS ---
            if enviar_para_sheets(payload):
                st.toast("📊 Espelhado no Sheets!", icon="📈")
            else:
                st.warning("⚠️ Salvo no GitHub, mas Sheets não respondeu.")

        except Exception as e:
            st.error(f"⚠️ Erro ao salvar: {e}")
            # Botão de emergência caso o GitHub falhe
            st.download_button(
                label="📥 Baixar Backup de Emergência",
                data=json.dumps(payload, indent=4, ensure_ascii=False).encode('utf-8'),
                file_name=f"{nome_limpo}_EMERGENCIA.json",
                mime="application/json",
                use_container_width=True
            )
    # Versao_Final_06_04


if st.session_state.pagina == "analise":

    
    # --- MOTOR DE AUDITORIA (INTEGRADO AO GITHUB) ---
    st.markdown("---")
    st.title("🚨 Auditoria de Gargalos, Nexo de Coerência e Ranking de Inovação")

    # O segredo: perguntar ao session_state evita o erro técnico "NameError"
    base = st.session_state.get('base_auditoria', [])
    t_check = st.session_state.get('t')

    if not base or t_check is None:
        st.info("Carregue os dados na seção 'Visualizar Dados' no Menu para ativar a auditoria. 🔴🔴🔴 ATENÇÃO!!! VOCÊ DEVE CLICAR PRIMEIRO NO BOTÃO 'Visualizar Dados' E DEPOIS VOLTAR E CLICAR DE NOVO NO BOTÃO 'Análise Inteligente' ")
    else:
        

    
        # --- RANKING E MÉTRICAS DE IMPACTO (SINCRONIZADO COM GERCINO) ---
        st.markdown("## 📈 Dashboard de Impacto Estratégico Global")

        ranking_dados = []
        total_geral_ano_horas = 0.0

        # Processamento com a Lógica Exata do Gercino
        for f in base:
            # Busca nome de forma robusta
            n_r = (f.get('colaborador') or f.get('nome') or "Desconhecido").upper()
            # Busca sugestões (Universal)
            s_r = f.get('tabelas', {}).get('sugestoes', []) or f.get('sugestoes', [])
            
            horas_colaborador = 0.0
            for s in s_r:
                try:
                    # 1. Limpeza e Conversão (Igual ao Individual)
                    h_s = float(str(s.get('Horas', '0')).lower().replace('h', '').replace(',', '.').strip() or 0)
                    m_s = float(str(s.get('Minutos', '0')).lower().replace('min', '').replace(',', '.').strip() or 0)
                    
                    # 2. Frequência Dinâmica (A regra que cravou os R$ 7.805,42)
                    f_s = str(s.get('Frequência', 'M')).upper().strip()
                    mult = {'D': 220, 'S': 48, 'M': 12, 'T': 4, 'A': 1}.get(f_s, 12)
                    
                    # 3. Cálculo Bruto Integral (Sem fator de 0.50)
                    horas_colaborador += (h_s + (m_s / 60)) * mult
                except: 
                    continue
            
            total_geral_ano_horas += horas_colaborador
            ranking_dados.append({"Colaborador": n_r, "Sug.": len(s_r), "Economia": horas_colaborador})

        if ranking_dados:
            # Valor hora técnica unificado em R$ 65,00
            total_financeiro = total_geral_ano_horas * 65.0
            
            col_c1, col_c2, col_c3 = st.columns(3)
            
            with col_c1:
                st.metric(
                    label="⚡ Eficiência Recuperável", 
                    value=f"{total_geral_ano_horas:,.1f} h/ano",
                    help="Soma das horas brutas declaradas por todos os colaboradores."
                )
            
            with col_c2:
                st.metric(
                    label="💰 ROI Projetado (EBITDA)", 
                    value=f"R$ {total_financeiro:,.2f}",
                    delta="Expectativa Bruta",
                    help="Valor total anualizado (R$ 65,00/h)."
                )
                
            with col_c3:
                # Ganho em dias (base padrão 8h)
                dias_novos = total_geral_ano_horas / 8
                st.metric(
                    label="📅 Ganho de Capacidade", 
                    value=f"{dias_novos:,.1f} Dias",
                    help="Dias de trabalho recuperados por ano."
                )

            st.markdown("---")

            st.subheader("🏆 Ranking de Inovação: Conversão em Valor/Ano")

            ranking_dados = []
            
            # --- MOTOR SINCRONIZADO: Aplica a lógica ULTRA em toda a base ---
            for f in base:
                n_colab = (f.get('colaborador') or f.get('nome') or "CONSULTOR").upper()
                # Busca sugestões tanto no padrão 'tabelas' quanto direto
                s_lista = f.get('tabelas', {}).get('sugestoes', []) or f.get('sugestoes', [])
                
                t_h_poupadas = 0.0
                t_rs_recuperavel = 0.0
                
                for sug in s_lista:
                    t_sug = str(sug.get('Sugestão', '')).lower()
                    if t_sug in ["nenhuma", "nada", "n/a", "", "nenhuma melhoria"]: 
                        continue

                    try:
                        # Limpeza Identica ao Motor Ultra (Gercino)
                        m_val = int(''.join(filter(str.isdigit, str(sug.get('Minutos', '0')))) or 0)
                        h_val = int(''.join(filter(str.isdigit, str(sug.get('Horas', '0')))) or 0)
                        t_min_total = (h_val * 60) + m_val
                        
                        # Classificação de Impacto Ultra
                        if any(w in t_sug for w in ['sistema', 'automacao', 'ia', 'integrar', 'digitalizar', 'api', 'robo', 'python']):
                            potencial = 0.85
                        elif any(w in t_sug for w in ['padronizar', 'checklist', 'treinamento', 'pop', 'manual', 'procedimento']):
                            potencial = 0.45
                        else:
                            potencial = 0.20
                        
                        # Frequência e Cálculo
                        f_sug = str(sug.get('Frequência', 'M')).upper().strip()
                        m_anual = {'D': 220, 'S': 48, 'M': 12, 'T': 4, 'A': 1}.get(f_sug, 12)
                        
                        h_ano_bruto = (t_min_total * m_anual) / 60
                        h_ano_poupada = h_ano_bruto * potencial
                        
                        t_h_poupadas += h_ano_poupada
                        t_rs_recuperavel += (h_ano_poupada * 65.0)
                    except:
                        continue
                
                if t_h_poupadas > 0:
                    ranking_dados.append({
                        "Colaborador": n_colab,
                        "Qtd": len(s_lista),
                        "H_Recup": t_h_poupadas,
                        "ROI_Final": t_rs_recuperavel
                    })

            # Gerar DataFrame e ordenar pelo Valor Real
            df_r = pd.DataFrame(ranking_dados).sort_values(by="ROI_Final", ascending=False)

            # --- EXIBIÇÃO DOS CARDS GLOBAIS ---
            c_g1, c_g2 = st.columns(2)
            c_g1.metric("ROI Real Auditado (Global)", f"R$ {df_r['ROI_Final'].sum():,.2f}")
            c_g2.metric("Eficiência Recuperada", f"{df_r['H_Recup'].sum():,.1f} h/ano")

            # --- TABELA DE RANKING FORMATADA ---
            st.dataframe(
                df_r,
                column_config={
                    "Colaborador": st.column_config.TextColumn("Colaborador"),
                    "Qtd": st.column_config.NumberColumn("Sug."),
                    "H_Recup": st.column_config.NumberColumn("Horas/Ano", format="%.1f h"),
                    "ROI_Final": st.column_config.NumberColumn("Valor Auditado", format="R$ %.2f"),
                },
                column_order=("Colaborador", "Qtd", "H_Recup", "ROI_Final"),
                hide_index=True,
                use_container_width=True
            )

            st.caption("🛡️ **Auditoria Pericial:** Ranking sincronizado com o Motor Ultra (85% | 45% | 20%).")

            # --- LEGENDA TÉCNICA ATUALIZADA ---
            st.caption("""
            🔬 **Metodologia de Auditoria Pericial:** Projeções fundamentadas no Custo Total de Ocupação (R$ 65,00/h). 
            O sistema aplica Ponderação Dinâmica: 🤖 Transformação Digital (85%), 📈 Otimização (45%) 
            ou 💡 Incremental (20%), garantindo que o ROI reflita a complexidade técnica da entrega.
            """)

        st.markdown("---")

        # --- MOTOR DE AUDITORIA ---
        mapa_auditoria = {}
        for idx, f in enumerate(base):
            campos = f.get('campos', {}) if isinstance(f.get('campos'), dict) else {}
            n_extraido = (f.get('colaborador') or f.get('nome') or campos.get('nome') or f'Colaborador {idx}')
            nome_chave = str(n_extraido).upper().strip()
            mapa_auditoria[nome_chave] = f
    
        

        # --- MOTOR DE AUDITORIA (PROCESSAMENTO) ---
        mapa_auditoria = {}
        for idx, f in enumerate(base):
            campos = f.get('campos', {}) if isinstance(f.get('campos'), dict) else {}
            n_extraido = (f.get('colaborador') or f.get('nome') or campos.get('nome') or f'Colaborador {idx}')
            nome_chave = str(n_extraido).upper().strip()
            mapa_auditoria[nome_chave] = f

        st.markdown("---")
        # --- FIM DO RANKING ---

        # Se os dados existem na memória, o motor processa normalmente
        mapa_auditoria = {}
        for idx, f in enumerate(base):
            campos = f.get('campos', {}) if isinstance(f.get('campos'), dict) else {}
            n_extraido = (f.get('colaborador') or f.get('nome') or campos.get('nome') or f'Colaborador {idx}')
            nome_chave = str(n_extraido).upper().strip()
            mapa_auditoria[nome_chave] = f

        nomes_disponiveis = sorted(list(mapa_auditoria.keys()))
        colab_alvo = st.selectbox(f"🎯 Selecione para Auditoria ({len(nomes_disponiveis)} encontrados):", nomes_disponiveis)

        dados_alvo = mapa_auditoria[colab_alvo]
        t = dados_alvo.get('tabelas', {})
        
        # --- FUNÇÃO DE AUDITORIA (MATEMÁTICA CORRIGIDA S/5, M/20) ---
        def auditar_tabela_v2(lista):
            total_dia = 0.0
            detalhes = []
            for i in lista:
                try:
                    desc = i.get('Atividade') or i.get('Dificuldade') or i.get('Sugestão') or "Item"
                    h = float(str(i.get('Horas', '0')).lower().replace('h', '').replace(',', '.').strip() or 0)
                    m = float(str(i.get('Minutos', '0')).lower().replace('min', '').replace(',', '.').strip() or 0)
                    f = str(i.get('Frequência', 'D')).upper().strip()
                    
                    # Divisores exatos: Semanal por 5, Mensal por 20
                    divisores = {'D': 1, 'S': 5, 'M': 20, 'T': 60, 'A': 240}
                    divisor = divisores.get(f, 1)
                    
                    valor_diario = (h + (m / 60)) / divisor
                    total_dia += valor_diario
                    
                    detalhes.append({
                        "Descrição": desc,
                        "Relatado": f"{int(h)}h {int(m)}min ({f})",
                        "Impacto Real (Dia)": f"{valor_diario:.3f} h/dia"
                    })
                except: continue
            return total_dia, detalhes

        # --- PROCESSAMENTO DOS SUBTOTAIS (Identado dentro do if base_auditoria) ---
        h_alta, det_alta = auditar_tabela_v2(t.get('alta', []))
        h_norm, det_norm = auditar_tabela_v2(t.get('normal', []))
        h_baix, det_baix = auditar_tabela_v2(t.get('baixa', []))
        h_dif, det_dif = auditar_tabela_v2(t.get('dificuldades', []))
        
        h_total = h_alta + h_norm + h_baix + h_dif

        # --- MOTOR DE NEXO CAUSAL (ESCALA UNIFORME E SUAVE) ---
        h_calc = float(h_total) if h_total else 0.0
        jornada_referencia = 8.0
        
        if h_calc == 0:
            score = 0
        else:
            # Distância absoluta de 8h
            desvio = abs(h_calc - jornada_referencia)
            
            # Usando multiplicador de 6 pontos por hora de desvio.
            # Isso permite que mesmo alguém com 18h ainda tenha um score visível.
            score = 100 - (desvio * 6)

        # Trava o score entre 0 e 100
        score = max(0, min(100, float(score)))

        import plotly.graph_objects as go
        fig = go.Figure(go.Indicator(
            mode = "gauge+number", 
            value = score,
            title = {'text': f"Nexo Causal: {h_total:.2f}h/dia"},
            gauge = {
                'axis': {'range': [0, 100]},
                'bar': {'color': "#2E3192"},
                'steps': [
                    {'range': [0, 45], 'color': "#ff4b4b"},
                    {'range': [45, 80], 'color': "#ffa500"},
                    {'range': [80, 100], 'color': "#00c853"}
                ]
            }
        ))
        fig.update_layout(height=250, margin=dict(l=20, r=20, t=50, b=20))
        st.plotly_chart(fig, use_container_width=True)

        # --- DETALHAMENTO POR CATEGORIA (IDENTAÇÃO CORRIGIDA) ---
        st.subheader("📋 Detalhamento por Categoria")

        # 1. Pegamos os nomes das dificuldades para servir de filtro
        nomes_dificuldades = [str(d.get('Dificuldade', '')).strip().lower() for d in t.get('dificuldades', [])]

        secoes = [
            ("🔴 Alta Complexidade", det_alta, h_alta),
            ("🟡 Complexidade Normal", det_norm, h_norm),
            ("🟢 Baixa Complexidade", det_baix, h_baix),
            ("⚠️ Dificuldades", det_dif, h_dif)
        ]

        for titulo, dados, subtotal in secoes:
            with st.expander(f"{titulo} (Total: {subtotal:.2f}h/dia)"):
                if dados:
                    # 2. FILTRAGEM EM TEMPO REAL:
                    # Se não for a seção de Dificuldades, removemos o que for lixo ou duplicata
                    if "Dificuldades" not in titulo:
                        dados_limpos = [
                            d for d in dados 
                            if str(d.get('Descrição', '')).strip().lower() not in ["vazio", "vazio...", "", "none", "."]
                            and str(d.get('Descrição', '')).strip().lower() not in nomes_dificuldades
                        ]
                        if dados_limpos:
                            st.table(dados_limpos)
                        else:
                            st.write("Sem atividades relevantes nesta categoria.")
                    else:
                        # Na seção de Dificuldades, mostra tudo normal
                        st.table(dados)
                else:
                    st.write("Sem registros.")
                
        else:
            # Este else volta para a margem zero (alinhado com o IF inicial)
            st.info("💡 Por favor, carregue os dados na Visualização de Registros para ativar a auditoria.")

        def auditoria_super_inteligente(tabelas_dict, h_total_calculada):
            checklist = []
            
            # --- DICIONÁRIO DE PESOS SEMÂNTICOS ---
            pesos = {
                "estrategico": ["gerenciar", "estratégico", "implementar", "planejar", "reestruturação", "onboarding", "diretrizes"],
                "tecnico_pesado": ["auditoria", "fundamentação", "diagnosticar", "mitigação", "validar", "conduzir"],
                "operacional": ["atualizar", "organizar", "baixar", "enviar", "lembretes", "conferir checklists", "separar", "status"]
            }

            for cat, itens in tabelas_dict.items():
                for item in itens:
                    desc_pura = (item.get('Atividade') or item.get('Dificuldade') or "Vazio")
                    desc = desc_pura.lower()
                    freq = str(item.get('Frequência', 'D')).upper().strip()
                    
                    h = float(str(item.get('Horas', '0')).lower().replace('h', '').replace(',', '.').strip() or 0)
                    m = float(str(item.get('Minutos', '0')).lower().replace('min', '').replace(',', '.').strip() or 0)
                    
                    divisores = {'D': 1, 'S': 5, 'M': 20, 'T': 60, 'A': 240}
                    divisor = divisores.get(freq, 1)
                    tempo_execucao = h + (m / 60)
                    impacto_diario = tempo_execucao / divisor

                    alertas = []

                    # --- RIGOR 1: PADRONIZAÇÃO SUSPEITA (O EFEITO "1 HORA") ---
                    # Se o cara coloca 1h em tudo, ele não parou para pensar, ele apenas preencheu para acabar logo.
                    if h == 1 and m == 0:
                        alertas.append("Lançamento Padronizado (1h): Indica falta de precisão no preenchimento.")

                    # --- RIGOR 2: INCOERÊNCIA DE NATUREZA ---
                    # Atividades estratégicas/pesadas NÃO levam o mesmo tempo que operacionais.
                    if any(p in desc for p in pesos["estrategico"]) and tempo_execucao <= 1.0:
                        alertas.append("Complexidade Estratégica: Tempo relatado insuficiente para a profundidade da tarefa.")
                    
                    if any(p in desc for p in pesos["operacional"]) and tempo_execucao >= 1.0:
                        alertas.append("Inchaço Operacional: Tarefa simples consumindo muito tempo por execução.")

                    # --- RIGOR 3: CONTRADIÇÃO DE CATEGORIA ---
                    if cat == 'alta' and impacto_diario < 0.2:
                        alertas.append("Classificada como 'Alta', mas o impacto é de tarefa de suporte.")

                    # --- RIGOR 4: TAREFAS VAZIAS OU GENÉRICAS ---
                    if "vazio" in desc or len(desc_pura) < 5:
                        alertas.append("Descrição insuficiente para auditoria.")

                    status = "✅" if not alertas else "❌"
                    veredito = "Coerente" if not alertas else " | ".join(alertas)
                    
                    checklist.append({
                        "Status": status,
                        "Atividade": desc_pura,
                        "Impacto": f"{impacto_diario:.3f} h/dia",
                        "Análise Crítica": veredito
                    })

            return checklist

        
        
        # --- NOVO PARECER DE EXTRAPOLAÇÃO (FILTRADO E IDENTADO) ---
        # 1. Mapeia nomes de dificuldades para exclusão
        nomes_difs = [str(d.get('Dificuldade', '')).strip().lower() for d in t.get('dificuldades', [])]

        # 2. Cria dicionário limpo apenas com Atividades reais
        t_limpo = {
            cat: [
                i for i in t.get(cat, []) 
                if str(i.get('Atividade','')).strip().lower() not in nomes_difs 
                and str(i.get('Atividade','')).strip().lower() not in ["vazio", "vazio...", "", "none", "."]
            ] for cat in ['alta', 'normal', 'baixa']
        }

        # 3. Gera a tabela de Status/Análise Crítica sem as duplicatas
        res_final = auditoria_super_inteligente(t_limpo, h_total)
        
        if res_final:
            st.markdown("""
                <style>
                    div[data-testid="stTable"] {
                        width: fit-content !important;
                        margin: 0 auto !important;
                    }
                    /* ESTA É A MÁGICA: Espreme a segunda coluna (Atividade) */
                    table tr td:nth-child(2) {
                        max-width: 300px !important; 
                        white-space: normal !important;
                        word-wrap: break-word !important;
                    }
                    th, td {
                        padding: 3px 10px !important; 
                        font-size: 13px !important;
                    }
                </style>
            """, unsafe_allow_html=True)
            
            st.table(res_final)

        st.markdown("### 📝 Parecer do Perito Digital")
        
        # Lógica de análise de volume
        if h_total > 12:
            st.error(f"🚨 **INVIABILIDADE MATEMÁTICA:** O colaborador relata **{h_total:.2f}h** de trabalho por dia. É fisicamente impossível manter essa carga com qualidade. O formulário foi preenchido com superestimação de tempos ou erro crasso de frequência.")
        elif h_total > 9:
            st.warning(f"⚠️ **SOBRECARGA DETECTADA:** Carga de **{h_total:.2f}h/dia**. Excede a jornada legal e indica risco de burnout ou erros técnicos.")



    # --- MOTOR DE PERÍCIA DE DIFICULDADES (LINHA DURA - SEM FILTROS) ---
    st.markdown("---")
    st.subheader("⚠️ Auditoria de Gargalos e Nexo de Coerência")

    def analisar_dificuldades_rigoroso(dificuldades_lista, tabelas_dict, h_total_atividades):
        check_dif = []
        
        # Consolida atividades para cruzamento semântico
        texto_atividades = ""
        # 1. Cria lista de nomes das dificuldades para exclusão
        nomes_dificuldades = [str(d.get('Dificuldade', '')).strip().lower() for d in dificuldades_lista]

        for cat in ['alta', 'normal', 'baixa']:
            if cat in tabelas_dict:
                # FILTRAGEM ATIVA: Remove vazios e remove o que for dificuldade
                tabelas_dict[cat] = [
                    item for item in tabelas_dict[cat] 
                    if str(item.get('Atividade', '')).strip().lower() not in ["vazio", "vazio...", "", "none", "."]
                    and str(item.get('Atividade', '')).strip().lower() not in nomes_dificuldades
                ]
                
                # 2. Concatena apenas as atividades que sobraram
                for item in tabelas_dict[cat]:
                    texto_atividades += str(item.get('Atividade', '')).lower() + " "

        if not dificuldades_lista:
            return []

        for d in dificuldades_lista:
            # 1. Coleta de dados com limpeza
            desc_pura = str(d.get('Dificuldade') or d.get('Sugestão') or "Vazio")
            desc = desc_pura.lower()
            setor = str(d.get('Setor Envolvido') or "N/A").upper()
            
            # 2. Cálculo de Impacto (Rigor nos Divisores)
            h_d = float(str(d.get('Horas', '0')).lower().replace('h', '').replace(',', '.').strip() or 0)
            m_d = float(str(d.get('Minutos', '0')).lower().replace('min', '').replace(',', '.').strip() or 0)
            freq_raw = str(d.get('Frequência', 'M')).upper().strip()
            
            divisores = {'D': 1, 'S': 5, 'M': 20, 'T': 60, 'A': 240}
            divisor = divisores.get(freq_raw, 20)
            impacto_diario = (h_d + (m_d / 60)) / divisor

            alertas = []

            # --- TESTE 1: NEXO DE IMPACTO ---
            if impacto_diario < 0.05: # Menos de 3 min/dia
                alertas.append("Irrelevante: Impacto temporal muito baixo para ser considerado gargalo.")

            # --- TESTE 2: CONTRADIÇÃO TEMPORAL ---
            if any(p in desc for p in ["frequente", "constante", "sempre", "todo dia"]):
                if freq_raw not in ['D', 'S']:
                    alertas.append(f"Incoerência: Relata ser 'constante' mas a frequência é {freq_raw}.")

            # --- TESTE 3: MATRIZ DE CORRELAÇÃO SEMÂNTICA (AJUSTADO) ---
            matriz_nexo = {
                "retrabalho": ["conferência", "lançamento", "cálculo", "análise", "ajuste", "correção", "fechamento", "erro", "inconsistência"],
                "conferência": ["apuração", "análise", "conferir", "checklist", "verificação", "validação", "divergência", "detalhada"],
                "acúmulo": ["fechamento", "envio", "atendimento", "cadastro", "lançamento", "gestão", "operacional", "reduz"],
                "padronização": ["processos", "fluxos", "organização", "método", "rotina", "interno"],
                "legislação": ["aplicação", "consultivo", "clt", "convenção", "sindicato", "fiscal", "alteração"],
                "sistema": ["alterdata", "esocial", "plataforma", "instabilidade", "lento", "limitação"],
                "comunicação": ["atendimento", "esclarecimento", "dúvidas", "internas", "externas", "desafio"],
                "informação": ["cadastro", "atualização", "dados", "planilha", "setor", "dependência"]
            }

            achou_nexo = False
            
            # 1. Contagem de volume para validar 'Acúmulo' automaticamente
            total_atividades = sum(len(tabelas_dict.get(cat, [])) for cat in ['alta', 'normal', 'baixa'])

            # 2. Lógica de busca na Matriz
            for palavra_chave, correlatas in matriz_nexo.items():
                if palavra_chave in desc:
                    # Se a palavra-chave OU qualquer correlata dela estiver nas atividades
                    if any(corr in texto_atividades for corr in correlatas) or palavra_chave in texto_atividades:
                        achou_nexo = True
                        break
                    
            # 3. Regra de Ouro para 'Acúmulo' e 'Volume'
            if any(p in desc for p in ["acúmulo", "volume", "quantidade", "muita"]):
                if total_atividades > 12: # Se tem mais de 12 tarefas, o nexo de acúmulo é real
                    achou_nexo = True

            # 4. Validação de Rigor de Tempo vs. Frequência
            if impacto_diario > 0.8 and len(desc_pura) < 20:
                alertas.append("Subdetalhamento: Impacto crítico para uma descrição tão curta.")

            if not achou_nexo and len(desc) > 15:
                if not any(x in desc for x in ["equipe", "gestão", "processo", "comunicação", "demanda"]):
                    alertas.append("Desconexão: Esta dificuldade não possui correlação técnica com as tarefas listadas.")

            # --- TESTE 4: RECLAMAÇÃO DE VOLUME ---
            if any(p in desc for p in ["acúmulo", "sobrecarga", "volume"]):
                if h_total_atividades < 7.0:
                    alertas.append(f"Falso Alerta: Reclama de volume, mas o Nexo de Atividades é de apenas {h_total_atividades:.2f}h.")

            # Montagem do Resultado
            status = "🚩" if alertas else "✅"
            veredito = "Nexo Causal Confirmado" if not alertas else " | ".join(alertas)

            check_dif.append({
                "Status": status,
                "Setor": setor,
                "Dificuldade": desc_pura[:1000] + "...",
                "Impacto Diário": f"{impacto_diario:.3f} h/dia",
                "Análise do Perito": veredito
            })
        
        return check_dif

    # --- CHAMADA E EXIBIÇÃO ---

    # 1. Verifica se 't' existe e se é um dicionário antes de tentar o .get()
    t_valida = locals().get('t')

    if isinstance(t_valida, dict):
        # Agora é seguro usar o .get()
        lista_dif = locals().get('t', {}).get('dificuldades', [])
        h_v = locals().get('h_total', 0)

        res_dificuldades = analisar_dificuldades_rigoroso(lista_dif, t_valida, h_v)

        if res_dificuldades:
            st.table(res_dificuldades)
        else:
            st.info("ℹ️ Nenhuma dificuldade encontrada para este colaborador.")
    else:
        # Se 't' não existe ou não é dicionário, mostra o seu alerta amarelo
        st.info("⚠️ ATENÇÃO ACIMA ☝️")


import pandas as pd
import streamlit as st
 
def motor_pericia_ultra(tabelas, dificuldades, sugestoes):
    # Consolida contexto das atividades para analise
    todas_atv = tabelas.get('alta', []) + tabelas.get('normal', []) + tabelas.get('baixa', [])
    contexto_atv = " ".join([a.get('Atividade', '').lower() for a in todas_atv])
     
    analise_detalhada = []
 
    for sug in sugestoes:
        texto_sug = str(sug.get('Sugestão', '')).lower()
        if texto_sug in ["nenhuma", "nada", "n/a", "", "nenhuma melhoria"]: 
            continue
 
        freq = sug.get('Frequência', 'D').upper()
        # Limpeza robusta de strings para numeros
        m = int(''.join(filter(str.isdigit, str(sug.get('Minutos', '0')))) or 0)
        h = int(''.join(filter(str.isdigit, str(sug.get('Horas', '0')))) or 0)
        tempo_min_atual = (h * 60) + m
         
        # --- INTELIGENCIA DE CLASSIFICACAO ---
        if any(w in texto_sug for w in ['sistema', 'automacao', 'ia', 'integrar', 'digitalizar', 'api', 'robo', 'python']):
            potencial = 0.85
            categoria = "TRANSFORMACAO DIGITAL"
            cor_status = "ALTO IMPACTO"
        elif any(w in texto_sug for w in ['padronizar', 'checklist', 'treinamento', 'pop', 'manual', 'procedimento']):
            potencial = 0.45
            categoria = "OTIMIZACAO DE PROCESSO"
            cor_status = "ESTRUTURAL"
        else:
            potencial = 0.20
            categoria = "MELHORIA INCREMENTAL"
            cor_status = "OPERACIONAL"
 
        # --- ENGENHARIA DE VALOR ---
        mult = {'D': 220, 'S': 48, 'M': 12, 'T': 4, 'A': 1}.get(freq, 1)
        h_ano_atual = (tempo_min_atual * mult) / 60
        h_poupadas = h_ano_atual * potencial
        valor_financeiro = h_poupadas * 65.0 
 
        analise_detalhada.append({
            "ESTRATEGIA": categoria,
            "SUGESTAO ANALISADA": sug.get('Sugestão', '').upper(),
            "H_FLOAT": h_poupadas,
            "RS_FLOAT": valor_financeiro,
            "ECONOMIA PROJETADA": f"- {h_poupadas:.1f} h/ano",
            "VALOR RECUPERAVEL": f"R$ {valor_financeiro:,.2f}",
            "PARECER": cor_status
        })
     
    return pd.DataFrame(analise_detalhada)
 
# --- EXIBICAO NO DASHBOARD ---
if st.session_state.get("pagina") == "analise":
    st.markdown("---")
 
    with st.status("Processando analise pericial...", expanded=True):
        st.header("🔬 Central de Inteligencia e Auditoria de Processos")
 
        # Busca t_base de forma segura (contexto local ou session)
        t_base = locals().get('t') or st.session_state.get('t_selecionado')
 
        if isinstance(t_base, dict):
            sug_lista = t_base.get('sugestoes', [])
             
            if sug_lista:
                # Verificacao de Engajamento
                sug_primeira = str(sug_lista[0].get('Sugestão', '')).lower().strip()
                if sug_primeira in ["nenhuma", "nada", "nenhuma melhoria", "", "n/a"]:
                    st.error("🚨 ALERTA DE GESTAO: O colaborador optou por nao sugerir aperfeicoamentos.")
                else:
                    st.subheader(f"📊 Business Case: {t_base.get('colaborador', 'Consultor')}")
                     
                    df_analise = motor_pericia_ultra(t_base, [], sug_lista)
                     
                    if not df_analise.empty:
                        # --- CALCULOS TOTAIS ---
                        total_h_ano = df_analise['H_FLOAT'].sum()
                        total_valor = df_analise['RS_FLOAT'].sum()
                        
                        # PERSISTÊNCIA NA SESSÃO (PARA O CARD FINAL)
                        st.session_state['v_audit_final'] = total_valor
                        st.session_state['h_audit_final'] = total_h_ano
 
                        # metricas principais
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Capacidade Recuperada", f"{total_h_ano:.1f} h/ano")
                        c2.metric("ROI Operacional Est.", f"R$ {total_valor:,.2f}")
                        c3.metric("Impacto em Dias", f"{total_h_ano/8:.1f} dias")
 
                        # --- CARD DE VIABILIDADE PERICIAL (CALIBRADO) ---
                        st.markdown("---")
                        st.subheader("🛡️ Verificação de Viabilidade Pericial")
                        
                        # --- INÍCIO DA CORREÇÃO ---
                        # 1. Primeiro, definimos quem são as sugestões (evita o NameError)
                        sugestoes_lista = t_base.get('sugestoes', []) if 't_base' in locals() else t.get('sugestoes', [])
                        
                        # 2. Calculamos o valor BRUTO sem descontos
                        v_bruto_final = 0.0
                        for s in sugestoes_lista:
                            try:
                                # Limpa o texto das horas e minutos
                                h_limpo = float(str(s.get('Horas', '0')).lower().replace('h','').strip() or 0)
                                m_limpo = float(str(s.get('Minutos', '0')).lower().replace('min','').strip() or 0)
                                
                                # Pega a frequência e define o multiplicador anual
                                f_tipo = str(s.get('Frequência', 'M')).upper().strip()
                                m_anual = {'D': 220, 'S': 48, 'M': 12, 'T': 4, 'A': 1}.get(f_tipo, 12)
                                
                                # Soma no Bruto (R$ 65/hora)
                                v_bruto_final += ((h_limpo + (m_limpo / 60)) * m_anual * 65.0)
                            except:
                                continue

                        # 3. Exibimos os cards na tela
                        c_bruto, c_roi = st.columns(2)
                        with c_bruto:
                            st.metric("📢 Expectativa (Bruto)", f"R$ {v_bruto_final:,.2f}")
                            st.caption("Soma integral das sugestões.")

                        with c_roi:
                            # Tenta pegar o total_valor que seu motor calculou
                            v_final_auditado = locals().get('total_valor', 0.0)
                            st.metric("💎 ROI Real Auditado", f"R$ {v_final_auditado:,.2f}")
                            st.caption("Valor aprovado após auditoria.")
                        # --- FIM DA CORREÇÃO ---
                        
                        # --- PARECER DETALHADO DO PERITO ---
                        st.markdown("#### 📝 Parecer Técnico de Viabilidade")
                        
                        # 1. Cálculo do Ajuste (Evita NameError e divisão por zero)
                        # Comparamos o ROI Real (total_valor) com a Expectativa Bruta (v_bruto)
                        ajuste = ((v_final_auditado / v_bruto_final) - 1) * 100 if v_bruto_final > 0 else 0

                        # 2. Lógica de Parecer Dinâmico
                        if ajuste > -30:
                            st.success("✅ **VIABILIDADE ALTA:** As sugestões possuem alto índice de aproveitamento técnico (85%+). O risco de implementação é baixo e o retorno é estrutural.")
                        elif ajuste > -60:
                            st.warning("⚠️ **VIABILIDADE MODERADA:** Requer ajustes de processos manuais antes da automação. ROI garantido após saneamento de dados.")
                        else:
                            st.error("📉 **VIABILIDADE CRÍTICA:** Muitas tarefas operacionais de baixo valor agregado. Foco deve ser em eliminar a tarefa, não em automatizar.")

                        st.info(f"💡 **Nota do Perito:** Aplicado multiplicador de frequência dinâmico e custo técnico de R$ 65,00/h.")

                        # --- TABELA FINAL ---
                        st.markdown("### 📋 Detalhamento das Oportunidades")
                        
                        # Tratamento seguro das colunas da tabela
                        if 'df_analise' in locals():
                            cols_remover = [c for c in ['H_FLOAT', 'RS_FLOAT'] if c in df_analise.columns]
                            st.table(df_analise.drop(columns=cols_remover))
                        else:
                            st.error("Erro ao carregar a tabela de análise.")





# =================================================================
# --- BLOCO INTEGRAL: GERAÇÃO DO LAUDO PERICIAL 2026 (FINAL) ---
# =================================================================

if st.session_state.get("pagina") == "analise":

    # 1. INICIALIZAÇÃO DE SEGURANÇA E FONTES
    linhas_atividades_html = ""
    linhas_gargalos_html = ""
    ranking_final_html = ""
    linhas_sugestoes_html = ""
    horas_totais_ano = ganho_capacidade_dias = roi_real_auditado = 0.0

    # 2. CÁLCULO DE MÉTRICAS ESTRATÉGICAS (INDIVIDUALIZADO)
    if 'ranking_dados' in locals() and ranking_dados:
        df_rank_calc = pd.DataFrame(ranking_dados)
        colab_atual = colab_alvo if 'colab_alvo' in locals() else "CONSOLIDADO"
        
        if colab_atual != "CONSOLIDADO":
            df_perito = df_rank_calc[df_rank_calc['Colaborador'] == colab_atual]
        else:
            df_perito = df_rank_calc

        horas_totais_ano = df_perito['H_Recup'].sum()
        roi_real_auditado = df_perito['ROI_Final'].sum()
        ganho_capacidade_dias = horas_totais_ano / 8

    # 3. EXTRAÇÃO DE DIFICULDADES E ATIVIDADES (NEXO CAUSAL)
    
    # --- Atividades (Integração com res_final) ---
    if 'res_final' in locals() and res_final:
        for item in res_final:
            status_ativ = item.get('Status', '✅')
            cor_st = "#52c41a" if status_ativ == "✅" else "#ff4d4f"
            ativ_nome = item.get('Atividade', 'N/A')
            impacto_ativ = item.get('Impacto', 'N/A')
            analise_ativ = item.get('Análise Crítica', 'N/A')
            
            linhas_atividades_html += f"""
            <tr>
                <td style='color:{cor_st}; text-align:center;'>{status_ativ}</td>
                <td>{ativ_nome}</td>
                <td>{impacto_ativ}</td>
                <td>{analise_ativ}</td>
            </tr>"""

    # --- Gargalos (Sincronizado com o Motor Linha Dura) ---
    dificuldades_lista_processada = res_dificuldades if 'res_dificuldades' in locals() else []
    gargalos_vistos = set()
    
    for d in dificuldades_lista_processada:
        dif_txt = d.get('Dificuldade', 'N/A').replace('...', '')
        if dif_txt not in gargalos_vistos and dif_txt.lower() not in ['n/a', 'nenhuma', 'nenhuma dificuldade']:
            gargalos_vistos.add(dif_txt)
            
            status_icone = d.get('Status', '⚠️')
            setor_dif = d.get('Setor', 'OPERACIONAL')
            analise_perito = d.get('Análise do Perito', 'Nexo Causal Confirmado')
            impacto_txt = d.get('Impacto Diário', '0.000 h/dia')
            
            linhas_gargalos_html += f"""
            <tr>
                <td style='text-align:center;'>{status_icone}</td>
                <td>{setor_dif}</td>
                <td>{dif_txt}</td>
                <td><b>{analise_perito}</b> (Carga: {impacto_txt})</td>
            </tr>"""

    # 4. GERAÇÃO DAS TABELAS DE SUGESTÕES E RANKING
    if 'df_analise' in locals() and not df_analise.empty:
        for _, row in df_analise.iterrows():
            linhas_sugestoes_html += f"""
            <tr>
                <td>{row.get('ESTRATEGIA')}</td>
                <td>{row.get('SUGESTAO ANALISADA')}</td>
                <td style='text-align:center;'>{row.get('ECONOMIA PROJETADA')}</td>
                <td style='text-align:right;'><b>{row.get('VALOR RECUPERAVEL')}</b></td>
            </tr>"""

    if 'ranking_dados' in locals() and ranking_dados:
        df_ranking_auditado = df_rank_calc.sort_values(by="ROI_Final", ascending=False)
        for i, (_, row) in enumerate(df_ranking_auditado.iterrows(), 1):
            cor_pos = "#FFD700" if i == 1 else "#1B1E5D"
            ranking_final_html += f"""
            <tr>
                <td style='text-align:center; font-weight:bold; color:{cor_pos};'>{i}º</td>
                <td><b>{row['Colaborador']}</b></td>
                <td style='text-align:center;'>{row['Qtd']}</td>
                <td style='text-align:right;'>{row['H_Recup']:.1f}h</td>
                <td style='text-align:right; font-weight:bold;'>R$ {row['ROI_Final']:,.2f}</td>
            </tr>"""


    # SEGURANÇA ANTIFALHA:
    colab_nome_exibicao = colab_atual if 'colab_atual' in locals() else "COLABORADOR"

    # 5. MONTAGEM DA ESTRUTURA HTML FINAL
    # Alinhado com o 'if' anterior
    t_base_safe = t_base if isinstance(t_base, dict) else {}
    nome_para_exibicao = locals().get('colab_atual') or t_base_safe.get('colaborador') or "COLABORADOR"

    html_final = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset='utf-8'>
        <style>
            body {{ font-family: sans-serif; padding: 30px; color: #333; }}
            .header-banner {{ background: #1B1E5D; color: white; padding: 20px; text-align: center; border-radius: 10px; }}
            .container-metrics {{ display: flex; gap: 10px; margin: 20px 0; }}
            .metric-box {{ background: #f4f4f4; padding: 15px; flex: 1; text-align: center; border-bottom: 4px solid #1B1E5D; border-radius: 8px; }}
            .section-title {{ background: #1B1E5D; color: white; padding: 8px; margin-top: 25px; font-weight: bold; border-radius: 4px; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
            th {{ background: #eee; padding: 10px; text-align: left; border-bottom: 2px solid #1B1E5D; }}
            td {{ padding: 10px; border-bottom: 1px solid #ddd; font-size: 12px; }}
        </style>
    </head>
    <body>
        <div class='header-banner'>
            <h1>LAUDO PERICIAL: AUDITORIA ESTRATEGICA</h1>
            <h2 style='text-transform: uppercase;'>{nome_para_exibicao}</h2>
        </div>

        <div class='container-metrics'>
            <div class='metric-box'><b>ROI AUDITADO</b><br>R$ {roi_real_auditado:,.2f}</div>
            <div class='metric-box'><b>EFICIENCIA</b><br>{horas_totais_ano:.1f} h/ano</div>
            <div class='metric-box'><b>LIBERACAO</b><br>{ganho_capacidade_dias:.1f} Dias</div>
        </div>

        <div class='section-title'>OPORTUNIDADES E SUGESTOES</div>
        <table>
            <thead><tr><th>ESTRATEGIA</th><th>SUGESTAO</th><th>ECONOMIA</th><th>VALOR</th></tr></thead>
            <tbody>{linhas_sugestoes_html}</tbody>
        </table>

        <div class='section-title'>AUDITORIA DE NEXO CAUSAL</div>
        <table>
            <thead><tr><th>ST</th><th>ATIVIDADE</th><th>IMPACTO</th><th>ANALISE</th></tr></thead>
            <tbody>{linhas_atividades_html}</tbody>
        </table>

        <div class='section-title'>GARGALOS IDENTIFICADOS</div>
        <table>
            <thead><tr><th>ST</th><th>SETOR</th><th>DIFICULDADE</th><th>ANALISE PERITO</th></tr></thead>
            <tbody>{linhas_gargalos_html}</tbody>
        </table>

        <div class='section-title'>RANKING PERFORMANCE</div>
        <table>
            <thead><tr><th>POS</th><th>COLABORADOR</th><th>SUG.</th><th>HORAS</th><th>ROI</th></tr></thead>
            <tbody>{ranking_final_html}</tbody>
        </table>
    </body>
    </html>
    """

    # 6. EXIBIÇÃO DO BOTÃO (RESOLVE NAMEERROR DEFINITIVAMENTE)
    st.divider()
    
    # Usa a variável segura 'nome_para_exibicao' que já foi validada no Bloco 5
    nome_seguro_arquivo = nome_para_exibicao.replace(" ", "_")
    
    st.download_button(
        label=f"📥 BAIXAR LAUDO PERICIAL: {nome_para_exibicao}",
        data=html_final,
        file_name=f"Laudo_Pericial_{nome_seguro_arquivo}.html",
        mime="text/html",
        use_container_width=True,
        key="btn_laudo_final_2026"
    )


# ============================================================
# 🚀 COMANDO MASTER DO LUCIANO: PROCESSAMENTO, SOMA E DIAGNÓSTICO
# ============================================================

if st.session_state.pagina == "disc":
    st.title("📊 Panorama Estratégico da Equipe")
    st.subheader(f"Operador: Luciano")

    # BOTÃO MASTER: O TIRO ÚNICO
    if st.button("🔥 GERAR E SOMAR TODOS OS PERFIS DA EQUIPE", key="btn_master_soma_coletiva"):
        
        with st.spinner("Sincronizando com GitHub e processando DNA da equipe..."):
            try:
                # 1. Garante a Conexão com o Repo
                repo = st.session_state.get('repo_conectado')
                if repo is None:
                    g = Github(st.secrets["DB_TOKEN"])
                    repo = g.get_repo("lucianohcl/formulario-colaborador")
                    st.session_state.repo_conectado = repo

                # 2. Carrega todos os arquivos da pasta /dados/
                lista_fresca = carregar_todos_formularios(repo)
                
                if not lista_fresca:
                    st.error("❌ Nenhum formulário encontrado no repositório.")
                else:
                    resultados_acumulados = []
                    
                    # Mapa de tradução DISC
                    mapa_disc = {"A": "D", "B": "I", "C": "S", "D": "C"}

                    # 3. LOOP DE PROCESSAMENTO
                    for f in lista_fresca:
                        respostas_raw = f.get("disc", {})
                        if respostas_raw:
                            # Traduz as respostas A,B,C,D para D,I,S,C
                            respostas_traduzidas = {k: mapa_disc[v] for k, v in respostas_raw.items() if v in mapa_disc}
                            
                            # Calcula os percentuais individuais
                            percentuais, _ = calcular_disc(respostas_traduzidas)
                            resultados_acumulados.append(percentuais)

                    # 4. CONSOLIDAÇÃO E MATEMÁTICA
                    if resultados_acumulados:
                        # Criamos o DataFrame com todos os colaboradores
                        df_total = pd.DataFrame(resultados_acumulados).apply(pd.to_numeric).dropna()
                        
                        # Médias para força do grupo e Somas para peso total
                        medias = df_total.mean()
                        somas = df_total.sum()
                        
                        st.divider()
                        st.success(f"✅ SUCESSO! {len(resultados_acumulados)} perfis processados e somados.")

                        # 5. EXIBIÇÃO DOS RESULTADOS VISUAIS
                        col_info, col_chart = st.columns([1, 2])

                        with col_info:
                            st.markdown("### 📈 Peso Total da Equipe")
                            st.write(somas.rename("Pontuação Acumulada"))
                            
                            dominante_grupo = medias.idxmax()
                            st.info(f"**Cultura Dominante: {dominante_grupo}**")
                            st.caption(f"Análise baseada em {len(resultados_acumulados)} colaboradores.")

                        with col_chart:
                            fig = px.bar(
                                medias.reset_index(), 
                                x="index", y=0, 
                                color="index",
                                text_auto='.1f',
                                labels={'index': 'Fator DISC', '0': 'Amplitude Média %'},
                                color_discrete_map={"D":"#FF4136", "I":"#FF851B", "S":"#2ECC40", "C":"#0074D9"}
                            )
                            fig.update_layout(height=400, showlegend=False, template="plotly_white")
                            st.plotly_chart(fig, use_container_width=True)

                        # ============================================================
                        # 🧠 DIAGNÓSTICO AUTOMÁTICO (TEXTO AJUSTADO PELO LUCIANO)
                        # ============================================================
                        st.markdown("---")
                        st.subheader(f"🔍 Diagnóstico da Cultura do Grupo: {dominante_grupo}")

                        explicações = {
                            "D": {
                                "titulo": "Cultura de Resultados e Velocidade",
                                "texto": "O time é focado em metas, direto e competitivo. Possui entrega rápida e facilidade para lidar com crises, mas pode faltar paciência com processos longos e detalhes burocráticos."
                            },
                            "I": {
                                "titulo": "Cultura de Conexão e Entusiasmo",
                                "texto": "O time é comunicativo, otimista e focado em pessoas. Excelente clima organizacional e criatividade, mas pode ter dificuldade com a organização de dados e rotinas repetitivas."
                            },
                            "S": {
                                "titulo": "Cultura de Colaboração e Persistência",
                                "texto": "O time é leal, constante e trabalha em harmonia. Garante processos seguros e baixa rotatividade, mas pode mostrar resistência a mudanças bruscas."
                            },
                            "C": {
                                "titulo": "Cultura de Precisão e Qualidade (Alerta de Inconsistência)",
                                "texto": (
                                    "O time tem DNA analítico e detalhista, pautado por regras e lógica. "
                                    "**NOTA CRÍTICA:** Se este grupo apresenta altos índices de erro, o diagnóstico indica 'Pane Operacional'. "
                                    "Perfis de Conformidade erram quando submetidos a processos ambíguos, falta de checklists claros ou pressão extrema por velocidade em detrimento da qualidade. "
                                    "O erro aqui não é falta de atenção, é sobrecarga cognitiva. Para corrigir: implemente processos padronizados (POPs) "
                                    "e elimine a subjetividade das tarefas. O time precisa de trilhos, não de velocidade cega."
                                )
                            }
                        }

                        exp = explicações.get(dominante_grupo)
                        if exp:
                            st.warning(f"### {exp['titulo']}")
                            st.write(exp['texto'])

                    else:
                        st.warning("⚠️ Nenhum dado DISC válido encontrado nos arquivos.")

            except Exception as e:
                st.error(f"❌ Erro crítico no processamento: {e}")

# --- Rodapé Final ---
st.markdown("---")




import streamlit as st
import pandas as pd
import json
import os
from openai import OpenAI

import streamlit as st


if st.session_state.get("pagina") == "parecer":

    @st.cache_data(show_spinner=True)
    def buscar_benchmark_ia_estrategico(cargo, funcao, objetivo, qualificacoes):
        try:
            client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"].strip())
            
            contexto_resumido = f"{objetivo[:700]}... Qualificações: {qualificacoes[:300]}"
            
            # PROMPT DE OURO: ENGENHARIA DE PROCESSOS E CRONOANÁLISE FORENSE
            prompt = f"""
            Aja como um Engenheiro de Processos Sênior e Especialista em Cronoanálise Forense.
            Sua missão é realizar uma DECOMPOSIÇÃO ESTRATÉGICA DE CARGA HORÁRIA para o cargo '{cargo}' com foco em '{funcao}'.

            CONTEXTO DO ALVO:
            {contexto_resumido}

            REGRAS RÍGIDAS DE ARQUITETURA DE DADOS:
            1. JORNADA TOTAL: Você deve preencher exatamente 480 minutos de impacto diário.
            2. MULTIFREQUÊNCIA: Distribua as atividades em DIÁRIA, SEMANAL e MENSAL.
            3. CONVERSÃO DE IMPACTO (CÁLCULO): 
            - Atividade Semanal: (Tempo Total / 5 dias).
            - Atividade Mensal: (Tempo Total / 22 dias).
            4. CRITÉRIO DE VALOR: Priorize atividades de ALTO VALOR AGREGADO (Gestão, Auditoria, Estratégia, Compliance). Elimine tarefas braçais.
            5. OBJETIVOS TÉCNICOS: Cada 'meta' deve ser um KPI ou um marco de controle auditável.

            FORMATO DE SAÍDA (ESTRITAMENTE JSON):
            {{
                "NOME_DA_ATIVIDADE": {{
                    "tempo": minutos_inteiros,
                    "freq": "DIÁRIA/SEMANAL/MENSAL",
                    "meta": "Descrição técnica do objetivo e métrica de sucesso",
                    "complexidade": "ALTA/MÉDIA/BAIXA"
                }}
            }}

            Pense passo a passo: Identifique as rotinas críticas, calcule o tempo em alta performance e ajuste os pesos para totalizar 480min de impacto diário.
            """

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": "Você é um auditor forense de processos sênior."},
                        {"role": "user", "content": prompt}],
                response_format={ "type": "json_object" }
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            st.error(f"Erro na Inteligência: {e}")
            return None

    # ==============================================================================
    # 🛡️ MOTOR DE AUDITORIA NETEXAME: TRÍPTICO PERICIAL
    # ==============================================================================

    @st.cache_data(show_spinner="Analisando dados e economizando créditos...")
    def processar_parecer_com_cache(dados_json_str):
        """
        Se o JSON for o mesmo, o Streamlit retorna o laudo da memória.
        Custo: 0 créditos nas repetições.
        """
        # CHAME AQUI A SUA FUNÇÃO QUE CONECTA COM A OPENAI
        # Certifique-se de que 'realizar_pericia_direta' existe no seu script
        return realizar_pericia_direta(dados_json_str)

    def mostrar_pagina_parecer():
        st.title("🛡️ NetExame: Auditoria Forense Estratégica")
        st.markdown("---")
        
        caminho_dados = "dados"
        if not os.path.exists(caminho_dados):
            st.error("Pasta de 'dados' não encontrada.")
            return

        arquivos = [f for f in os.listdir(caminho_dados) if f.endswith('.json')]
        
        if arquivos:
            colaborador_file = st.selectbox("🎯 Selecione o Alvo da Auditoria:", arquivos)
            
            if st.button("🚀 Gerar Laudo de Eficiência Avançado"):
                with open(os.path.join(caminho_dados, colaborador_file), 'r', encoding='utf-8') as f:
                    colab = json.load(f)
                
                # Execução do Cérebro IA
                pop_ia = buscar_benchmark_ia_estrategico(
                    colab['campos'].get('cargo', 'N/A').upper(),
                    colab['campos'].get('funcao', 'GESTÃO').upper(),
                    colab['campos'].get('objetivo', ''),
                    colab['campos'].get('cursos', '')
                )

                if pop_ia:
                    # --------------------------------------------------------------
                    # BLOCO 1: O DEVER (POP PADRÃO IA)
                    # --------------------------------------------------------------
                    st.header("📚 [A] POP Padrão IA (Carga Diária 480m)")
                    dados_ia = []
                    total_ia_diario = 0
                    for ativ, info in pop_ia.items():
                        t, f = info['tempo'], info['freq'].upper()
                        imp = t if "DIÁRIA" in f else (t/5 if "SEMANAL" in f else t/22)
                        total_ia_diario += imp
                        dados_ia.append({
                            "Atividade": ativ,
                            "Freq": f,
                            "Tempo Base": f"{t}m",
                            "Impacto Diário Convertido": f"{imp:.1f}m",
                            "Eficiência vs 480m": f"{(imp/480)*100:.1f}%",
                            "Meta Auditável": info['meta']
   
                            
                        })
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Carga Alvo", "480 min")
                    c2.metric("Ocupação POP IA", f"{total_ia_diario:.1f} min")
                    c3.metric("Eficiência Teórica", f"{(total_ia_diario/480)*100:.1f}%")
                    st.table(pd.DataFrame(dados_ia))

                    # --- GERADOR DE HTML PARA DOWNLOAD ---
                    html_content = f"""
                    <html>
                    <head><meta charset="UTF-8"><title>POP Padrão IA</title></head>
                    <body style="font-family: sans-serif; padding: 20px;">
                        <h2>📚 [A] POP Padrão IA (Carga Diária 480m)</h2>
                        <table border="1" style="border-collapse: collapse; width: 100%;">
                            <tr style="background-color: #f2f2f2;">
                                <th>Atividade</th><th>Freq</th><th>Tempo Base</th><th>Impacto Diário</th><th>Eficiência</th><th>Meta Auditável</th>
                            </tr>
                    """
                    for d in dados_ia:
                        html_content += f"""
                            <tr>
                                <td>{d['Atividade']}</td><td>{d['Freq']}</td><td>{d['Tempo Base']}</td>
                                <td>{d['Impacto Diário Convertido']}</td><td>{d['Eficiência vs 480m']}</td><td>{d['Meta Auditável']}</td>
                            </tr>"""

                    html_content += f"""
                        </table>
                        <p><strong>Carga Alvo:</strong> 480 min | <strong>Ocupação:</strong> {total_ia_diario:.1f} min | <strong>Eficiência:</strong> {(total_ia_diario/480)*100:.1f}%</p>
                    </body>
                    </html>
                    """

                    st.download_button(
                        label="📥 Baixar POP em HTML",
                        data=html_content,
                        file_name="pop_ia_netexame.html",
                        mime="text/html"
                    )

                    # --------------------------------------------------------------
                    # BLOCO 2: O SER (ATIVIDADES REALMENTE LISTADAS)
                    # --------------------------------------------------------------
                    st.header("📝 [B] Realidade Relatada (Análise de Esforço)")
                    atividades_relatadas = colab['tabelas'].get('alta', []) + \
                                        colab['tabelas'].get('normal', []) + \
                                        colab['tabelas'].get('baixa', [])
                    
                    dados_reais = []
                    total_real_diario = 0
                    for item in atividades_relatadas:
                        h = int(str(item.get('Horas', '0')).split()[0])
                        m = int(str(item.get('Minutos', '0')).split()[0])
                        t_bruto = (h * 60) + m
                        f_real = item.get('Frequência', '').upper()
                        
                        # Normalização de frequência para cálculo
                        if "D" in f_real: imp = t_bruto
                        elif "S" in f_real: imp = t_bruto / 5
                        elif "M" in f_real: imp = t_bruto / 22
                        else: imp = 0
                        
                        total_real_diario += imp
                        dados_reais.append({
                            "Atividade Listada": item.get('Atividade'),
                            "Freq": f_real,
                            "Tempo Bruto": f"{t_bruto}m",
                            "Tempo Diário Convertido": f"{imp:.1f}m",
                            "Peso na Jornada": f"{(imp/480)*100:.1f}%"
                        })
                    
                    r1, r2, r3 = st.columns(3)
                    r1.metric("Total Real Relatado", f"{total_real_diario:.1f} min")
                    r2.metric("Eficiência Diária Alvo", f"{(total_real_diario/480)*100:.1f}%")
                    r3.metric("Gap/Ociosidade", f"{480 - total_real_diario:.1f} min")
                    st.table(pd.DataFrame(dados_reais))

                    # --------------------------------------------------------------
                    # BLOCO 3: O CRUZAMENTO (CONFRONTO E NEXO CAUSAL)
                    # --------------------------------------------------------------
                    st.header("⚖️ [C] Cruzamento e Veredito Pericial")
                    confronto = []
                    mapeadas = set()

                    for ativ_ia, info_ia in pop_ia.items():
                        f_ia = info_ia['freq'].upper()
                        imp_ia = info_ia['tempo'] if "DIÁRIA" in f_ia else (info_ia['tempo']/5 if "SEMANAL" in f_ia else info_ia['tempo']/22)
                        
                        imp_real_vinculado = 0
                        status = "❌ AUSENTE"
                        chave = ativ_ia.split()[0].upper()
                        
                        for item in atividades_relatadas:
                            desc = str(item.get('Atividade', '')).upper()
                            if chave in desc:
                                h, m = int(str(item.get('Horas', '0')).split()[0]), int(str(item.get('Minutos', '0')).split()[0])
                                t_b, f_r = (h * 60) + m, item.get('Frequência', '').upper()
                                imp_v = t_b if "D" in f_r else (t_b/5 if "S" in f_r else t_b/22)
                                imp_real_vinculado += imp_v
                                status = "✅ IDENTIFICADO"
                                mapeadas.add(desc)

                        confronto.append({
                            "Atividade": ativ_ia,
                            "Origem": "POP PADRÃO",
                            "Impacto POP": f"{imp_ia:.1f}m",
                            "Impacto Real": f"{imp_real_vinculado:.1f}m",
                            "Divergência": f"{imp_real_vinculado - imp_ia:+.1f}m",
                            "Status": status
                        })

                    # Adiciona o que sobrou (Desvios)
                    for item in atividades_relatadas:
                        desc = str(item.get('Atividade', '')).upper()
                        if desc not in mapeadas:
                            h, m = int(str(item.get('Horas', '0')).split()[0]), int(str(item.get('Minutos', '0')).split()[0])
                            t_b, f_r = (h * 60) + m, item.get('Frequência', '').upper()
                            imp_extra = t_b if "D" in f_r else (t_b/5 if "S" in f_r else t_b/22)
                            confronto.append({
                                "Atividade": desc,
                                "Origem": "DESVIO",
                                "Impacto POP": "0.0m",
                                "Impacto Real": f"{imp_extra:.1f}m",
                                "Divergência": f"+{imp_extra:.1f}m",
                                "Status": "⚠️ FORA DO PADRÃO"
                            })

                    f1, f2, f3 = st.columns(3)
                    f1.metric("Aderência ao Cargo", f"{(total_ia_diario/total_real_diario*100 if total_real_diario > 0 else 0):.1f}%")
                    f2.metric("Perda/Ganho Eficiência", f"{total_ia_diario - total_real_diario:+.1f} min")
                    f3.metric("Risco Operacional", "ALTO" if total_real_diario > 480 else "BAIXO")
                    st.table(pd.DataFrame(confronto))

    # --- INICIALIZAÇÃO ---
    if 'pagina' not in st.session_state:
        st.session_state.pagina = "parecer"

    mostrar_pagina_parecer()



    import streamlit as st
    import pandas as pd
    import base64
    import json
    from openai import OpenAI

    # ==============================================================================
    # 1. CONFIGURAÇÕES INICIAIS E SEGURANÇA
    # ==============================================================================
    # Certifique-se de ter a chave nas configurações do Streamlit Cloud (Secrets)
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    # Inicializa as variáveis no estado da sessão para evitar NameError
    if 'resultado_parecer_gpt' not in st.session_state:
        st.session_state['resultado_parecer_gpt'] = "Aguardando processamento da análise pericial..."

    if 'analise_concluida' not in st.session_state:
        st.session_state['analise_concluida'] = False

    # ==============================================================================
    # 2. MOTOR DE INTELIGÊNCIA FORENSE (OPENAI)
    # ==============================================================================
    def realizar_pericia_ia(nome_colaborador, cargo, atividades_relatadas):
        """
        Chama o GPT-4o-mini para analisar o nexo causal e gerar o POP Universal.
        """
        prompt = f"""
        Aja como um Auditor Forense de Processos Sênior. 
        Analise o cargo '{cargo}' para o colaborador '{nome_colaborador}'.
        O colaborador relatou as seguintes atividades: {atividades_relatadas}

        SUA MISSÃO:
        1. Reafirme o POP PADRÃO UNIVERSAL (480 min totais).
        2. Escreva um PARECER TÉCNICO focado em:
        - Inconsistência de carga horária (se houver).
        - Desvios de função (operacional vs estratégico).
        - Riscos de omissão de tarefas críticas (como conferência de folha).

        RESPONDA EXCLUSIVAMENTE NO FORMATO JSON ABAIXO:
        {{
            "parecer_pericial": "Seu texto de parecer aqui detalhado.",
            "pop_universal": [
                {{"Atividade": "A", "Freq": "D", "Tempo": "Xm", "Impacto": "Ym", "Peso": "Z%"}}
            ]
        }}
        """
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Você é um auditor especializado em eficiência operacional e cronoanálise."},
                    {"role": "user", "content": prompt}
                ],
                response_format={ "type": "json_object" }
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            st.error(f"Erro na API da OpenAI: {e}")
            return None


    # ==============================================================================
    # 3. GERADOR DE ARTEFATOS (HTMLS SEPARADOS)
    # ==============================================================================

    def gerar_html_laudo_puro(nome_colab, parecer_ia):
        """Gera o HTML contendo APENAS o cabeçalho e o parecer técnico (sem a tabela)."""
        html = f"""
        <html>
        <head><meta charset="utf-8"><style>
            body {{ font-family: 'Segoe UI', sans-serif; background-color: #f4f7f6; padding: 40px; color: #333; }}
            .container {{ background: white; padding: 40px; border-radius: 15px; max-width: 900px; margin: auto; border-top: 10px solid #d90429; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }}
            .header {{ background: #0d1b2a; color: white; padding: 20px; border-radius: 8px; margin-bottom: 30px; }}
            .parecer {{ background: #fff5f5; border-left: 5px solid #d90429; padding: 25px; border-radius: 0 8px 8px 0; font-style: italic; line-height: 1.6; }}
            footer {{ text-align: center; font-size: 11px; color: #999; margin-top: 40px; border-top: 1px solid #eee; padding-top: 20px; }}
        </style></head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>🛡️ NetExame: Parecer Pericial</h1>
                    <p>Análise Forense de Processos | Colaborador: {nome_colab}</p>
                </div>
                <h3>🔍 Diagnóstico Técnico Pericial</h3>
                <div class="parecer">{parecer_ia}</div>
                <footer>Gerado por NetExame Auditoria & IA Forense 2026</footer>
            </div>
        </body>
        </html>
        """
        return html



    # ==============================================================================
    # 4. INTERFACE E EXECUÇÃO
    # ==============================================================================
    st.title("⚖️ NetExame: Perícia Forense de Processos")

    # Inputs (No seu caso, isso viria do JSON que você já carrega)
    nome_alvo = st.text_input("Nome do Colaborador", "ADSON")
    cargo_alvo = st.text_input("Cargo", "GESTOR DE DP")
    relato_exemplo = "Atendimento a clientes, auditoria de folha, suporte técnico, organizar arquivos, etc."

    if st.button("🚀 INICIAR PERÍCIA TÉCNICA"):
        with st.spinner("IA analisando nexo causal e eficiência..."):
            resultado = realizar_pericia_ia(nome_alvo, cargo_alvo, relato_exemplo)
            
            if resultado:
                # SALVA NO ESTADO DA SESSÃO
                st.session_state['resultado_parecer_gpt'] = resultado['parecer_pericial']
                st.session_state['pop_universal_ia'] = resultado['pop_universal']
                st.session_state['analise_concluida'] = True
                
                st.success("Análise Concluída!")
                st.markdown(f"**Parecer:** {resultado['parecer_pericial']}")

    # --- SEÇÃO DE DOWNLOAD DO LAUDO ---
    if st.session_state['analise_concluida']:
        st.markdown("---")
        st.subheader("🏁 Finalização e Entrega")
        
        col1, col2 = st.columns(2)

        with col1:
            if st.button("📥 BAIXAR LAUDO PERICIAL"):
                # HTML SEM O POP (Apenas Parecer e Cabeçalho)
                html_laudo = gerar_html_laudo_puro(
                    nome_alvo, 
                    st.session_state['resultado_parecer_gpt']
                )
                b64 = base64.b64encode(html_laudo.encode('utf-8')).decode()
                href = f'<a href="data:text/html;base64,{b64}" download="LAUDO_{nome_alvo}.html" style="text-decoration:none;"><button style="background-color:#d90429;color:white;padding:15px;border:none;border-radius:10px;cursor:pointer;font-weight:bold;width:100%;">📄 BAIXAR PARECER</button></a>'
                st.markdown(href, unsafe_allow_html=True)




    import streamlit as st
    import json
    import os
    import glob
    from openai import OpenAI

    # 1. SETUP E CONEXÃO
    st.set_page_config(page_title="IA Auditor Pro 360", layout="wide")

    # Inicializa o cliente OpenAI (Certifique-se de configurar a chave no Streamlit Secrets ou env)
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    def solicitar_pericia_360(dados):
        """
        O Cérebro do Sistema: Envia o JSON e recebe o laudo estratégico.
        """
        prompt = f"""
        ESTRUTURA DE DADOS BRUTOS:
        {json.dumps(dados, indent=2, ensure_ascii=False)}

        SUA MISSÃO: Realizar uma perícia técnica 360° (Consultor Sênior).
        
        CRUZAMENTOS OBRIGATÓRIOS:
        1. FORMAÇÃO vs. OPERAÇÃO: O colaborador tem cursos de Pós e Especializações (ver em 'cursos') mas gasta tempo em tarefas de 'baixa' complexidade?
        2. ANÁLISE DE CARGA (MATEMÁTICA): Some os tempos e frequências. Se ultrapassar 480 min/dia e ele disse 'nenhuma dificuldade', identifique o risco de Burnout Mascarado.
        3. HIBRIDISMO DISC: Analise se o perfil é híbrido (ex: A com D). Explique como isso gera conflito entre 'entrega rápida' e 'perfeccionismo exagerado'.
        4. ALINHAMENTO ESTRATÉGICO: O 'Objetivo' dele condiz com as tarefas de 'Alta' ou ele é um gestor sequestrado pelo operacional?
        5. VEREDITO: O que ele deve DELEGAR, PARAR e FOCO IMEDIATO.

        SAÍDA: Use Markdown, seja direto, crítico e focado em ROI humano.
        """

        response = client.chat.completions.create(
            model="gpt-4o", # Ou gpt-4-turbo
            messages=[
                {"role": "system", "content": "Você é um Auditor Forense de Performance Humana."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2
        )
        return response.choices[0].message.content

    # 2. INTERFACE UNIVERSAL
    def main():
        st.title("🛡️ Motor de Perícia e Auditoria IA 360°")
        st.info("Este sistema analisa o DNA do colaborador: Formação, Tempo, Objetivo e Comportamento.")

        # Busca arquivos JSON
        arquivos = glob.glob("**/dados/*.json", recursive=True) + glob.glob("*.json")
        
        if not arquivos:
            st.warning("📂 Nenhum JSON encontrado. Salve os dados na pasta /dados.")
            return

        lista_colab = {os.path.basename(f): f for f in arquivos}
        escolha = st.selectbox("🎯 Selecione o Colaborador para Auditoria:", list(lista_colab.keys()))

        if escolha:
            with open(lista_colab[escolha], "r", encoding="utf-8") as f:
                dados_lidos = json.load(f)
                
                # Layout de colunas para dados rápidos
                c1, c2, c3 = st.columns(3)
                c1.metric("Colaborador", dados_lidos.get('colaborador'))
                c2.metric("Cargo", dados_lidos.get('campos', {}).get('cargo'))
                c3.metric("Unidade", dados_lidos.get('campos', {}).get('unidade'))

                if st.button("🚀 GERAR LAUDO FORENSE 360°"):
                    with st.spinner("IA processando cruzamento de dados..."):
                        laudo = solicitar_pericia_360(dados_lidos)
                        st.divider()
                        st.markdown(laudo)
                        
                        # 1. Botão Original (Markdown)
                        st.download_button(
                            "Exportar Laudo (.md)", 
                            laudo, 
                            file_name=f"pericia_{dados_lidos.get('colaborador')}.md"
                        )

                        # 2. SEU NOVO BOTÃO HTML (Agora protegido dentro do IF)
                        html_template = f"""
                        <!DOCTYPE html>
                        <html lang="pt-BR">
                        <head>
                            <meta charset="UTF-8">
                            <style>
                                body {{ font-family: Arial, sans-serif; margin: 30px; line-height: 1.6; color: #333; }}
                                h1, h2, h3 {{ color: #2c3e50; border-bottom: 2px solid #ef233c; padding-bottom: 5px; }}
                                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                                th, td {{ border: 1px solid #bdc3c7; padding: 12px; text-align: left; }}
                                th {{ background-color: #ecf0f1; font-weight: bold; }}
                                .footer {{ margin-top: 30px; font-size: 0.8em; color: #7f8c8d; text-align: center; }}
                            </style>
                        </head>
                        <body>
                            {laudo.replace('\n', '<br>')}
                            <div class="footer">Gerado automaticamente por Motor de Perícia IA 360°</div>
                        </body>
                        </html>
                        """
                        
                        st.download_button(
                            label="🌐 Exportar em HTML (Visual Profissional)",
                            data=html_template,
                            file_name=f"pericia_{dados_lidos.get('colaborador', 'doc')}.html",
                            mime="text/html",
                            use_container_width=True
                        )


    if __name__ == "__main__":
        main()

                

    # ==============================================================================
    # 🧠 MOTOR DE INTELIGÊNCIA: PERÍCIA FORENSE 360° (NEXO CAUSAL TOTAL)
    # ==============================================================================
    @st.cache_data(show_spinner="Recuperando perícia do cache...")
    def realizar_super_pericia_ia(dados):
        """
        Realiza a auditoria cruzada: JSON + Benchmark + Nexo Causal + DISC.
        """
        # Preparação dos dados para o Prompt
        contexto_dados = json.dumps(dados, indent=2, ensure_ascii=False)
        
        prompt = f"""
        Aja como um Perito Auditor Forense e Engenheiro de Processos Sênior. 
        Sua missão é realizar uma DECOMPOSIÇÃO ESTRATÉGICA E ANÁLISE DE NEXO CAUSAL.

        DADOS BRUTOS DO ALVO:
        {contexto_dados}

        SUA AUDITORIA DEVE CRUZAR OBRIGATORIAMENTE:
        1. NEXO CAUSAL: O Objetivo do cargo condiz com as tarefas de 'Alta' complexidade relatadas? 
        2. CAPACIDADE vs. ENTREGA: O colaborador possui cursos de Pós/Especialização (ver em 'cursos') que estão sendo desperdiçados em tarefas operacionais?
        3. HIBRIDISMO COMPORTAMENTAL: Analise o DISC. Como o perfil (Ex: Dominante/Conformidade) impacta a velocidade e a precisão nestas tarefas?
        4. CRONOANÁLISE CRÍTICA: Se a soma das tarefas (D/S/M) converter para mais de 480min/dia, aponte risco de Burnout Mascarado e Erro Forense.
        5. VEREDITO ESTRATÉGICO: O que ele deve PARAR de fazer, o que deve DELEGAR e onde deve ser o FOCO IMEDIATO.

        FORMATO DE SAÍDA (ESTRITAMENTE JSON):
        {{
            "parecer_executivo": "Texto profundo, crítico e técnico (estilo consultoria sênior) detalhando os desvios e o nexo causal.",
            "pop_benchmark": [
                {{"Atividade": "Nome da Tarefa", "Freq": "D/S/M", "Tempo": "X min", "Meta": "KPI de Sucesso"}}
            ],
            "veredito_final": "Resumo executivo do plano de ação."
        }}
        """
        
        try:
            client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Você é um auditor sênior especializado em eficiência operacional e capital humano."},
                    {"role": "user", "content": prompt}
                ],
                response_format={ "type": "json_object" },
                temperature=0.3
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            st.error(f"Erro na Perícia IA: {e}")
            return None



    import streamlit as st
    import pandas as pd
    import json
    import os
    import glob
    import base64
    from openai import OpenAI

    # ==============================================================================
    # 1. SETUP E MOTOR DE INTELIGÊNCIA (NEXO CAUSAL 360°)
    # ==============================================================================
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    @st.cache_data(show_spinner="Consultando base de dados e IA...", ttl=86400)
    def realizar_super_pericia_ia(dados):
        """
        O Cérebro do Sistema: Cruza JSON, Benchmark e Gera Parecer Forense.
        """
        contexto_puro = json.dumps(dados, indent=2, ensure_ascii=False)
        
        prompt = f"""
        Aja como um Perito Auditor Forense e Engenheiro de Processos Sênior. 
        Sua missão é realizar uma DECOMPOSIÇÃO ESTRATÉGICA E ANÁLISE DE NEXO CAUSAL.

        DADOS BRUTOS DO ALVO:
        {contexto_puro}

        SUA AUDITORIA DEVE CRUZAR OBRIGATORIAMENTE:
        1. NEXO CAUSAL: O Objetivo do cargo condiz com as tarefas relatadas? 
        2. CAPACIDADE vs. ENTREGA: Há desperdício de formação (cursos) em tarefas operacionais?
        3. HIBRIDISMO COMPORTAMENTAL: Como o perfil DISC impacta a execução?
        4. CRONOANÁLISE: Se > 480min/dia, aponte risco de Burnout Mascarado.
        5. VEREDITO: O que PARAR, DELEGAR e FOCO IMEDIATO.

        FORMATO DE SAÍDA (ESTRITAMENTE JSON):
        {{
            "parecer_executivo": "Texto profundo e técnico sobre desvios e nexo causal.",
            "pop_benchmark": [
                {{"Atividade": "Nome", "Freq": "D/S/M", "Tempo": "X min", "Meta": "KPI"}}
            ],
            "veredito_final": "Resumo executivo do plano de ação."
        }}
        """
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "Você é um auditor sênior."},
                        {"role": "user", "content": prompt}],
                response_format={ "type": "json_object" }
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            st.error(f"Erro na IA: {e}")
            return None



    import streamlit as st
    import json
    import glob
    import base64
    from openai import OpenAI

    # ==============================================================================
    # 1. MOTOR DE PERÍCIA (RESILIENTE)
    # ==============================================================================
    @st.cache_data(show_spinner=False)
    def realizar_pericia_direta(dados_json_str):
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Você é um Perito Forense em RH. Gere um laudo técnico focado em PERFIL e NEXO CAUSAL."},
                    {"role": "user", "content": f"Una o Perfil DISC, as competências técnicas e os POPs necessários para este colaborador. Retorne APENAS um JSON com 'analise_perfil_nexo' (texto) e 'pop_estrategico' (texto): {dados_json_str}"}
                ],
                response_format={ "type": "json_object" }
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {"error": str(e)}

    # ==============================================================================
    # 2. GERADOR DE HTML (ESTRUTURA LIMPA)
    # ==============================================================================
    def construir_html_pericial(dados, analise_ia):
        nome = dados.get('colaborador', 'N/A')
        cargo = dados.get('campos', {}).get('cargo', 'N/A').upper()
        
        # Pegando os textos da IA com proteção contra None
        corpo_laudo = analise_ia.get('analise_perfil_nexo', 'Dados de análise não disponíveis.')
        pop_texto = analise_ia.get('pop_estrategico', 'POP não definido.')

        return f"""
        <html>
        <head><meta charset="UTF-8">
            <style>
                body {{ font-family: 'Segoe UI', sans-serif; padding: 40px; color: #2b2d42; line-height: 1.6; }}
                .container {{ background: white; padding: 35px; border-radius: 12px; border-left: 10px solid #d90429; box-shadow: 0 4px 20px rgba(0,0,0,0.08); max-width: 800px; margin: auto; }}
                h1 {{ border-bottom: 2px solid #eee; padding-bottom: 10px; color: #d90429; }}
                .section-title {{ font-weight: bold; text-transform: uppercase; margin-top: 25px; color: #2b2d42; background: #f4f7f6; padding: 5px 10px; }}
                .content {{ margin-top: 10px; text-align: justify; }}
                .footer {{ margin-top: 40px; font-size: 12px; color: #777; border-top: 1px solid #eee; padding-top: 10px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>LAUDO TÉCNICO: {nome}</h1>
                <p><strong>CARGO:</strong> {cargo} | <strong>AUDITORIA DE PROCESSO</strong></p>
                
                <div class="section-title">I. ANÁLISE COMPORTAMENTAL E NEXO CAUSAL</div>
                <div class="content">{corpo_laudo}</div>
                
                <div class="section-title">II. DIRETRIZES DE POP (PROCEDIMENTO OPERACIONAL PADRÃO)</div>
                <div class="content">{pop_texto}</div>
                
                <div class="footer">Documento gerado para fins de auditoria interna e otimização de processos.</div>
            </div>
        </body>
        </html>
        """

    # ==============================================================================
    # 3. INTERFACE
    # ==============================================================================
    def main():
        st.title("🛡️ NetExame: Perícia & POP")
        
        arquivos = glob.glob("**/dados/*.json", recursive=True) + glob.glob("*.json")
        escolha = st.selectbox("🎯 Selecionar Perfil:", arquivos)

        if escolha:
            with open(escolha, 'r', encoding='utf-8') as f:
                dados_alvo = json.load(f)

            if st.button("🚀 GERAR PERÍCIA TÉCNICA", use_container_width=True):
                dados_str = json.dumps(dados_alvo, ensure_ascii=False)
                
                with st.spinner("Correlacionando Perfil e Processos..."):
                    resultado = realizar_pericia_direta(dados_str)
                    
                    if "error" not in resultado:
                        html_final = construir_html_pericial(dados_alvo, resultado)
                        b64 = base64.b64encode(html_final.encode('utf-8')).decode()
                        nome_colab = dados_alvo.get('colaborador')
                        
                        st.markdown(f'''
                            <div style="text-align:center; margin-top:20px;">
                                <a href="data:text/html;base64,{b64}" download="PERICIA_{nome_colab}.html">
                                    <button style="background:#d90429; color:white; padding:20px; border-radius:12px; cursor:pointer; width:100%; font-weight:bold; border:none; font-size:18px;">
                                        📥 BAIXAR LAUDO TÉCNICO (HTML)
                                    </button>
                                </a>
                            </div>
                        ''', unsafe_allow_html=True)
                        st.success("Perícia concluída com foco em Nexo e POP!")

    if __name__ == "__main__":
        main()


import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import json
import base64
import os
from openai import OpenAI
from datetime import datetime


import streamlit as st
import hashlib
import json

def buscar_benchmark_ia_estrategico(cargo, funcao, objetivo, cursos):
    """Gera o POP Estratégico baseado no cargo."""
    prompt = f"Gere um POP para o cargo {cargo}, função {funcao}. Objetivo: {objetivo}. Cursos: {cursos}. Saída em JSON com tempo, freq e meta."
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "Especialista em Processos e POPs."},
                      {"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except: return None

def definir_kpis_inteligentes_pop(pop_texto):
    """Gera os 5 KPIs baseados no POP gerado."""
    prompt = f"Com base neste POP: {pop_texto}, defina 5 KPIs críticos em JSON: nome, benchmark, legenda, evidencia."
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "CPO - Chief Productivity Officer."},
                  {"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def auditar_performance_ia(kpi, relato, arquivos_nomes):
    """Audita as evidências enviadas."""
    prompt = f"Audite: KPI {kpi['nome']}, Relato: {relato}, Provas: {arquivos_nomes}. Retorne JSON: realizado (0-100), parecer, status."
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)



# ==============================================================================
# MOTOR DE CACHE (A ECONOMIA DE DINHEIRO ESTÁ AQUI)
# ==============================================================================

def gerar_hash_dados(dados):
    """Cria uma assinatura única para o conjunto de dados."""
    return hashlib.md5(json.dumps(dados, sort_keys=True).encode()).hexdigest()

@st.cache_data(show_spinner=False, ttl=3600) # Cache por 1 hora
def chamar_ia_com_cache(tipo_chamada, payload):
    """
    Esta função só gasta sua chave se os dados mudarem.
    Se o payload for o mesmo, ele traz a resposta instantaneamente do cache.
    """
    if tipo_chamada == "DEFINIR_KPIS":
        # Aqui vai sua lógica de: client.chat.completions.create(...)
        # return definir_kpis_inteligentes_pop(payload)
        pass
    
    elif tipo_chamada == "AUDITAR_KPI":
        # return auditar_performance_ia(payload['kpi'], payload['relato'], payload['arquivos'])
        pass

# ==============================================================================
# IMPLEMENTAÇÃO NO SEU APP
# ==============================================================================

def aba_produtividade_com_economia():
    st.title("💰 Gestão com Otimização de Créditos")
    
    pop_texto = st.text_area("Insira o POP:")
    
    # Criamos um identificador único para este POP
    hash_pop = gerar_hash_dados(pop_texto)

    if st.button("🚀 Gerar KPIs (Com Cache)"):
        # O Streamlit verifica se já rodou essa função com esse 'hash_pop' antes
        # Se sim, ele não gasta sua chave OpenAI!
        config = chamar_ia_com_cache("DEFINIR_KPIS", pop_texto)
        st.success("Resultado obtido (via Cache ou API)")


# 1. MOTOR DE INTELIGÊNCIA (PROMPTS DE ALTA PERFORMANCE)
# Certifique-se de ter a chave configurada nos Secrets do Streamlit (OPENAI_API_KEY)
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def definir_kpis_inteligentes_pop(pop_texto):
    """
    Prompt Profissional: Analisa o POP e define 5 KPIs com métricas ideais.
    """
    prompt = f"""
    Aja como um Chief Productivity Officer (CPO). 
    Com base no POP (Procedimento Operacional Padrão) abaixo, defina os 5 indicadores (KPIs) 
    mais críticos para garantir a excelência operacional deste cargo.
    
    POP: {pop_texto}
    
    PARA CADA KPI, VOCÊ DEVE DEFINIR:
    1. Nome do Indicador.
    2. Métrica Ideal (Benchmark %).
    3. Descrição do que ele mede (Legenda).
    4. Tipo de evidência documental necessária para auditoria.

    SAÍDA EXCLUSIVA EM JSON:
    {{
      "kpis": [
        {{"nome": "Nome", "benchmark": 95, "legenda": "Desc", "evidencia": "Tipo"}}
      ]
    }}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "Especialista em Gestão de Performance e KPIs."},
                  {"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

def auditar_performance_ia(kpi, relato, arquivos_nomes):
    """
    Prompt Profissional: Audita a evidência e gera o valor real vs meta.
    """
    prompt = f"""
    Aja como um Auditor de Governança Corporativa. 
    Analise a entrega do colaborador para o KPI: {kpi['nome']}.
    Métrica Ideal: {kpi['benchmark']}%
    Relato do Colaborador: {relato}
    Evidências (Arquivos Enviados): {arquivos_nomes}

    MISSÃO:
    1. Atribua o 'Valor Realizado' (0-100) baseado no nexo causal entre o relato e as provas.
    2. Gere um Parecer Técnico justificando a nota de forma executiva.
    3. Determine o Status: 'Auditado e Validado' ou 'Inconsistência Identificada'.

    SAÍDA EXCLUSIVA EM JSON:
    {{
      "realizado": 80,
      "parecer": "Texto analítico curto",
      "status": "Status"
    }}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

# 2. CONSTRUTOR DE ARTEFATOS (HTML)
def gerar_html_executivo(colaborador, df_kpi, parecer_geral):
    style = """
    <style>
        body { font-family: 'Segoe UI', sans-serif; background: #f0f2f6; padding: 40px; }
        .card { background: white; padding: 30px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }
        .header { border-bottom: 5px solid #1e3a8a; margin-bottom: 20px; }
        .kpi-row { display: flex; justify-content: space-between; padding: 10px; border-bottom: 1px solid #eee; }
    </style>
    """
    html = f"<html><head>{style}</head><body><div class='card'>"
    html += f"<div class='header'><h1>Laudo Executivo de Performance</h1><p>{colaborador}</p></div>"
    for _, row in df_kpi.iterrows():
        html += f"<div class='kpi-row'><b>{row['KPI']}</b> <span>Meta: {row['Meta']}% | Real: {row['Real']}%</span></div>"
        html += f"<p style='font-size:12px; color:#666;'><i>Legenda: {row['Legenda']}</i></p>"
    html += f"<div style='background:#1e3a8a; color:white; padding:20px; border-radius:10px; margin-top:20px;'><h3>Parecer da Auditoria</h3>{parecer_geral}</div>"
    html += "</div></body></html>"
    return html

def gerar_html_manual_governanca():
    style = """
    <style>
        body { font-family: 'Segoe UI', sans-serif; line-height: 1.6; color: #333; padding: 40px; }
        h1 { color: #1e3a8a; border-bottom: 2px solid #1e3a8a; }
        .rule-box { border: 1px solid #e2e8f0; padding: 20px; border-radius: 10px; background: #f8fafc; }
    </style>
    """
    return f"<html><head>{style}</head><body><h1>Manual de Governança NetExame</h1><div class='rule-box'><h2>Regras de Performance</h2><p>1. KPIs extraídos do POP. 2. Upload múltiplo obrigatório. 3. Auditoria via GPT-4o. 4. Medição Mensal.</p></div></body></html>"

# 3. INTERFACE STREAMLIT
def aba_produtividade_inteligente():
    st.title("📈 Gestão de Performance: Tríptico Pericial")
    
    t1, t2, t3, t4 = st.tabs(["📥 Entrada & Auditoria", "📊 Dashboard Executivo", "🏆 Ranking Global", "📖 Manual"])

    with t1:
        st.subheader("🛡️ Auditoria Forense Estratégica")
        
        caminho_dados = "dados"
        if not os.path.exists(caminho_dados):
            st.error("Pasta de 'dados' não encontrada.")
            return

        arquivos = [f for f in os.listdir(caminho_dados) if f.endswith('.json')]
        
        if arquivos:
            colaborador_file = st.selectbox("🎯 Selecione o Alvo da Auditoria:", arquivos)
            
            with open(os.path.join(caminho_dados, colaborador_file), 'r', encoding='utf-8') as f:
                colab = json.load(f)
            
            # Extração Automática de Dados
            nome_colab = colab['campos'].get('nome_colaborador', colab['campos'].get('nome', 'Alvo'))
            cargo_colab = colab['campos'].get('cargo', 'N/A').upper()
            funcao_colab = colab['campos'].get('funcao', 'GESTÃO').upper()
            objetivo_colab = colab['campos'].get('objetivo', '')
            cursos_colab = colab['campos'].get('cursos', '')

            st.info(f"**Auditando:** {nome_colab} | **Cargo:** {cargo_colab}")

            if st.button("🚀 Gerar POP e Laudo de Eficiência Avançado"):
                with st.spinner("Analisando dados e economizando créditos..."):
                    pop_ia = buscar_benchmark_ia_estrategico(
                        cargo_colab, funcao_colab, objetivo_colab, cursos_colab
                    )

                    if pop_ia:
                        st.session_state['pop_ia_atual'] = pop_ia
                        config_kpis = definir_kpis_inteligentes_pop(str(pop_ia))
                        st.session_state['kpis_pop'] = config_kpis['kpis']
                        st.success("POP e KPIs gerados com sucesso!")

            if 'pop_ia_atual' in st.session_state:
                st.header("📚 [A] POP Padrão IA (Carga Diária 480m)")
                pop_ia = st.session_state['pop_ia_atual']
                dados_ia = []
                total_ia_diario = 0
                for ativ, info in pop_ia.items():
                    # Tenta ler 'tempo', se não achar tenta 'tempo_estimado', se não achar usa 0
                    t = info.get('tempo', info.get('tempo_estimado', 0))
                    # Garante que frequencia seja string antes do .upper()
                    f = str(info.get('freq', info.get('frequencia', 'DIÁRIA'))).upper()
                    
                    # Cálculo de impacto (Base 480min/dia)
                    impacto = t if "DIÁR" in f else (t/5 if "SEMAN" in f else t/22)
                    
                    dados_ia.append({
                        "Atividade": ativ, 
                        "Freq": f, 
                        "Tempo Base": f"{t}m",
                        "Impacto Diário": f"{impacto:.1f}m", 
                        "Meta": info.get('meta', '100%')
                    })
                
                c1, c2 = st.columns(2)
                c1.metric("Ocupação POP IA", f"{total_ia_diario:.1f} min")
                c2.metric("Eficiência Teórica", f"{(total_ia_diario/480)*100:.1f}%")
                st.table(pd.DataFrame(dados_ia))

            if 'kpis_pop' in st.session_state:
                st.divider()
                st.subheader("🔍 Auditoria de Evidências")
                for i, kpi in enumerate(st.session_state['kpis_pop']):
                    with st.expander(f"KPI: {kpi['nome']} (Meta: {kpi['benchmark']}%)"):
                        files = st.file_uploader(f"Provas para {kpi['nome']}", accept_multiple_files=True, key=f"up_{i}")
                        relato = st.text_area("Relato da Entrega", key=f"rel_{i}")
                        
                        if st.button("Auditar Indicador", key=f"bt_{i}"):
                            with st.spinner("Perícia em curso..."):
                                nomes = ", ".join([f.name for f in files]) if files else "Nenhum arquivo"
                                auditoria = auditar_performance_ia(kpi, relato, nomes)
                                st.session_state[f"res_{i}"] = {
                                    "KPI": kpi['nome'], "Meta": kpi['benchmark'], 
                                    "Real": auditoria['realizado'], "Parecer": auditoria['parecer'],
                                    "Legenda": kpi['legenda'], "Status": auditoria['status']
                                }
                                st.success(f"Auditado: {auditoria['realizado']}%")

    with t2:
        if 'res_0' in st.session_state:
            st.header(f"Dashboard: {nome_colab}")
            data = [st.session_state[f"res_{i}"] for i in range(5) if f"res_{i}" in st.session_state]
            df = pd.DataFrame(data)
            fig = go.Figure()
            fig.add_trace(go.Bar(x=df['KPI'], y=df['Meta'], name='Meta', marker_color='#cbd5e1'))
            fig.add_trace(go.Bar(x=df['KPI'], y=df['Real'], name='Real', marker_color='#1e3a8a'))
            st.plotly_chart(fig, use_container_width=True)
            
            html_ex = gerar_html_executivo(nome_colab, df, "Auditoria confirmada via Tríptico Pericial.")
            st.download_button("📥 Baixar Laudo Executivo", html_ex, file_name=f"Laudo_{nome_colab}.html", mime="text/html")

    with t3:
        st.header("🏆 Ranking de Produtividade Global")
        # Ajustado de 'nome' para 'nome_colab' para evitar erro
        mock_rank = pd.DataFrame([{"Nome": nome_colab, "Eficiência": 88}, {"Nome": "Maria", "Eficiência": 94}]).sort_values("Eficiência", ascending=False)
        st.table(mock_rank)

    with t4:
        st.header("📖 Manual de Governança")
        st.components.v1.html(gerar_html_manual_governanca(), height=400)

# Iniciar
if st.session_state.pagina == "produtividade":
    aba_produtividade_inteligente()



