# ============================================================
# IMPORTS
# ============================================================

import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from statistics import mean

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from datetime import datetime
import pytz
import time


# ============================================================

# CONFIGURAÇÃO E INICIALIZAÇÃO ÚNICA

# ============================================================

st.set_page_config(

    page_title="Sistema de Análise de Tarefas",

    page_icon="📊",

    layout="wide",

    initial_sidebar_state="expanded"

)



# Inicialização centralizada

if "logged_in" not in st.session_state: st.session_state.logged_in = False

if "pagina" not in st.session_state: st.session_state.pagina = "home"

if "formularios" not in st.session_state: st.session_state["formularios"] = []



# Leitura da URL (Prioridade total para permitir acesso ao formulário)

query_params = st.query_params

if "page" in query_params:

    st.session_state.pagina = query_params["page"]

st.markdown("""
    <style>
    /* Oculta a coluna de índice do data_editor */
    div[data-testid="stDataEditor"] > div > div > div > div:first-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)


# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]

# --- FUNÇÕES DE EXPORTAÇÃO (COLE NO TOPO DO SEU ARQUIVO) ---
from docx import Document
from fpdf import FPDF
import io

def gerar_word(form):
    doc = Document()
    
    # Nome e Data (Chaves minúsculas)
    doc.add_heading(f"Relatório: {form.get('nome', 'Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('data_envio', 'N/A')}")
    
    # Identificação
    doc.add_heading("Informações de Identificação", level=1)
    campos = ['setor', 'departamento', 'cargo', 'chefe', 'empresa', 'escolaridade', 'cursos_obrigatorios_ou_diferenciais', 'trabalho_e_principal_objetivo']
    for c in campos:
        label = c.replace('_', ' ').capitalize()
        doc.add_paragraph(f"{label}: {form.get(c, 'N/A')}")

    # DISC (Loop minimalista)
    doc.add_heading("Respostas DISC", level=1)
    disc = form.get('disc', {})
    for k, v in disc.items():
        doc.add_paragraph(f"{k}: {v}")

    # Tabelas
    secoes = {"atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
              "dificuldades": ["Dificuldade", "Frequência", "Setor/Parceiro Envolvido", "Tempo Perdido"],
              "sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]}
    
    for chave, cols in secoes.items():
        doc.add_heading(f"📋 {chave.capitalize()}", level=1)
        dados = form.get(chave, [])
        if dados:
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = 'Table Grid'
            for i, col in enumerate(cols): table.rows[0].cells[i].text = col
            for item in dados:
                row = table.add_row().cells
                for i, col in enumerate(cols):
                    row[i].text = str(item.get(col, '')) if isinstance(item, dict) else ""
        else:
            doc.add_paragraph("Sem dados.")

    # Retorno para Streamlit
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def gerar_pdf(form):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    # 1. Título e Data (Chaves corrigidas para minúsculo)
    nome_pdf = form.get('nome', 'Colaborador')
    data_pdf = form.get('data_envio', 'N/A')
    elementos.append(Paragraph(f"Relatório: {nome_pdf}", styles['Title']))
    elementos.append(Paragraph(f"Data: {data_pdf}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    # 2. Informações Gerais
    elementos.append(Paragraph("Informações Gerais", styles['Heading2']))
    campos_gerais = [
        'setor', 'departamento', 'cargo', 'chefe', 'empresa', 
        'escolaridade', 'cursos_obrigatorios_ou_diferenciais', 'trabalho_e_principal_objetivo'
    ]

    for campo in campos_gerais:
        label = campo.replace('_', ' ').capitalize()
        texto = f"<b>{label}:</b> {form.get(campo, 'N/A')}"
        elementos.append(Paragraph(texto, styles['Normal']))
        elementos.append(Spacer(1, 6))

    # 3. Tabelas (Chaves corrigidas para minúsculo)
    secoes = {
        "atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for chave, colunas in secoes.items():
        elementos.append(Paragraph(chave.capitalize(), styles['Heading2']))
        dados = form.get(chave, [])
        
        if isinstance(dados, list) and len(dados) > 0:
            data = [colunas]
            for item in dados:
                if isinstance(item, dict):
                    data.append([str(item.get(c, '')) for c in colunas])
            
            tabela = Table(data, repeatRows=1)
            tabela.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTSIZE', (0,0), (-1,-1), 8)
            ]))
            elementos.append(tabela)
        else:
            elementos.append(Paragraph("Nenhum dado preenchido.", styles['Normal']))
        elementos.append(Spacer(1, 12))

    # 4. DISC (Lógica para buscar dentro do dicionário 'disc')
    elementos.append(Paragraph("Avaliação DISC", styles['Heading2']))
    dados_disc = form.get('disc', {})
    
    if dados_disc:
        for i, pergunta in enumerate(perguntas_disc, 1):
            # Busca a resposta salva como disc_1, disc_2...
            valor_resposta = dados_disc.get(f"disc_{i}", "Não respondido")
            elementos.append(Paragraph(f"<b>{i}. {pergunta}</b>", styles['Normal']))
            elementos.append(Paragraph(f"Resposta: {valor_resposta}", styles['Normal']))
            elementos.append(Spacer(1, 4))
    else:
        elementos.append(Paragraph("Dados DISC não encontrados.", styles['Normal']))

    doc.build(elementos)
    buffer.seek(0)
    return buffer

# ============================================================
# DEFINIÇÃO E CARREGAMENTO DO BANCO DE DADOS (AJUSTADO)
# ============================================================
import streamlit as st
import pandas as pd
import os
import json
import sys

import os
import sys
import json
import streamlit as st

# --- DEFINIÇÃO DE CAMINHO À PROVA DE ERROS ---
if getattr(sys, 'frozen', False):
    base_dir = os.path.dirname(sys.executable)
else:
    base_dir = os.path.dirname(os.path.abspath(__file__))

# Definimos o diretório de dados como absoluto
dados_dir = os.path.join(base_dir, "dados")

# Criamos a pasta 'dados' se ela não existir
os.makedirs(dados_dir, exist_ok=True)

# --- FUNÇÃO DE CARREGAMENTO DINÂMICO ---
def carregar_todos_formularios():
    """
    Lê todos os arquivos .json da pasta 'dados' individualmente.
    """
    lista_formularios = []
    # Usamos a variável global dados_dir definida acima
    if os.path.exists(dados_dir):
        for nome_arquivo in os.listdir(dados_dir):
            if nome_arquivo.endswith(".json"):
                caminho_completo = os.path.join(dados_dir, nome_arquivo)
                try:
                    with open(caminho_completo, "r", encoding="utf-8") as f:
                        dados = json.load(f)
                        if isinstance(dados, dict):
                            lista_formularios.append(dados)
                except Exception:
                    continue
    return lista_formularios

# --- CARREGAMENTO INICIAL ---
# Agora chamamos a função que criamos para ler os arquivos individuais
if "formularios" not in st.session_state:
    st.session_state["formularios"] = carregar_todos_formularios()
# ============================================================
# LOGIN (Com Bypass para o Formulário)
# ============================================================
# Só bloqueia o acesso se NÃO estiver logado E NÃO for a página de formulário
if not st.session_state.logged_in and st.session_state.pagina != "formulario":
    st.title("🔐 Acesso")
    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar", key="login_button"):
        if (usuario == "admin" and senha == "admin123") or (usuario == "Luciano" and senha == "123"):
            st.session_state.logged_in = True
            st.session_state.user_nome = usuario
            st.session_state.is_admin = True
            
            # ATUALIZAÇÃO: Definimos a variável que o painel de exportação espera
            if usuario == "Luciano":
                st.session_state["usuario_logado"] = "Luciano 123"
            else:
                st.session_state["usuario_logado"] = usuario
                
            st.rerun()
        else:
            st.error("Usuário ou senha incorretos")
    
    st.stop()

# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.title("📌 Menu de Navegação")

btn_home = st.sidebar.button("🏠 Home")
btn_analise = st.sidebar.button("📊 Análise Inteligente")
btn_comparar = st.sidebar.button("⚖️ Comparar Colaboradores")
btn_disc = st.sidebar.button("🧠 Perfil DISC")
btn_parecer = st.sidebar.button("📄 Parecer Estratégico")
btn_visualizar = st.sidebar.button("👁️ Visualizar Dados")
btn_produtividade = st.sidebar.button("🚀 Produtividade")


st.sidebar.markdown("---")

btn_logout = st.sidebar.button("🚪 Logout")

pagina_anterior = st.session_state.pagina

if btn_home:
    st.session_state.pagina = "home"
elif btn_analise:
    st.session_state.pagina = "analise"
elif btn_comparar:
    st.session_state.pagina = "comparar"
elif btn_disc:
    st.session_state.pagina = "disc"
elif btn_parecer:
    st.session_state.pagina = "parecer"
elif btn_visualizar:
    st.session_state.pagina = "visualizar"
# O elif abaixo verifica a URL sem precisar de botão
elif st.session_state.pagina == "formulario":
    pass # Este comando é obrigatório para não dar erro de sintaxe
elif btn_logout:
    st.session_state.logged_in = False
    st.session_state.pagina = "home"

if pagina_anterior != st.session_state.pagina:
    st.rerun()

import streamlit as st
import pandas as pd
import os
import json
import sys

# ============================================================
# CONFIGURAÇÃO DE DIRETÓRIO E CARREGAMENTO
# ============================================================

# Define o diretório base e a pasta de dados
base_dir = os.path.dirname(os.path.abspath(__file__))
dados_dir = os.path.join(base_dir, "dados")
os.makedirs(dados_dir, exist_ok=True)

# Função para carregar todos os JSONs da pasta 'dados'
def carregar_todos_formularios():
    lista_formularios = []
    if os.path.exists(dados_dir):
        for nome_arquivo in os.listdir(dados_dir):
            if nome_arquivo.endswith(".json"):
                caminho_completo = os.path.join(dados_dir, nome_arquivo)
                try:
                    with open(caminho_completo, "r", encoding="utf-8") as f:
                        dados = json.load(f)
                        if isinstance(dados, dict):
                            lista_formularios.append(dados)
                except Exception:
                    continue
    return lista_formularios

# Inicializa o estado da sessão com os dados carregados
if "formularios" not in st.session_state:
    st.session_state["formularios"] = carregar_todos_formularios()

# --- BLOCO DE CSS PARA OCULTAÇÃO ---
if st.query_params.get("page") == "formulario":
    st.markdown("""
    <style>
        [data-testid="stSidebar"] {display: none !important;}
        #MainMenu, footer, header {visibility: hidden !important;}
    </style>
    """, unsafe_allow_html=True)
# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]


# --- FORMULÁRIO ---
if st.query_params.get("page") == "formulario":
    st.title("📋 Formulário Completo do Colaborador")
    
    # Listas padronizadas (devem vir antes do form)
    lista_horas = [f"{i} h" for i in range(25)]
    lista_minutos = [f"{i} min" for i in range(0, 60, 5)]
    lista_frequencia = ["DVD", "D", "S", "Q", "M", "T", "A"]
    
    # ÚNICO BLOCO DO FORMULÁRIO
    with st.form("form_colaborador"):
        # Dados de Identificação
        col1, col2 = st.columns(2)
        nome = col1.text_input("Nome do colaborador")
        setor = col2.text_input("Setor")
        cargo = col1.text_input("Cargo")
        chefe = col2.text_input("Chefe imediato")
        departamento = col1.text_input("Departamento")
        empresa = col2.text_input("Empresa / Unidade")
        escolaridade = col1.text_input("Escolaridade")
        devolucao = col2.text_input("Devolver preenchido em")
        
        cursos = st.text_area("Cursos obrigatórios ou diferenciais")
        objetivo = st.text_area("Trabalho e principal objetivo")
        
        
        
        # --- SEÇÃO DE ATIVIDADES ---
        st.markdown("---")
        
        # Mude para 3 colunas
        col1, col2, col3 = st.columns(3)
        
        # Supondo que você tenha definido col1, col2 e col3 anteriormente
        with col1:
            st.info("""
            **📋 LEGENDA DE FREQUÊNCIA:**
            * **DVD**: Diário Várias Vezes
            * **D**: Diário | **S**: Semanal
            * **Q**: Quinzenal | **M**: Mensal
            * **T**: Trimestral | **A**: Anual
            """)

        with col2:
            st.warning("""
            **⏱️ COMO REGISTRAR O TEMPO:**
            * **Horas e Minutos**: Selecione o valor em cada coluna.
            * **Menos de 1 hora?**: Selecione **0 h** e o tempo real em minutos.
            * **Não se aplica?**: Selecione **0 h** e **0 min** em ambos.
            """)
            
        with col3:
            st.error("""
            **⚠️ DETALHE:**
            * A numeração lateral (nones) é um comportamento nativo do sistema que polui a página.
            * Ignore-a e preencha normalmente; isso não afeta em nada os seus dados.
            """)        
        
        
        
        st.subheader("🔹 Atividades Executadas")
        
        edit_ativ = st.data_editor(
            pd.DataFrame({
                "Atividade Descrita": [""] * 20,
                "Frequência": [""] * 20,
                "Horas": [""] * 20,
                "Minutos": [""] * 20
            }).reset_index(drop=True), # Limpeza do índice
            column_config={
                "Frequência": st.column_config.SelectboxColumn("Frequência", options=lista_frequencia),
                "Horas": st.column_config.SelectboxColumn("Horas", options=lista_horas),
                "Minutos": st.column_config.SelectboxColumn("Minutos", options=lista_minutos),
            },
            hide_index=True,
            num_rows="fixed",
            use_container_width=True
        )

        # --- SEÇÃO DE DIFICULDADES ---
        st.markdown("---")
        st.subheader("⚠️ Dificuldades e Bloqueios")
        
        # 1. INICIALIZAÇÃO SEGURA DO ESTADO (O Cofre)
        if 'df_dificuldades' not in st.session_state:
            st.session_state.df_dificuldades = pd.DataFrame({
                "Dificuldade": [""] * 20,
                "Setor/Parceiro Envolvido": [""] * 20,
                "Horas Perdidas": [""] * 20,
                "Minutos Perdidos": [""] * 20,
                "Frequência": [""] * 20
            })

        

        lista_frequencia = ["DVD", "D", "S", "Q", "M", "T", "A"]

        # 3. EDITOR DE DADOS (Conectado ao Session State)
        # Se estiver dentro de um st.form, os dados só "descem" para o código após o Submit
        edit_dif = st.data_editor(
            st.session_state.df_dificuldades,
            column_config={
                "Horas Perdidas": st.column_config.SelectboxColumn("Horas Perdidas", options=lista_horas),
                "Minutos Perdidos": st.column_config.SelectboxColumn("Minutos Perdidos", options=lista_minutos),
                "Frequência": st.column_config.SelectboxColumn(
                    "Frequência", 
                    options=lista_frequencia,
                    help="Selecione a sigla conforme a legenda acima"
                ),
            },
            hide_index=True,
            num_rows="fixed",
            use_container_width=True,
            key="dif_editor"
        )

        # 4. SALVAMENTO AUTOMÁTICO DO ESTADO
        st.session_state.df_dificuldades = edit_dif


        # Salva o estado para persistir os dados
        st.session_state.df_dificuldades = edit_dif

        # --- SEÇÃO DE SUGESTÕES ATUALIZADA ---
        st.markdown("---")
        st.subheader("💡 Sugestões de Melhoria e Impacto")
        
        edit_sug = st.data_editor(
            pd.DataFrame({
                "Sugestão de Melhoria": [""] * 20,
                "Impacto Esperado": [""] * 20,
                "Redução Horas": [""] * 20,
                "Redução Minutos": [""] * 20,
                "Frequência do Impacto": [""] * 20
            }).reset_index(drop=True),
            column_config={
                "Redução Horas": st.column_config.SelectboxColumn(
                    "Redução Horas", 
                    options=lista_horas
                ),
                "Redução Minutos": st.column_config.SelectboxColumn(
                    "Redução Minutos", 
                    options=lista_minutos
                ),
                "Frequência do Impacto": st.column_config.SelectboxColumn(
                    "Frequência do Impacto", 
                    options=lista_frequencia
                ),
            },
            hide_index=True,
            num_rows="fixed",
            use_container_width=True,
            key="sug_editor"
        )

        

        

        st.markdown("---")
        st.subheader("📊 Questionário DISC")
        for i, pergunta in enumerate(perguntas_disc, 1):
            st.radio(
                label=f"{i}. {pergunta}", 
                options=["A", "B", "C", "D"], 
                key=f"disc_{i}", 
                horizontal=True, 
                index=None
            )
        # BOTÃO DO FORMULÁRIO
        enviar = st.form_submit_button("🚀 ENVIAR FORMULÁRIO FINAL")
          
        # -------------------------------------------------
        # VALIDAÇÕES E PROCESSAMENTO
        # -------------------------------------------------
        
        if enviar:

            # 1. GERA A DATA E HORA DE BRASÍLIA
            fuso_brasilia = pytz.timezone('America/Sao_Paulo')
            data_hoje = datetime.now(fuso_brasilia).strftime('%d/%m/%Y %H:%M:%S')
            # Criamos uma lista com todos os campos que NÃO podem estar vazios
            # Verifique se os nomes das variáveis (cursos, trabalho, objetivo) são esses mesmos
            campos_obrigatorios = [
                nome, setor, cargo, chefe, departamento, empresa,
                cursos, objetivo
            ]

            # 1. VALIDAÇÃO DE CAMPOS (Identificação + Cursos/Trabalho/Objetivo)
            # O strip() remove espaços vazios para garantir que haja texto real
            if any(not str(campo).strip() for campo in campos_obrigatorios):
                st.error("⚠️ Erro: Preencha todos os campos obrigatórios!")

            # 2. VALIDAÇÃO DO DISC
            elif any(st.session_state.get(f"disc_{i}") is None for i in range(1, 25)):
                st.error("⚠️ Erro: Responda todas as perguntas do DISC!")

            else:
                import os
                import json

                base_dir = os.path.dirname(os.path.abspath(__file__))
                dados_dir = os.path.join(base_dir, "dados")
                os.makedirs(dados_dir, exist_ok=True)

                # 3. EVITAR DUPLICIDADE
                nome_limpo = nome.strip().replace(" ", "_")
                arquivos_existentes = [f for f in os.listdir(dados_dir) if f.startswith(nome_limpo)]

                if arquivos_existentes:
                    st.error(f"⚠️ Já existe um formulário enviado para '{nome}'.")

                else:

                    # 4. CONFIRMAÇÃO
                    if not st.session_state.get("confirmado", False):

                        st.warning(
                            "⚠️ Revise o formulário. Clique novamente no botão para confirmar o envio."
                        )

                        st.session_state["confirmado"] = True

                    else:
                        st.success("✅ Formulário enviado com sucesso!")

                        dados = {
                            "nome": nome,
                            "data_envio": data_hoje,
                            "setor": setor,
                            "cargo": cargo,
                            "chefe": chefe,
                            "departamento": departamento,
                            "empresa": empresa,
                            "escolaridade": escolaridade,
                            "devolucao": devolucao,
                            "cursos_obrigatorios_ou_diferenciais": cursos,
                            "trabalho_e_principal_objetivo": objetivo,
                            "atividades": edit_ativ.to_dict() if hasattr(edit_ativ, 'to_dict') else edit_ativ,
                            "dificuldades": edit_dif.to_dict() if hasattr(edit_dif, 'to_dict') else edit_dif,
                            "sugestoes": edit_sug.to_dict() if hasattr(edit_sug, 'to_dict') else edit_sug,
                            "disc": {
                                f"disc_{i}": st.session_state.get(f"disc_{i}")
                                for i in range(1, 25)
                            }
                        }
                        
                        # Reseta a confirmação para um próximo envio
                        st.session_state["confirmado"] = False

                        caminho = os.path.join(dados_dir, f"{nome_limpo}.json")

                        with open(caminho, "w", encoding="utf-8") as f:
                            json.dump(dados, f, ensure_ascii=False, indent=4)

                        st.session_state["confirmado"] = False


# --- VISUALIZAÇÃO ---
if st.session_state.get("pagina") == "visualizar":
    st.title("👁️ Visualização de Registros")
    
    # 1. Carrega os dados frescos do disco
    lista_de_arquivos = carregar_todos_formularios()
    
    # 2. Se a sua função carregar_todos_formularios() já retorna a lista, 
    # apenas certifique-se de que não estamos adicionando isso ao session_state de forma acumulativa.
    if not lista_de_arquivos:
        st.warning("⚠️ Nenhum formulário encontrado.")
    else:
        # Mostra o total para conferência
        st.success(f"Foram encontrados {len(lista_de_arquivos)} formulários.")
        
        # 3. Exibição limpa
        for idx, form in enumerate(lista_de_arquivos, 1):
            nome_exibir = str(form.get('Nome', f'Colaborador {idx}')).upper()
            
            with st.expander(f"👤 FORMULÁRIO DE: {nome_exibir} ({form.get('DataEnvio', 'Sem Data')})"):
                # [Aqui você mantém o seu código de exibição de dados]
                
            
            
            
                # 1. Cabeçalho Completo
                st.subheader("📝 Informações de Identificação")
                col1, col2 = st.columns(2)
                
                # Buscando as chaves minúsculas do seu dicionário 'dados'
                col1.write(f"**Data de Envio:** {form.get('data_envio', 'N/A')}")
                col2.write(f"**Devolver em:** {form.get('devolucao', 'N/A')}")
                
                col_a, col_b = st.columns(2)
                col_a.write(f"**Setor:** {form.get('setor', 'N/A')}")
                col_b.write(f"**Departamento:** {form.get('departamento', 'N/A')}")
                col_a.write(f"**Cargo:** {form.get('cargo', 'N/A')}")
                col_b.write(f"**Chefe Imediato:** {form.get('chefe', 'N/A')}")
                col_a.write(f"**Empresa/Unidade:** {form.get('empresa', 'N/A')}")
                col_b.write(f"**Escolaridade:** {form.get('escolaridade', 'N/A')}")
                
                # Buscando os nomes longos e específicos que definimos
                st.text_area("Cursos obrigatórios ou diferenciais", value=form.get('cursos_obrigatorios_ou_diferenciais', 'N/A'), disabled=True, key=f"cursos_{idx}")
                st.text_area("Trabalho e principal objetivo", value=form.get('trabalho_e_principal_objetivo', 'N/A'), height=150, disabled=True, key=f"obj_{idx}")
                
                # 2. Tabelas Dinâmicas
                secoes = {
                    "atividades": "📋 Atividades Executadas",
                    "dificuldades": "⚠️ Dificuldades e Bloqueios",
                    "sugestoes": "💡 Sugestões de Melhoria"
                }
                
                for chave, titulo in secoes.items():
                    st.markdown("---")
                    st.subheader(titulo)
                    if chave in form and form[chave]:
                        df = pd.DataFrame(form[chave])
                        df = df.replace("", None).dropna(how='all')
                        if not df.empty:
                            st.table(df)
                        else:
                            st.write("Nenhum dado preenchido nesta seção.")
                    else:
                        st.write("Seção não encontrada ou vazia.")
                
                # 3. Questionário DISC
                st.markdown("---")
                st.subheader("📊 Avaliação DISC (Perguntas e Respostas)")
                
                dados_disc = form.get("disc", {}) # Entra na "pasta" disc do JSON
                
                for i, pergunta in enumerate(perguntas_disc, 1):
                    valor_resposta = dados_disc.get(f"disc_{i}", "Não respondido")
                    st.write(f"**{i}. {pergunta}**")
                    st.info(f"Resposta selecionada: **{valor_resposta}**")
                    st.markdown("---")

                # --- BLOCO DE EXPORTAÇÃO (SÓ WORD E PDF) ---
                if st.session_state.get("usuario_logado") == "Luciano 123":
                    st.markdown("---")
                    st.subheader("⚙️ Painel de Exportação")
                    
                    # Usamos 2 colunas para ficar mais harmônico
                    col1, col2 = st.columns(2)
                    
                    # Padronização do nome do arquivo para ambos
                    data_clean = form.get('DataEnvio', '').replace('/', '').replace(' ', '_').replace(':', '')
                    nome_clean = form.get('Nome', 'Colaborador').replace(' ', '_')
                    nome_arquivo = f"Relatorio_{nome_clean}_{data_clean}"
                    
                    with col1:
                        st.download_button(
                            label="📄 Baixar em Word",
                            data=gerar_word(form),
                            file_name=f"{nome_arquivo}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    with col2:
                        st.download_button(
                            label="📑 Baixar em PDF",
                            data=gerar_pdf(form),
                            file_name=f"{nome_arquivo}.pdf",
                            mime="application/pdf"
                        )
                # --- FIM DO BLOCO ---



        # Botão de Limpeza
        st.markdown("---")
        if st.button("🗑️ LIMPAR TODOS OS FORMULÁRIOS"):
            for arquivo in os.listdir(dados_dir):
                if arquivo.endswith(".json"): 
                    os.remove(os.path.join(dados_dir, arquivo))
            st.session_state["formularios"] = []
            st.success("✅ Banco de dados limpo!"); st.rerun()

# ============================================================
# CALCULAR DISC PERCENTUAL E DOMINANTE
# ============================================================

def calcular_disc(respostas_disc):
    contagem = {"D":0, "I":0, "S":0, "C":0}
    for r in respostas_disc.values():
        if r in contagem:
            contagem[r] += 1
    total = sum(contagem.values())
    if total > 0:
        percentuais = {k: round(v/total*100,1) for k,v in contagem.items()}
        dominante = max(percentuais, key=percentuais.get)
    else:
        percentuais = contagem
        dominante = None
    return percentuais, dominante

# ============================================================
# SCORE DISC PONDERADO
# ============================================================

def score_disc(disc):
    pesos = {"D":1.0,"I":0.9,"S":0.85,"C":0.95}
    total = sum(disc.values())
    if total == 0:
        return 0
    calculo = sum(disc[k]*pesos.get(k,1) for k in disc)
    return round((calculo/total)*100,2)

# ============================================================
# CALCULAR CARGA HORÁRIA
# ============================================================

def calcular_carga(atividades):
    total_min = 0
    for at in atividades:
        try:
            tempo = float(at.get("tempo","0"))
        except:
            tempo = 0
        freq = at.get("frequencia","semanal").lower()
        if freq == "diaria":
            total_min += tempo * 5
        elif freq == "mensal":
            total_min += tempo / 4
        else:
            total_min += tempo
    horas = total_min / 60
    status = "Adequado"
    if horas > 44: status = "Sobrecarga"
    elif horas < 30: status = "Subutilização"
    return round(horas,2), status

# ============================================================
# GERAR ATIVIDADES IDEAIS (GPT)
# ============================================================

def gerar_atividades_ideais(cargo, setor, client=None):
    if client is None:
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]
    
    prompt = f"""
    Gere 12 atividades ideais para:
    Cargo: {cargo}
    Setor: {setor}
    Para cada atividade informe:
      - nome_atividade
      - descricao
      - frequencia_ideal
      - tempo_medio_minutos
      - justificativa_tecnica
    Responda SOMENTE JSON válido.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.3
        )
        return json.loads(response.choices[0].message.content)
    except:
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]

# ============================================================
# COMPARAÇÃO SEMÂNTICA
# ============================================================

def comparar_semanticamente(reais, ideais, client=None):
    if client is None:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

    prompt = f"""
    Compare semanticamente:
    Atividades reais: {reais}
    Atividades ideais: {ideais}
    Retorne JSON com:
      - score_aderencia (0-100)
      - tempo_gap_medio_percentual
      - atividades_desvio
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

# ============================================================
# CLASSIFICAR DIFICULDADES
# ============================================================

def classificar_dificuldades_gpt(dificuldades, client=None):
    if client is None:
        return {}
    
    prompt = f"""
    Classifique semanticamente as dificuldades abaixo em:
    - Processo
    - Tempo
    - Comunicação
    - Estrutura
    - Liderança
    - Sistema
    Retorne JSON com contagem por categoria.
    Dificuldades: {dificuldades}
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {}

# ============================================================
# ÍNDICE GERAL DO CARGO
# ============================================================

def indice_geral(score_aderencia, score_disc, status_carga):
    fator_carga = 100
    if status_carga == "Sobrecarga": fator_carga = 70
    elif status_carga == "Subutilização": fator_carga = 75
    return round(mean([score_aderencia, score_disc, fator_carga]),2)

# ============================================================
# MOTOR PRINCIPAL COMPLETO – ANÁLISE CORPORATIVA
# ============================================================

def gerar_analise_corporativa(dados, client=None):
    """
    Gera análise completa de um colaborador com base em:
    - Atividades reais
    - Perfil DISC
    - Dificuldades
    Retorna:
    - parecer (texto)
    - indicadores (dict)
    """
    # 1️⃣ Atividades ideais
    ideais = gerar_atividades_ideais(dados["cargo"], dados["setor"], client)

    # 2️⃣ Comparação semântica (reais x ideais)
    comparacao = comparar_semanticamente(dados["atividades"], ideais, client)

    # 3️⃣ Carga horária
    horas, status_carga = calcular_carga(dados["atividades"])

    # 4️⃣ Score DISC
    disc_score = score_disc(dados["disc"])

    # 5️⃣ Classificação de dificuldades
    dificuldades_classificadas = classificar_dificuldades_gpt(dados["dificuldades"], client)

    # 6️⃣ Score de aderência
    score_aderencia = comparacao.get("score_aderencia",0)

    # 7️⃣ Índice geral
    indice = indice_geral(score_aderencia, disc_score, status_carga)

    # 8️⃣ Classificação de risco
    risco = "Baixo" if indice < 60 else "Moderado" if indice < 75 else "Alto"

    # 9️⃣ Prompt final para parecer estratégico
    prompt_final = f"""
    Gere parecer estratégico completo considerando:
    - Score aderência: {score_aderencia}
    - Horas semanais: {horas}
    - Status carga: {status_carga}
    - Score DISC: {disc_score}
    - Dificuldades: {dificuldades_classificadas}
    - Índice geral do cargo: {indice}
    - Classificação de risco: {risco}
    
    Inclua:
    - Diagnóstico estrutural
    - Análise de desvios
    - Avaliação comportamental
    - Riscos organizacionais
    - Recomendação detalhada de redistribuição
    - Atividades corretas para o cargo com tempo e frequência ideais
    - Conclusão executiva
    """

    # 10️⃣ Obter parecer do GPT
    parecer = ""
    try:
        if client:
            resposta = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt_final}],
                temperature=0.3
            )
            parecer = resposta.choices[0].message.content
        else:
            parecer = "GPT não disponível. Retorno padrão: análise resumida."
    except:
        parecer = "Erro ao gerar parecer com GPT."

    # 11️⃣ Indicadores
    indicadores = {
        "score_aderencia": score_aderencia,
        "horas_semanais": horas,
        "status_carga": status_carga,
        "score_disc": disc_score,
        "indice_geral": indice,
        "risco": risco
    }

    return parecer, indicadores

# ============================================================
# GERAR PDF DO PARECER
# ============================================================

def gerar_pdf(parecer, nome):
    """
    Recebe:
    - parecer (texto)
    - nome do colaborador
    Cria arquivo PDF pronto para download
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    nome_arquivo = f"{nome}_parecer.pdf"
    doc = SimpleDocTemplate(nome_arquivo)
    elements = []
    styles = getSampleStyleSheet()

    # Título
    elements.append(Paragraph("PARECER ESTRATÉGICO ORGANIZACIONAL", styles["Title"]))
    elements.append(Spacer(1, 0.5*inch))

    # Conteúdo linha a linha
    for linha in parecer.split("\n"):
        if linha.strip():  # evita parágrafos vazios
            elements.append(Paragraph(linha, styles["Normal"]))
            elements.append(Spacer(1, 0.2*inch))

    doc.build(elements)
    return nome_arquivo

# ============================================================
# PASTA BASE PARA FORMULÁRIOS (JSON)
# ============================================================
# Usamos 'dados_dir' para manter o padrão que já criamos
json_master = os.path.join(dados_dir, "formularios.json")

# Inicializa arquivo JSON se não existir
if not os.path.exists(json_master):
    with open(json_master, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=4)

# ============================================================
# FUNÇÃO PARA SALVAR FORMULÁRIO EM JSON
# ============================================================
def salvar_formulario_json(formulario):
    """
    Recebe um dicionário do formulário preenchido, salva no arquivo 
    JSON único dentro da pasta 'dados' e atualiza a sessão para 
    espelhamento imediato na interface.
    """
    # 1. Tenta carregar os dados existentes ou cria uma lista vazia
    try:
        with open(json_master, "r", encoding="utf-8") as f:
            dados_existentes = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        # Se o arquivo não existir ou estiver vazio/inválido, inicia como lista vazia
        dados_existentes = []

    # 2. Adiciona o novo formulário à lista
    dados_existentes.append(formulario)

    # 3. Salva a lista completa de volta no arquivo
    with open(json_master, "w", encoding="utf-8") as f:
        json.dump(dados_existentes, f, ensure_ascii=False, indent=4)

    # 4. Atualiza o estado da sessão do Streamlit para refletir a mudança instantaneamente
    st.session_state["formularios"] = dados_existentes


